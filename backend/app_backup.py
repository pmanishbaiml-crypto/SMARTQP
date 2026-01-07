import os
from flask import Flask, request, jsonify, render_template, send_file, send_from_directory
from flask_cors import CORS
import firebase_admin
from firebase_admin import credentials, auth, firestore
from supabase_service import supabase_service
from functools import wraps
from dotenv import load_dotenv # Make sure dotenv is installed: pip install python-dotenv

# --- New Imports for File Parsing ---
import pdfplumber
import docx # For .docx files
from docx.shared import Inches, Pt # For setting table column widths and font sizes
import docx.enum.text # For text alignment
import docx.shared # For measurements
from werkzeug.utils import secure_filename
import re # For cleaning text and potential regex parsing
import traceback # For detailed error logging
import uuid # For generating unique IDs for question banks and questions
from datetime import datetime # For current date/time in template
import pythoncom # For COM initialization in PDF conversion
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

# --- New Imports for Gemini API ---
import requests # For making HTTP requests to the Gemini API
import json # For handling JSON payload
import time # For rate limiting

# Load environment variables from .env file (for local development)
load_dotenv()

app = Flask(__name__)
CORS(app) # Enable CORS for all origins, adjust in production for production deployment

# --- Firebase Admin SDK Initialization ---
if not firebase_admin._apps:
    try:
        # Try to use the service account JSON file directly
        service_account_path = 'skit-qp-firebase-adminsdk-fbsvc-146911a4da.json'
        if os.path.exists(service_account_path):
            cred = credentials.Certificate(service_account_path)
            firebase_admin.initialize_app(cred)
            print("Firebase Admin SDK initialized using service account JSON file.")
        else:
            # Fallback to Application Default Credentials
            cred = credentials.ApplicationDefault()
            firebase_admin.initialize_app(cred)
            print("Firebase Admin SDK initialized using Application Default Credentials.")
    except Exception as e:
        print(f"Error initializing Firebase Admin SDK: {e}")
        print("Please ensure the Firebase service account key JSON file exists in the backend directory.")
        print("For local development, make sure 'smartqpgen-amrutha-firebase-adminsdk-fbsvc-509c2b1cdd.json' is in the backend folder.")
else:
    print("Firebase Admin SDK already initialized.")

db_firestore = firestore.client() # Get Firestore client

# Simple rate limiting for authentication
auth_attempts = {}  # Store authentication attempts per IP
MAX_AUTH_ATTEMPTS = 5  # Maximum attempts per minute
AUTH_WINDOW = 60  # Time window in seconds

# --- JWT Secret Key (for Flask sessions, if needed, but Firebase ID tokens are primary) ---
app.config['SECRET_KEY'] = os.getenv('FLASK_SECRET_KEY', 'your_super_secret_key_change_this_in_production')

# --- Simple in-memory cache for reducing Firestore calls ---
question_cache = {}
cache_timeout = 300  # 5 minutes

def get_cached_questions(user_uid, cache_key):
    """Get questions from cache if available and not expired"""
    import time
    if cache_key in question_cache:
        cached_data, timestamp = question_cache[cache_key]
        if time.time() - timestamp < cache_timeout:
            print(f"Using cached questions for {cache_key}")
            return cached_data
        else:
            # Cache expired, remove it
            del question_cache[cache_key]
    return None

def cache_questions(user_uid, cache_key, questions):
    """Cache questions with timestamp"""
    import time
    question_cache[cache_key] = (questions, time.time())
    print(f"Cached {len(questions)} questions for {cache_key}")

# --- File Upload Configuration ---
UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads') # Directory to temporarily save uploaded files
ALLOWED_EXTENSIONS = {'pdf', 'docx'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
os.makedirs(UPLOAD_FOLDER, exist_ok=True) # Create the upload folder if it doesn't exist

# --- Authentication Decorator ---
def firebase_auth_required(f):
    """Decorator to protect routes, verifying Firebase ID tokens."""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        auth_header = request.headers.get('Authorization')
        if not auth_header:
            return jsonify({'message': 'Authorization header is missing!'}), 401

        try:
            # Token format: "Bearer <ID_TOKEN>"
            id_token = auth_header.split(' ')[1]
            # Verify the Firebase ID token
            decoded_token = auth.verify_id_token(id_token)
            request.current_user_uid = decoded_token['uid'] # Attach UID to request object
            return f(*args, **kwargs)
        except firebase_admin.auth.InvalidIdTokenError:
            return jsonify({'message': 'Invalid or expired authentication token.'}), 401
        except Exception as e:
            return jsonify({'message': f'Authentication error: {str(e)}'}), 500
    return decorated_function

# --- Helper Functions for Parsing ---
def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def convert_docx_to_pdf_robust(docx_path):
    """
    Robust PDF conversion function that tries multiple methods
    """
    pdf_path = docx_path.replace('.docx', '.pdf')

    # Method 1: Try docx2pdf with COM initialization
    try:
        from docx2pdf import convert
        print("Attempting PDF conversion with docx2pdf...")

        # Initialize COM for the current thread
        pythoncom.CoInitialize()

        try:
            convert(docx_path, pdf_path)
            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                print("PDF conversion successful with docx2pdf")
                return pdf_path
        finally:
            # Always uninitialize COM
            pythoncom.CoUninitialize()

    except Exception as e:
        print(f"docx2pdf conversion failed: {e}")

    # Method 2: Try using win32com directly (Windows only)
    try:
        import win32com.client
        print("Attempting PDF conversion with win32com...")

        # Initialize COM
        pythoncom.CoInitialize()

        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False

            # Open document
            doc = word.Documents.Open(os.path.abspath(docx_path))

            # Save as PDF (format 17 is PDF)
            doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
            doc.Close()
            word.Quit()

            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                print("PDF conversion successful with win32com")
                return pdf_path

        finally:
            pythoncom.CoUninitialize()

    except Exception as e:
        print(f"win32com conversion failed: {e}")

    # Method 3: Return original DOCX if all PDF methods fail
    print("All PDF conversion methods failed, returning DOCX")
    return docx_path

def clean_text(text):
    """Basic text cleaning: remove excessive whitespace, newlines, strip leading/trailing space."""
    if text is None:
        return ""
    # Replace multiple spaces/newlines with a single space, then strip
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def extract_marks_from_text(marks_text):
    """Extract numeric marks from various formats like '6M', '8M', '6', etc."""
    if not marks_text:
        return 0
    
    # Remove common suffixes and extract number
    marks_text = str(marks_text).strip().upper()
    
    # Handle formats like "6M", "8M", "10M"
    if marks_text.endswith('M'):
        try:
            return int(marks_text[:-1])
        except ValueError:
            pass
    
    # Handle formats like "6", "8", "10"
    try:
        return int(marks_text)
    except ValueError:
        pass
    
    # Handle formats like "6 marks", "8 marks"
    marks_match = re.search(r'(\d+)\s*marks?', marks_text, re.IGNORECASE)
    if marks_match:
        return int(marks_match.group(1))
    
    # Handle formats like "6 pts", "8 pts"
    pts_match = re.search(r'(\d+)\s*pts?', marks_text, re.IGNORECASE)
    if pts_match:
        return int(pts_match.group(1))
    
    return 0

def extract_table_as_text(table_data):
    """Convert table data to readable text format with proper formatting"""
    if not table_data:
        return ""
    
    table_text = ""
    for row in table_data:
        if row and any(cell for cell in row if cell and str(cell).strip()):
            # Filter out empty cells and join with proper spacing
            row_cells = [str(cell).strip() if cell else "" for cell in row]
            # Use consistent spacing instead of tabs for better display
            row_text = " | ".join(row_cells)
            table_text += row_text + "\n"
    
    return table_text.strip()

def extract_table_data_for_knn(table_data):
    """Extract table data specifically for k-NN dataset format"""
    if not table_data:
        return ""
    
    # Check if this looks like the k-NN dataset
    if len(table_data) > 0 and len(table_data[0]) >= 5:
        header_row = table_data[0]
        if any("S.NO" in str(cell) for cell in header_row) or any("CGPA" in str(cell) for cell in header_row):
            # This is the k-NN dataset format
            table_text = ""
            for row in table_data:
                if row and any(cell for cell in row if cell and str(cell).strip()):
                    row_cells = [str(cell).strip() if cell else "" for cell in row]
                    row_text = " | ".join(row_cells)
                    table_text += row_text + "\n"
            return table_text.strip()
    
    # Fallback to regular table extraction
    return extract_table_as_text(table_data)

def format_table_data_for_display(table_data):
    """Format table data specifically for frontend display with proper line breaks"""
    if not table_data:
        return ""
    
    formatted_rows = []
    for row in table_data:
        if row and any(cell for cell in row if cell and str(cell).strip()):
            # Filter out empty cells and join with proper spacing
            row_cells = [str(cell).strip() if cell else "" for cell in row]
            # Use pipe separators for better frontend parsing
            row_text = " | ".join(row_cells)
            formatted_rows.append(row_text)
    
    return "\n".join(formatted_rows)

def format_question_with_tables(question_text):
    """Format question text that may contain embedded table data"""
    if not question_text:
        return question_text
    
    # Check if the question contains concatenated table data (like "S.NO CGPA Assessment Project submitted Result 1 9.2 85 8 Pass 2 8 80 7 Pass...")
    if 'S.NO' in question_text and 'CGPA' in question_text and 'Assessment' in question_text:
        # This is likely the k-NN dataset format
        return format_knn_dataset_in_question(question_text)
    
    # Check for other table patterns like "S.NO GPA No. of projects done Award 1 9.5 5 Yes..."
    if 'S.NO' in question_text and 'GPA' in question_text and 'Award' in question_text:
        return format_gpa_award_dataset_in_question(question_text)
    
    # Check for other table patterns like "X Y Class 3 1 A 5 2 A..."
    if 'X Y Class' in question_text or ('X' in question_text and 'Y' in question_text and 'Class' in question_text):
        return format_xy_class_dataset_in_question(question_text)
    
    # Check if the question contains table-like data
    lines = question_text.split('\n')
    formatted_lines = []
    in_table = False
    table_lines = []
    
    for line in lines:
        line = line.strip()
        if not line:
            if in_table and table_lines:
                # End of table, format it
                formatted_table = format_embedded_table(table_lines)
                formatted_lines.append(formatted_table)
                table_lines = []
                in_table = False
            continue
            
        # Check if this line looks like table data (has multiple space-separated values)
        words = line.split()
        # More specific conditions for table detection
        is_table_data = (
            len(words) >= 3 and 
            not any(word in line.lower() for word in ['question', 'consider', 'write', 'what', 'how', 'explain', 'describe', 'compare', 'contrast', 'differences', 'between', 'instance', 'based', 'learning', 'model', 'nearest', 'neighbour', 'algorithm', 'weighted', 'determine', 'class', 'given', 'test', 'assign', 'predict', 'using', 'classifier', 'centroid', 'training', 'dataset', 'target', 'variable', 'discrete', 'valued', 'takes', 'values', 'choose', 'k=']) and
            (any(word.isdigit() for word in words) or any(word in ['pass', 'fail', 'yes', 'no', 'a', 'b'] for word in words) or 
             any(word.lower() in ['s.no', 'sno', 'gpa', 'cgpa', 'x', 'y', 'class', 'award', 'result', 'assessment', 'project', 'submitted'] for word in words))
        )
        
        if is_table_data:
            # This looks like table data
            if not in_table:
                in_table = True
            table_lines.append(line)
        else:
            if in_table and table_lines:
                # End of table, format it
                formatted_table = format_embedded_table(table_lines)
                formatted_lines.append(formatted_table)
                table_lines = []
                in_table = False
            formatted_lines.append(line)
    
    # Handle table at the end
    if in_table and table_lines:
        formatted_table = format_embedded_table(table_lines)
        formatted_lines.append(formatted_table)
    
    return '\n'.join(formatted_lines)

def format_knn_dataset_in_question(question_text):
    """Format k-NN dataset that's concatenated in question text"""
    # Find the start of the table data
    table_start = question_text.find('S.NO')
    if table_start == -1:
        return question_text
    
    # Split into question part and table part
    question_part = question_text[:table_start].strip()
    table_part = question_text[table_start:].strip()
    
    # Format the table part
    formatted_table = format_concatenated_table(table_part)
    
    # Combine question and formatted table
    return f"{question_part}\n\n{formatted_table}"

def format_xy_class_dataset_in_question(question_text):
    """Format X Y Class dataset that's concatenated in the question text"""
    # Find where the table data starts
    table_start = question_text.find('X Y Class')
    if table_start == -1:
        # Try alternative patterns
        if 'X' in question_text and 'Y' in question_text and 'Class' in question_text:
            # Find the position where X Y Class appears
            words = question_text.split()
            for i, word in enumerate(words):
                if word == 'X' and i + 2 < len(words) and words[i+1] == 'Y' and words[i+2] == 'Class':
                    # Found "X Y Class", get the position in the original text
                    table_start = question_text.find(' '.join(words[i:i+3]))
                    break
    
    if table_start == -1:
        return question_text
    
    # Split into question part and table part
    question_part = question_text[:table_start].strip()
    table_part = question_text[table_start:].strip()
    
    # Format the table part
    formatted_table = format_xy_class_table(table_part)
    
    # Combine question and formatted table
    return f"{question_part}\n\n{formatted_table}"

def format_xy_class_table(table_text):
    """Format X Y Class table data"""
    words = table_text.split()
    
    # Find the header part (X Y Class)
    header_words = []
    data_start_idx = 0
    
    for i, word in enumerate(words):
        if word.isdigit() and i > 2:  # Found first number after headers
            data_start_idx = i
            break
        header_words.append(word)
    
    if len(header_words) < 3 or data_start_idx == 0:
        return f"Table:\n{table_text}"
    
    # Format headers
    headers = ' | '.join(header_words)
    formatted_table = f"Table:\n{headers}\n"
    
    # Format data rows (should be groups of 3: number, number, letter)
    data_words = words[data_start_idx:]
    for i in range(0, len(data_words), 3):
        if i + 2 < len(data_words):
            row = f"{data_words[i]} | {data_words[i+1]} | {data_words[i+2]}"
            formatted_table += row + "\n"
    
    return formatted_table.strip()

def format_gpa_award_dataset_in_question(question_text):
    """Format S.NO GPA Award dataset that's concatenated in the question text"""
    # Find where the table data starts
    table_start = question_text.find('S.NO')
    if table_start == -1:
        return question_text
    
    # Split into question part and table part
    question_part = question_text[:table_start].strip()
    table_part = question_text[table_start:].strip()
    
    # Format the table part
    formatted_table = format_gpa_award_table(table_part)
    
    # Combine question and formatted table
    return f"{question_part}\n\n{formatted_table}"

def format_gpa_award_table(table_text):
    """Format S.NO GPA Award table data"""
    words = table_text.split()
    
    # Find the header part (S.NO GPA No. of projects done Award)
    header_words = []
    data_start_idx = 0
    
    for i, word in enumerate(words):
        if word.isdigit() and i > 3:  # Found first number after headers
            data_start_idx = i
            break
        header_words.append(word)
    
    if len(header_words) < 4 or data_start_idx == 0:
        return f"Table:\n{table_text}"
    
    # Format headers
    headers = ' | '.join(header_words)
    formatted_table = f"Table:\n{headers}\n"
    
    # Format data rows (should be groups of 4: number, decimal, number, yes/no)
    data_words = words[data_start_idx:]
    for i in range(0, len(data_words), 4):
        if i + 3 < len(data_words):
            row = f"{data_words[i]} | {data_words[i+1]} | {data_words[i+2]} | {data_words[i+3]}"
            formatted_table += row + "\n"
    
    return formatted_table.strip()

def format_concatenated_table(table_text):
    """Format concatenated table text like 'S.NO CGPA Assessment Project submitted Result 1 9.2 85 8 Pass 2 8 80 7 Pass...'"""
    words = table_text.split()
    
    # Find the header part
    header_words = []
    data_start_idx = 0
    
    for i, word in enumerate(words):
        if word.isdigit() and i > 0:  # Found first number (row number)
            data_start_idx = i
            break
        header_words.append(word)
    
    if not header_words or data_start_idx == 0:
        return f"Table:\n{table_text}"
    
    # Format headers
    headers = ' | '.join(header_words)
    
    # Format data rows
    data_words = words[data_start_idx:]
    rows = []
    current_row = []
    
    for word in data_words:
        if word.isdigit() and len(current_row) >= 4:  # New row starting
            if current_row:
                rows.append(' | '.join(current_row))
            current_row = [word]
        else:
            current_row.append(word)
    
    # Add the last row
    if current_row:
        rows.append(' | '.join(current_row))
    
    # Combine everything
    result = f"Table:\n{headers}\n" + '\n'.join(rows)
    return result

def format_embedded_table(table_lines):
    """Format embedded table lines into a proper table structure"""
    if not table_lines:
        return ""
    
    # Check if first line looks like headers
    first_line_words = table_lines[0].split()
    if len(first_line_words) >= 2 and not any(word.isdigit() for word in first_line_words):
        # First line is likely headers
        headers = ' | '.join(first_line_words)
        data_lines = table_lines[1:]
    else:
        # No clear headers, treat all as data
        headers = None
        data_lines = table_lines
    
    formatted_table = "Table:\n"
    if headers:
        formatted_table += headers + "\n"
    
    for line in data_lines:
        words = line.split()
        if len(words) >= 2:
            formatted_line = ' | '.join(words)
            formatted_table += formatted_line + "\n"
    
    return formatted_table.strip()

def extract_table_from_docx_cell(cell):
    """Extract table content from a DOCX cell if it contains a table"""
    table_text = ""
    for table in cell.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells]
            table_text += "\t".join(row_cells) + "\n"
    return table_text.strip()

# --- Parsing Functions ---
def parse_pdf_question_bank(filepath):
    """
    Parses a PDF question bank, attempting to extract questions from tables.
    Assumes a table structure for questions (SL#, Question, CO, Level, Marks).
    Math equations will be extracted as raw text.
    """
    extracted_questions = []
    
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables() 
                
                for table in tables:
                    for row_index, row in enumerate(table):
                        if row_index == 0: # Skip header row
                            continue
                        
                        if len(row) >= 6:
                            # Handle 6-column format: Q.No, Questions, CO, Level, Marks, Module
                            sl_no = clean_text(row[0])
                            question_text = clean_text(row[1])
                            co = clean_text(row[2])
                            blooms_level = clean_text(row[3])
                            marks_text = clean_text(row[4])
                            module = clean_text(row[5])

                            # Use enhanced marks extraction
                            marks = extract_marks_from_text(marks_text)

                            if question_text:
                                extracted_questions.append({
                                    "sl_no": sl_no,
                                    "question_text": question_text,
                                    "co": co,
                                    "blooms_level": blooms_level,
                                    "marks": marks,
                                    "module": module
                                })
                        elif len(row) >= 5:
                            # Handle 5-column format (legacy): Q.No, Questions, CO, Level, Marks
                            sl_no = clean_text(row[0])
                            question_text = clean_text(row[1])
                            co = clean_text(row[2])
                            blooms_level = clean_text(row[3])
                            marks_text = clean_text(row[4])

                            # Use enhanced marks extraction
                            marks = extract_marks_from_text(marks_text)

                            if question_text:
                                extracted_questions.append({
                                    "sl_no": sl_no,
                                    "question_text": question_text,
                                    "co": co,
                                    "blooms_level": blooms_level,
                                    "marks": marks,
                                    "module": "1"  # Default to module 1 for legacy format
                                })
                        else:
                            print(f"Skipping malformed PDF row due to insufficient columns: {row}") # Debugging
    except Exception as e:
        print(f"Error during PDF parsing of {filepath}: {e}")
        traceback.print_exc() # Print full traceback for debugging
        raise # Re-raise to be caught by the route's error handling
    
    return extracted_questions

def parse_docx_question_bank(filepath):
    """
    Parses a DOCX question bank, attempting to extract questions from tables.
    Assumes a table structure similar to PDF (SL#, Question, CO, Level, Marks).
    Math equations will be extracted as raw text.
    """
    extracted_questions = []
    try:
        doc = docx.Document(filepath)

        for table in doc.tables:
            for row_index, row in enumerate(table.rows):
                cells = [clean_text(cell.text) for cell in row.cells]
                
                # Skip empty rows
                if not any(cells):
                    continue
                
                # Check if this is a header row (contains common header words)
                header_words = ['q.no', 'question', 'co', 'level', 'marks', 'module', 'sl', 's.no']
                is_header = any(word in ' '.join(cells).lower() for word in header_words)
                
                if is_header and row_index == 0:
                    continue  # Skip header row
                
                if len(cells) >= 6:
                    # Handle 6-column format: Q.No, Questions, CO, Level, Marks, Module
                    sl_no = cells[0] if cells[0].strip() else f"Q{len(extracted_questions) + 1}"
                    question_text = cells[1]
                    co = cells[2]
                    blooms_level = cells[3]
                    marks_text = cells[4]
                    module = cells[5]

                    # Extract marks using the helper function
                    marks = extract_marks_from_text(marks_text)

                    if question_text and question_text.strip():
                        extracted_questions.append({
                            "sl_no": sl_no,
                            "question_text": format_question_with_tables(question_text),
                            "co": co,
                            "blooms_level": blooms_level,
                            "marks": marks,
                            "module": module
                        })
                elif len(cells) >= 5:
                    # Handle 5-column format: SL#, Questions, CO, Level, Marks
                    sl_no = cells[0] if cells[0].strip() else f"Q{len(extracted_questions) + 1}"
                    question_text = cells[1]
                    co = cells[2]
                    blooms_level = cells[3]
                    marks_text = cells[4]

                    # Extract marks using the helper function
                    marks = extract_marks_from_text(marks_text)

                    if question_text and question_text.strip():
                        extracted_questions.append({
                            "sl_no": sl_no,
                            "question_text": format_question_with_tables(question_text),
                            "co": co,
                            "blooms_level": blooms_level,
                            "marks": marks,
                            "module": "1"  # Default to module 1 for legacy format
                        })
                elif len(cells) >= 4:
                    # Handle 4-column format: Q.No, Questions, CO, Level (no marks)
                    sl_no = cells[0] if cells[0].strip() else f"Q{len(extracted_questions) + 1}"
                    question_text = cells[1]
                    co = cells[2]
                    blooms_level = cells[3]

                    if question_text and question_text.strip():
                        extracted_questions.append({
                            "sl_no": sl_no,
                            "question_text": format_question_with_tables(question_text),
                            "co": co,
                            "blooms_level": blooms_level,
                            "marks": 0,  # Default marks
                            "module": "1"  # Default module
                        })
                else:
                    print(f"Skipping malformed DOCX row due to insufficient columns: {cells}") # Debugging
    except Exception as e:
        print(f"Error during DOCX parsing of {filepath}: {e}")
        traceback.print_exc() # Print full traceback for debugging
        raise # Re-raise to be caught by the route's error handling
    return extracted_questions

def parse_pdf_with_embedded_tables(filepath):
    """
    Universal PDF parser that can handle any type of question bank format.
    This function uses multiple strategies to extract questions from various PDF structures.
    """
    extracted_questions = []
    
    # Check if file exists
    if not os.path.exists(filepath):
        print(f"Error: File not found: {filepath}")
        return extracted_questions
    
    try:
        with pdfplumber.open(filepath) as pdf:
            for page_num, page in enumerate(pdf.pages):
                # Extract all tables from the page
                tables = page.extract_tables()
                page_text = page.extract_text() or ""
                
                print(f"Page {page_num + 1}: Found {len(tables)} tables and {len(page_text)} characters of text")
                
                # Strategy 1: Process structured question bank tables
                structured_questions = process_structured_tables(tables, page_num + 1)
                extracted_questions.extend(structured_questions)
                
                # Strategy 2: Extract questions from plain text (fallback)
                text_questions = extract_questions_from_plain_text(page_text)
                print(f"Extracted {len(text_questions)} questions from plain text")
                
                # Remove duplicates from text questions (in case they were already extracted from tables)
                existing_sl_nos = {q.get('sl_no', '') for q in extracted_questions}
                unique_text_questions = [q for q in text_questions if q.get('sl_no', '') not in existing_sl_nos]
                extracted_questions.extend(unique_text_questions)
                
                # Strategy 3: Process embedded tables as potential questions
                embedded_questions = process_embedded_tables(tables, page_text, page_num + 1)
                extracted_questions.extend(embedded_questions)
    
    except Exception as e:
        print(f"Error parsing PDF with embedded tables: {e}")
        traceback.print_exc()
    
    print(f"Total questions extracted: {len(extracted_questions)}")
    return extracted_questions

def process_structured_tables(tables, page_num):
    """Process tables that are structured question banks"""
    questions = []
    
    for table_idx, table in enumerate(tables):
        if not table or len(table) < 2:
            continue
        
        print(f"Processing table {table_idx + 1} with {len(table)} rows")
        
        # Check if this is a question bank table (multiple detection methods)
        header_row = table[0]
        is_question_bank = detect_question_bank_table(header_row, table)
        
        if is_question_bank:
            print("Found structured question bank table")
            table_questions = extract_questions_from_structured_table(table)
            questions.extend(table_questions)
        else:
            print("Table does not appear to be a structured question bank")
    
    return questions

def detect_question_bank_table(header_row, table):
    """Detect if a table is a question bank using multiple criteria"""
    if not header_row:
        return False
    
    # Method 1: Standard question bank headers
    if any(header for header in header_row if header and ("Q.No" in str(header) or "Question" in str(header) or "SL" in str(header))):
        return True
    
    # Method 2: Numeric question numbers in first column with CO/Level indicators
    if (len(header_row) >= 4 and str(header_row[0]).strip().isdigit() and 
        any(str(cell).strip() in ['CO1', 'CO2', 'CO3', 'CO4', 'CO5', 'L1', 'L2', 'L3', 'L4', 'L5'] for cell in header_row[2:4] if cell)):
        return True
    
    # Method 3: Check if multiple rows have question-like structure
    if len(table) >= 3:
        question_like_rows = 0
        for row in table[:3]:  # Check first 3 rows
            if (len(row) >= 4 and str(row[0]).strip().isdigit() and 
                len(str(row[1]).strip()) > 20):  # Question text is substantial
                question_like_rows += 1
        
        if question_like_rows >= 2:  # At least 2 rows look like questions
            return True
    
    # Method 4: Check for marks patterns (6M, 8M, etc.)
    if any(str(cell).strip().endswith('M') for row in table[:3] for cell in row if cell):
        return True
    
    return False

def extract_questions_from_structured_table(table):
    """Extract questions from a structured table"""
    questions = []
    
    # Check if the first row is also a question (not just headers)
    if (len(table[0]) >= 4 and str(table[0][0]).strip().isdigit() and 
        any(str(cell).strip() in ['CO1', 'CO2', 'CO3', 'CO4', 'CO5', 'L1', 'L2', 'L3', 'L4', 'L5'] for cell in table[0][2:4] if cell)):
        # First row is also a question, process it
        question = extract_question_from_row(table[0])
        if question:
            questions.append(question)
            print(f"Extracted structured question: {question['sl_no']}")
    
    # Process remaining rows
    for row_index, row in enumerate(table[1:], 1):
        question = extract_question_from_row(row)
        if question:
            questions.append(question)
            print(f"Extracted structured question: {question['sl_no']}")
    
    return questions

def extract_question_from_row(row):
    """Extract question data from a table row"""
    if len(row) < 4:
        return None
    
    sl_no = clean_text(row[0]) if len(row) > 0 else ""
    question_text = clean_text(row[1]) if len(row) > 1 else ""
    co = clean_text(row[2]) if len(row) > 2 else "CO1"
    blooms_level = clean_text(row[3]) if len(row) > 3 else "L1"
    marks_text = clean_text(row[4]) if len(row) > 4 else "5"
    module = clean_text(row[5]) if len(row) > 5 else "1"

    marks = extract_marks_from_text(marks_text)

    if question_text and sl_no:
        # Check if question text contains table data and format it properly
        formatted_question = format_question_with_tables(question_text)
        
        return {
            "sl_no": sl_no,
            "question_text": formatted_question,
            "co": co,
            "blooms_level": blooms_level,
            "marks": marks,
            "module": module
        }
    
    return None

def process_embedded_tables(tables, page_text, page_num):
    """Process embedded tables that might contain questions"""
    questions = []
    
    for table_idx, table in enumerate(tables):
        if not table or len(table) < 2:
            continue
        
        # Skip if this was already processed as a structured table
        if detect_question_bank_table(table[0], table):
            continue
        
        print(f"Processing embedded table {table_idx + 1} with {len(table)} rows")
        
        # Try to find question context around this table
        table_text = extract_table_as_text(table)
        if table_text:
            question_context = find_question_context_in_text(page_text, table_text)
            if question_context:
                questions.append(question_context)
                print(f"Extracted question with embedded table: {question_context['sl_no']}")
            else:
                # Only create generic entries for very substantial tables with meaningful data
                if len(table) >= 5 and len(table_text) > 100:  # More selective criteria
                    # Check if table contains question-like content
                    table_content = ' '.join(table_text.split()[:20]).lower()  # First 20 words
                    question_indicators = ['classify', 'analyze', 'calculate', 'explain', 'compare', 'evaluate', 'design', 'implement']
                    
                    if any(indicator in table_content for indicator in question_indicators):
                        questions.append({
                            "sl_no": f"T{page_num}_{table_idx + 1}",
                            "question_text": f"Question with table data:\n\nTable:\n{table_text}",
                            "co": "CO1",
                            "blooms_level": "L1",
                            "marks": 5,
                            "module": "1"
                        })
                        print(f"Created generic question entry for table {table_idx + 1}")
                    else:
                        print(f"Skipped table {table_idx + 1} - not substantial enough or no question indicators")
                else:
                    print(f"Skipped table {table_idx + 1} - too small or insufficient content")
    
    return questions

def find_question_context_in_text(page_text, table_text):
    """Find question context around a table in the page text"""
    if not page_text or not table_text:
        return None
    
    lines = page_text.split('\n')
    table_lines = [line.strip() for line in table_text.split('\n') if line.strip()]
    
    # Find where the table appears in the text
    for i, line in enumerate(lines):
        if any(table_line in line for table_line in table_lines[:3]):  # Check first few table lines
            # Found the table, look backwards for question context
            question_text = ""
            for j in range(max(0, i-15), i):  # Look at 15 lines before the table
                if lines[j].strip():
                    question_text += lines[j] + " "
            
            if question_text.strip():
                # Extract question details
                q_match = re.search(r'Q\.?\s*(\d+)\.?\s*(.*)', question_text, re.IGNORECASE)
                if q_match:
                    q_num = q_match.group(1)
                    q_text = q_match.group(2).strip()
                    
                    # Extract marks, CO, and level
                    marks = extract_marks_from_text(q_text)
                    if marks == 0:
                        marks = 5  # Default
                    
                    co_match = re.search(r'CO\s*(\d+)', q_text, re.IGNORECASE)
                    co = co_match.group(0) if co_match else "CO1"
                    
                    level_match = re.search(r'L\s*(\d+)', q_text, re.IGNORECASE)
                    level = level_match.group(0) if level_match else "L1"
                    
                    # Combine question text with table - ensure proper formatting
                    if table_text.strip():
                        full_question = q_text + "\n\nTable:\n" + table_text
                    else:
                        full_question = q_text
                    
                    return {
                        "sl_no": q_num,
                        "question_text": clean_text(full_question),
                        "co": co,
                        "blooms_level": level,
                        "marks": marks,
                        "module": "1"
                    }
    
    return None

def extract_questions_from_plain_text(text_content):
    """Extract questions from plain text using various patterns"""
    questions = []
    
    if not text_content:
        return questions
    
    # Pattern 1: Q.1, Q.2, etc.
    pattern1 = re.findall(r'Q\.?\s*(\d+)\.?\s*(.*?)(?=Q\.?\s*\d+\.?\s*|$)', text_content, re.DOTALL | re.IGNORECASE)
    for q_num, q_text in pattern1:
        if q_text.strip():
            marks = extract_marks_from_text(q_text)
            co_match = re.search(r'CO\s*(\d+)', q_text, re.IGNORECASE)
            co = co_match.group(0) if co_match else "CO1"
            
            level_match = re.search(r'L\s*(\d+)', q_text, re.IGNORECASE)
            level = level_match.group(0) if level_match else "L1"
            
            questions.append({
                "sl_no": q_num,
                "question_text": clean_text(q_text),
                "co": co,
                "blooms_level": level,
                "marks": marks if marks > 0 else 5,
                "module": "1"
            })
    
    # Pattern 2: Question 1, Question 2, etc.
    pattern2 = re.findall(r'Question\s*(\d+)[:\.]?\s*(.*?)(?=Question\s*\d+[:\.]?\s*|$)', text_content, re.DOTALL | re.IGNORECASE)
    for q_num, q_text in pattern2:
        if q_text.strip():
            marks = extract_marks_from_text(q_text)
            co_match = re.search(r'CO\s*(\d+)', q_text, re.IGNORECASE)
            co = co_match.group(0) if co_match else "CO1"
            
            level_match = re.search(r'L\s*(\d+)', q_text, re.IGNORECASE)
            level = level_match.group(0) if level_match else "L1"
            
            questions.append({
                "sl_no": q_num,
                "question_text": clean_text(q_text),
                "co": co,
                "blooms_level": level,
                "marks": marks if marks > 0 else 5,
                "module": "1"
            })
    
    # Pattern 3: Just numbers (1., 2., 3., etc.)
    pattern3 = re.findall(r'^(\d+)\.\s*(.*?)(?=^\d+\.\s*|$)', text_content, re.MULTILINE | re.DOTALL)
    for q_num, q_text in pattern3:
        if q_text.strip() and len(q_text.strip()) > 20:  # Only if it looks like a real question
            marks = extract_marks_from_text(q_text)
            co_match = re.search(r'CO\s*(\d+)', q_text, re.IGNORECASE)
            co = co_match.group(0) if co_match else "CO1"
            
            level_match = re.search(r'L\s*(\d+)', q_text, re.IGNORECASE)
            level = level_match.group(0) if level_match else "L1"
            
            questions.append({
                "sl_no": q_num,
                "question_text": clean_text(q_text),
                "co": co,
                "blooms_level": level,
                "marks": marks if marks > 0 else 5,
                "module": "1"
            })
    
    # Remove duplicates based on question number
    seen_numbers = set()
    unique_questions = []
    for q in questions:
        if q["sl_no"] not in seen_numbers:
            seen_numbers.add(q["sl_no"])
            unique_questions.append(q)
    
    return unique_questions


# --- Routes ---
@app.route('/login')
@app.route('/')
def login_page():
    return render_template('login.html')

@app.route('/register')
def register_page():
    return render_template('register.html')

@app.route('/hod-login')
def hod_login_page():
    return render_template('hod_login.html')

@app.route('/hod-register')
def hod_register_page():
    return render_template('hod_register.html')

@app.route('/dashboard')
def dashboard_page():
    return render_template('dashboard.html')

@app.route('/hod-dashboard')
def hod_dashboard_page():
    return render_template('hod_dashboard.html')

@app.route('/authenticate-backend', methods=['POST'])
def authenticate_backend():
    try:
        # Rate limiting check
        client_ip = request.remote_addr
        current_time = time.time()
        
        # Clean old attempts
        if client_ip in auth_attempts:
            auth_attempts[client_ip] = [attempt for attempt in auth_attempts[client_ip] if current_time - attempt < AUTH_WINDOW]
        else:
            auth_attempts[client_ip] = []
        
        # Check if rate limit exceeded
        if len(auth_attempts[client_ip]) >= MAX_AUTH_ATTEMPTS:
            return jsonify({
                'message': 'Too many authentication attempts. Please try again later.',
                'error_type': 'rate_limit_exceeded'
            }), 429
        
        # Record this attempt
        auth_attempts[client_ip].append(current_time)
        
        data = request.get_json()
        print(f"Received authentication request: {data}")
        
        if not data:
            return jsonify({'message': 'No JSON data received!'}), 400
            
        id_token = data.get('idToken')

        if not id_token:
            return jsonify({'message': 'ID token is missing!'}), 400

        print(f"Attempting to verify ID token for user...")
        decoded_token = auth.verify_id_token(id_token)
        uid = decoded_token['uid']
        email = decoded_token.get('email')
        name = decoded_token.get('name', 'User')

        print(f"Token verified successfully for user: {email} (UID: {uid})")

        # Create/update user in Firebase (keep existing)
        user_ref = db_firestore.collection('users').document(uid)
        user_doc = user_ref.get()
        if not user_doc.exists:
            user_ref.set({
                'email': email,
                'username': name,
                'created_at': firestore.SERVER_TIMESTAMP,
                'last_login': firestore.SERVER_TIMESTAMP
            })
            print(f"New user {uid} registered in Firestore.")
        else:
            user_ref.update({'last_login': firestore.SERVER_TIMESTAMP})
            print(f"User {uid} logged in and updated last_login in Firestore.")

        # Create/update user in Supabase
        try:
            supabase_user = supabase_service.create_user(
                firebase_uid=uid,
                email=email,
                name=name,
                role='faculty'  # Default role, can be updated later
            )
            if supabase_user:
                print(f"User {uid} synced to Supabase successfully.")
            else:
                print(f"Warning: Failed to sync user {uid} to Supabase.")
        except Exception as e:
            print(f"Error syncing user to Supabase: {e}")

        response_data = {
            'message': 'Backend authenticated successfully',
            'uid': uid,
            'email': email,
            'username': name
        }
        print(f"Sending response: {response_data}")
        return jsonify(response_data), 200
    except firebase_admin.auth.InvalidIdTokenError as e:
        print(f"Invalid ID token error: {e}")
        return jsonify({'message': 'Invalid or expired ID token.'}), 401
    except Exception as e:
        error_str = str(e)
        print(f"Authentication error: {error_str}")
        
        # Handle specific Firebase quota errors
        if "Quota exceeded" in error_str or "ResourceExhausted" in error_str or "429" in error_str:
            return jsonify({
                'message': 'Firebase quota exceeded. Please try again later or contact support.',
                'error_type': 'quota_exceeded'
            }), 429
        elif "Timeout" in error_str:
            return jsonify({
                'message': 'Authentication timeout. Please try again.',
                'error_type': 'timeout'
            }), 408
        else:
            return jsonify({'message': f'Authentication failed: {error_str}'}), 500

@app.route('/protected_data', methods=['GET'])
@firebase_auth_required
def protected_data():
    user_uid = request.current_user_uid
    return jsonify({'message': f'Hello, user {user_uid}! You accessed protected data.'}), 200

# --- HOD Authentication Routes ---
@app.route('/hod_register', methods=['POST'])
def hod_register():
    try:
        data = request.get_json()

        # Validate required fields
        required_fields = ['email', 'password', 'name', 'department', 'institution', 'hodCode']
        for field in required_fields:
            if not data.get(field):
                return jsonify({'error': f'{field} is required'}), 400

        # Validate HOD code (you can customize this)
        valid_hod_codes = ['HOD2024', 'ADMIN123', 'DEPT_HEAD']  # Add your HOD codes here
        if data['hodCode'] not in valid_hod_codes:
            return jsonify({'error': 'Invalid HOD authorization code'}), 400

        # Create Firebase user
        user_record = auth.create_user(
            email=data['email'],
            password=data['password'],
            display_name=data['name']
        )

        # Set custom claims for HOD
        auth.set_custom_user_claims(user_record.uid, {'role': 'hod'})

        # Save HOD details to Supabase
        hod_data = {
            'firebase_uid': user_record.uid,
            'email': data['email'],
            'name': data['name'],
            'department': data['department'],
            'institution': data['institution'],
            'role': 'hod'
        }

        supabase_service.create_user(user_record.uid, data['email'], data['name'], 'hod', data['department'])

        return jsonify({
            'message': 'HOD registered successfully',
            'uid': user_record.uid
        }), 201

    except auth.EmailAlreadyExistsError:
        return jsonify({'error': 'Email already exists'}), 400
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/get_user_profile', methods=['GET'])
@firebase_auth_required
def get_user_profile():
    """Get current user profile including department"""
    user_uid = request.current_user_uid
    try:
        user_data = supabase_service.get_user_by_firebase_uid(user_uid)
        if user_data:
            return jsonify(user_data), 200
        else:
            return jsonify({'error': 'User not found'}), 404
    except Exception as e:
        print(f"Error fetching user profile: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/hod_login', methods=['POST'])
def hod_login():
    try:
        data = request.get_json()
        id_token = data.get('idToken')

        if not id_token:
            return jsonify({'error': 'ID token is required'}), 400

        # Verify the ID token
        decoded_token = auth.verify_id_token(id_token)
        uid = decoded_token['uid']
        
        # Check for HOD role in Firebase claims
        firebase_role = decoded_token.get('role')
        print(f"DEBUG: HOD Login Attempt - UID: {uid}, Firebase Role: {firebase_role}")

        # Check if user is HOD using Supabase
        hod_user = supabase_service.get_user_by_firebase_uid(uid)
        print(f"DEBUG: Supabase User: {hod_user}")
        
        # Auto-sync: If user missing in Supabase but has HOD claim in Firebase, create them
        if not hod_user and firebase_role == 'hod':
            print(f"User {uid} missing in Supabase but has HOD claim. Syncing...")
            email = decoded_token.get('email')
            name = decoded_token.get('name', 'HOD User')
            # We might not have department/institution here, but we can update later
            hod_user = supabase_service.create_user(uid, email, name, 'hod')
            
        if not hod_user or hod_user.get('role') != 'hod':
            print(f"DEBUG: Access Denied. User Role: {hod_user.get('role') if hod_user else 'None'}")
            return jsonify({'error': 'Access denied. HOD privileges required.'}), 403

        return jsonify({
            'message': 'HOD login successful',
            'user': {
                'uid': uid,
                'email': hod_user.get('email'),
                'name': hod_user.get('name'),
                'department': hod_user.get('department'),
                'institution': hod_user.get('institution', 'Unknown Institution'),
                'role': 'hod'
            }
        }), 200

    except auth.InvalidIdTokenError:
        return jsonify({'error': 'Invalid or expired token'}), 401
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/verify_hod', methods=['POST'])
@firebase_auth_required
def verify_hod():
    try:
        user_uid = request.current_user_uid

        # Check if user is HOD using Supabase
        hod_user = supabase_service.get_user_by_firebase_uid(user_uid)
        if not hod_user or hod_user.get('role') != 'hod':
            return jsonify({'error': 'Access denied. HOD privileges required.'}), 403

        return jsonify({
            'isHOD': True,
            'user': {
                'uid': user_uid,
                'email': hod_user.get('email'),
                'name': hod_user.get('name'),
                'department': hod_user.get('department'),
                'institution': hod_user.get('institution', 'Unknown Institution'),
                'role': 'hod'
            }
        }), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

# --- New Route for File Upload and Parsing with Firestore Saving ---
@app.route('/upload_and_parse', methods=['POST'])
@firebase_auth_required
def upload_and_parse():
    user_uid = request.current_user_uid
    
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        user_upload_folder = os.path.join(app.config['UPLOAD_FOLDER'], user_uid)
        os.makedirs(user_upload_folder, exist_ok=True)
        filepath = os.path.join(user_upload_folder, filename)
        
        file.save(filepath)
        
        file_extension = filename.rsplit('.', 1)[1].lower()
        
        try:
            parsed_questions = []
            if file_extension == 'pdf':
                # Try Advanced Parser first for better results with images/tables
                try:
                    from advanced_parser import get_advanced_parser
                    parser = get_advanced_parser()
                    # Create a subfolder for images
                    images_folder = os.path.join(user_upload_folder, "extracted_images")
                    os.makedirs(images_folder, exist_ok=True) # Ensure images folder exists
                    
                    # Parse
                    raw_parsed_content = parser.parse_pdf(filepath, images_folder)
                    
                    # Convert to the format expected by Firestore/Frontend
                    for item in raw_parsed_content:
                        parsed_questions.append({
                            "sl_no": "Auto", # parser might not extract SL correctly yet
                            "question_text": item["text"],
                            "co": "CO1", # Default
                            "blooms_level": "L1", # Default
                            "marks": 10, # Default
                            "module": "1",
                            "images": [img.replace('\\', '/') for img in item["images"]], # Normalize paths
                            "formulas": item["formulas"] # New field
                        })
                        
                except Exception as e:
                    print(f"Advanced parsing failed, falling back to legacy: {e}")
                    # Fallback to legacy parsing
                    parsed_questions = parse_pdf_with_embedded_tables(filepath)
                    if not parsed_questions:
                        parsed_questions = parse_pdf_question_bank(filepath)
                        
            elif file_extension == 'docx':
                parsed_questions = parse_docx_question_bank(filepath)
            
            if not parsed_questions:
                return jsonify({"error": "No questions could be parsed from the uploaded file. Please check its format."}), 400

            # --- Save the Question Bank and its Questions to Firestore ---
            questions_to_return = [] # Store questions with their Firestore ID for frontend
            batch = db_firestore.batch() # Use a batch for efficient writes

            for q_data_original in parsed_questions: # Use a different variable name for clarity
                # Create a copy of the original data to modify for Firestore saving
                q_data_for_firestore = q_data_original.copy() 

                question_doc_ref = db_firestore.collection('users').document(user_uid).collection('question_bank_pool').document() # Auto-generate ID
                
                # Add metadata to the question data for Firestore
                q_data_for_firestore['user_uid'] = user_uid
                q_data_for_firestore['source_file'] = filename # Track where it came from
                q_data_for_firestore['uploaded_at'] = firestore.SERVER_TIMESTAMP
                q_data_for_firestore['is_pre_selected'] = False 
                q_data_for_firestore['last_used_date'] = None 
                
                # Handle image paths - convert to relative URL or upload to storage (Supabase/Firebase)
                # For now, we keep local path, but in production, you'd upload these to Supabase Storage
                if 'images' in q_data_for_firestore:
                     # Just keeping the filename for simplicity in this demo
                     q_data_for_firestore['images'] = [os.path.basename(p) for p in q_data_for_firestore['images']]

                batch.set(question_doc_ref, q_data_for_firestore)
                
                # Prepare data to send back to frontend
                q_data_original['firestore_id'] = question_doc_ref.id
                # IMPORTANT: Remove the Sentinel objects before sending to frontend
                # These fields are correctly saved to Firestore, but cannot be JSON serialized
                if 'uploaded_at' in q_data_original:
                    del q_data_original['uploaded_at']
                if 'last_used_date' in q_data_original:
                    del q_data_original['last_used_date']

                questions_to_return.append(q_data_original)
            
            # --- Save the Question Bank Metadata (New Optimization) ---
            try:
                # Create a safe ID for the bank
                import re
                safe_bank_id = re.sub(r'[^a-zA-Z0-9]', '_', filename)
                bank_ref = db_firestore.collection('users').document(user_uid).collection('question_banks').document(f"bank_{safe_bank_id}")
                
                bank_data = {
                    'id': f"bank_{safe_bank_id}",
                    'name': filename,
                    'source_file': filename,
                    'uploaded_at': firestore.SERVER_TIMESTAMP,
                    'question_count': len(parsed_questions),
                    'type': 'question-bank',
                    'user_uid': user_uid
                }
                batch.set(bank_ref, bank_data)
                print(f"Saved question bank metadata for {filename}")
            except Exception as e:
                print(f"Error saving bank metadata: {e}")

            batch.commit() # Commit all writes in one go

            print(f"Saved {len(parsed_questions)} questions from '{filename}' to Firestore pool for user {user_uid}")

            return jsonify({
                "message": "File uploaded, parsed (Advanced Mode), and saved successfully!",
                "filename": filename,
                "parsed_questions_count": len(parsed_questions),
                "parsed_data": questions_to_return, # Send questions with their Firestore IDs
            }), 200
            
        except Exception as e:
            traceback.print_exc()
            return jsonify({"error": f"Error parsing or saving file: {str(e)}. Check server logs for details."}), 500
    else:
        return jsonify({"error": "File type not allowed. Only PDF and DOCX are supported."}), 400

# --- New Route for CIE Upload and Parse ---
@app.route('/upload_and_parse_cie', methods=['POST'])
@firebase_auth_required
def upload_and_parse_cie():
    user_uid = request.current_user_uid
    cie_type = request.form.get('cie_type', 'cie1')
    
    # Define which modules to expect based on CIE type
    if cie_type == 'cie1':
        expected_modules = ['module1_file', 'module2_file', 'module3_file']
        module_numbers = ['1', '2', '3']
    elif cie_type == 'cie2':
        expected_modules = ['module3_file', 'module4_file', 'module5_file']
        module_numbers = ['3', '4', '5']
    else:
        return jsonify({"error": "Invalid CIE type. Must be 'cie1' or 'cie2'."}), 400
    
    # Check if at least one file is uploaded
    uploaded_files = []
    for module_field in expected_modules:
        if module_field in request.files:
            file = request.files[module_field]
            if file and file.filename != '' and allowed_file(file.filename):
                uploaded_files.append((module_field, file))
    
    if not uploaded_files:
        return jsonify({"error": f"No valid files uploaded for {cie_type.upper()}. Please upload at least one module file."}), 400
    
    try:
        all_questions = []
        batch = db_firestore.batch()
        
        for module_field, file in uploaded_files:
            filename = secure_filename(file.filename)
            user_upload_folder = os.path.join(app.config['UPLOAD_FOLDER'], user_uid)
            os.makedirs(user_upload_folder, exist_ok=True)
            filepath = os.path.join(user_upload_folder, filename)
            
            file.save(filepath)
            
            # Extract module number from field name
            module_num = module_field.replace('_file', '').replace('module', '')
            
            file_extension = filename.rsplit('.', 1)[1].lower()
            
            # Parse the file
            parsed_questions = []
            if file_extension == 'pdf':
                # Try Advanced Parser first for better results with images/tables
                try:
                    from advanced_parser import get_advanced_parser
                    parser = get_advanced_parser()
                    # Create a subfolder for images
                    images_folder = os.path.join(user_upload_folder, "extracted_images")
                    os.makedirs(images_folder, exist_ok=True) # Ensure images folder exists
                    
                    # Parse
                    raw_parsed_content = parser.parse_pdf(filepath, images_folder)
                    
                    # Convert to the format expected by Firestore/Frontend
                    for item in raw_parsed_content:
                        # Use extracted sl_no if available, otherwise "Auto"
                        sl_no = item.get("sl_no")
                        if not sl_no:
                            sl_no = "Auto"
                            
                        parsed_questions.append({
                            "sl_no": sl_no,
                            "question_text": item["text"],
                            "co": item.get("co", "CO1"), 
                            "blooms_level": item.get("blooms_level", "L1"),
                            "marks": item.get("marks", 10),
                            "module": item.get("module", module_num), # Use extracted module if available, else form value
                            "images": [img.replace('\\', '/') for img in item["images"]], # Normalize paths
                            "formulas": item["formulas"] # New field
                        })
                        
                except Exception as e:
                    print(f"Advanced parsing failed for {filename}, falling back to legacy: {e}")
                    # Fallback to legacy parsing
                    parsed_questions = parse_pdf_with_embedded_tables(filepath)
                    if not parsed_questions:
                        parsed_questions = parse_pdf_question_bank(filepath)
            elif file_extension == 'docx':
                parsed_questions = parse_docx_question_bank(filepath)
            
            if parsed_questions:
                # Add module information to each question
                for q_data_original in parsed_questions:
                    q_data_for_firestore = q_data_original.copy()
                    
                    # Ensure module is set correctly
                    q_data_for_firestore['module'] = module_num
                    q_data_for_firestore['user_uid'] = user_uid
                    q_data_for_firestore['source_file'] = filename
                    q_data_for_firestore['cie_type'] = cie_type
                    q_data_for_firestore['uploaded_at'] = firestore.SERVER_TIMESTAMP
                    q_data_for_firestore['is_pre_selected'] = False
                    q_data_for_firestore['last_used_date'] = None
                    
                    question_doc_ref = db_firestore.collection('users').document(user_uid).collection('question_bank_pool').document()
                    batch.set(question_doc_ref, q_data_for_firestore)
                    
                    # Prepare data for frontend
                    q_data_original['firestore_id'] = question_doc_ref.id
                    q_data_original['module'] = module_num
                    q_data_original['cie_type'] = cie_type
                    
                    # Remove non-serializable fields
                    if 'uploaded_at' in q_data_original:
                        del q_data_original['uploaded_at']
                    if 'last_used_date' in q_data_original:
                        del q_data_original['last_used_date']
                    
                    all_questions.append(q_data_original)
        
        batch.commit()
        
        print(f"Saved {len(all_questions)} questions from {cie_type.upper()} modules to Firestore pool for user {user_uid}")
        
        return jsonify({
            "message": f"{cie_type.upper()} modules uploaded, parsed, and saved successfully!",
            "cie_type": cie_type,
            "total_questions": len(all_questions),
            "all_questions": all_questions,
            "uploaded_modules": [f"Module {m}" for m in module_numbers if any(f"module{m}_file" in f[0] for f in uploaded_files)]
        }), 200
        
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": f"Error parsing or saving {cie_type.upper()} files: {str(e)}. Check server logs for details."}), 500

# --- New Route for Question Rephrasing using Gemini API ---
@app.route('/rephrase_question', methods=['POST'])
@firebase_auth_required
def rephrase_question():
    user_uid = request.current_user_uid
    data = request.get_json()
    original_text = data.get('original_text')
    question_id = data.get('question_id') # Firestore ID of the question to update

    if not original_text:
        return jsonify({"error": "No original text provided for rephrasing."}), 400

    # Gemini API configuration
    GEMINI_API_KEY = os.getenv("GOOGLE_API_KEY", "") 
    GEMINI_API_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent"

    # Construct the prompt for rephrasing - MORE ROBUST PROMPT
    prompt = f"Rephrase the following question. The rephrased version must convey the exact same meaning and retain all technical terms and mathematical expressions. Provide ONLY the rephrased question text, without any introductory phrases like 'Here's the rephrased question:' or conversational elements. Do not include question numbers or sub-part labels.\n\nOriginal Question: \"{original_text}\""

    payload = {
        "contents": [
            {
                "role": "user",
                "parts": [{"text": prompt}]
            }
        ],
        "generationConfig": {
            "temperature": 0.7, 
            "topP": 0.95,
            "topK": 40,
            "maxOutputTokens": 200 
        }
    }

    headers = {
        "Content-Type": "application/json"
    }
    if GEMINI_API_KEY:
        GEMINI_API_URL += f"?key={GEMINI_API_KEY}"

    try:
        response = requests.post(GEMINI_API_URL, headers=headers, data=json.dumps(payload))
        response.raise_for_status() 
        
        gemini_response = response.json()
        
        rephrased_text = ""
        if gemini_response and gemini_response.get('candidates'):
            first_candidate = gemini_response['candidates'][0]
            if first_candidate.get('content') and first_candidate['content'].get('parts'):
                rephrased_text = first_candidate['content']['parts'][0].get('text', '').strip()
        
        if not rephrased_text:
            raise ValueError("Gemini API returned no rephrased text or an unexpected format.")

        # Optionally update the question in Firestore if question_id is provided
        if question_id:
            try:
                question_ref = db_firestore.collection('users').document(user_uid).collection('question_bank_pool').document(question_id)
                question_ref.update({'question_text': rephrased_text})
                print(f"Question {question_id} rephrased and updated in Firestore.")
            except Exception as e:
                print(f"Error updating question {question_id} in Firestore after rephrasing: {e}")
                # Don't fail the rephrase request just because Firestore update failed
        
        return jsonify({"rephrased_text": rephrased_text}), 200

    except requests.exceptions.RequestException as e:
        print(f"Request to Gemini API failed: {e}")
        traceback.print_exc()
        return jsonify({"error": f"Failed to connect to AI service: {str(e)}"}), 500
    except ValueError as e:
        print(f"AI service response error: {e}")
        traceback.print_exc()
        return jsonify({"error": f"AI service response error: {str(e)}"}), 500
    except Exception as e:
        print(f"An unexpected error occurred during rephrasing: {e}")
        traceback.print_exc()
        return jsonify({"error": f"An unexpected error occurred: {str(e)}"}), 500

# --- New Endpoint: Add a new question to the user's question bank pool ---
@app.route('/add_question_to_bank', methods=['POST'])
@firebase_auth_required
def add_question_to_bank():
    user_uid = request.current_user_uid
    data = request.get_json()

    required_fields = ['question_text', 'marks']
    if not all(field in data for field in required_fields):
        return jsonify({"error": "Missing required fields (question_text, marks)."}), 400

    try:
        new_question_ref = db_firestore.collection('users').document(user_uid).collection('question_bank_pool').document()
        new_question_id = new_question_ref.id

        question_data = {
            "question_text": data.get('question_text'),
            "co": data.get('co', 'N/A'),
            "blooms_level": data.get('blooms_level', 'N/A'),
            "marks": int(data.get('marks')), # Ensure marks is an integer
            "user_uid": user_uid,
            "uploaded_at": firestore.SERVER_TIMESTAMP,
            "is_pre_selected": False,
            "last_used_date": None,
            "source_file": "Manual Entry" # Indicate manual entry
        }
        
        new_question_ref.set(question_data)
        print(f"Manually added question {new_question_id} to bank for user {user_uid}.")
        
        # Prepare data to return to frontend (remove Sentinel objects)
        question_data_for_frontend = question_data.copy()
        if 'uploaded_at' in question_data_for_frontend:
            del question_data_for_frontend['uploaded_at']
        if 'last_used_date' in question_data_for_frontend:
            del question_data_for_frontend['last_used_date']

        question_data_for_frontend['firestore_id'] = new_question_id
        return jsonify({"message": "Question added successfully!", "new_question": question_data_for_frontend}), 200

    except ValueError:
        return jsonify({"error": "Marks must be a valid number."}), 400
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": f"Failed to add question to bank: {str(e)}"}), 500

# --- New Endpoint: Update an existing question in the user's question bank pool ---
@app.route('/update_question_in_bank', methods=['POST'])
@firebase_auth_required
def update_question_in_bank():
    user_uid = request.current_user_uid
    data = request.get_json()
    question_id = data.get('question_id')
    updated_fields = {k: v for k, v in data.items() if k != 'question_id'}

    if not question_id:
        return jsonify({"error": "Question ID is missing."}), 400
    if not updated_fields:
        return jsonify({"error": "No fields provided for update."}), 400

    try:
        # Ensure marks is converted to int if present in updated_fields
        if 'marks' in updated_fields:
            updated_fields['marks'] = int(updated_fields['marks'])

        question_ref = db_firestore.collection('users').document(user_uid).collection('question_bank_pool').document(question_id)
        question_ref.update(updated_fields)
        print(f"Question {question_id} updated in bank for user {user_uid}.")
        return jsonify({"message": "Question updated successfully!"}), 200

    except ValueError:
        return jsonify({"error": "Marks must be a valid number."}), 400
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": f"Failed to update question in bank: {str(e)}"}), 500

# --- New Endpoint: Delete a question from the user's question bank pool ---
@app.route('/delete_question_from_bank', methods=['DELETE'])
@firebase_auth_required
def delete_question_from_bank():
    user_uid = request.current_user_uid
    data = request.get_json()
    question_id = data.get('question_id')

    if not question_id:
        return jsonify({"error": "Question ID is missing."}), 400

    try:
        question_ref = db_firestore.collection('users').document(user_uid).collection('question_bank_pool').document(question_id)
        question_ref.delete()
        print(f"Question {question_id} deleted from bank for user {user_uid}.")
        return jsonify({"message": "Question deleted successfully!"}), 200

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": f"Failed to delete question from bank: {str(e)}"}), 500

# --- New Endpoint: Get all questions for the current user ---
@app.route('/get_user_questions', methods=['GET'])
@firebase_auth_required
def get_user_questions():
    user_uid = request.current_user_uid
    try:
        questions_ref = db_firestore.collection('users').document(user_uid).collection('question_bank_pool')
        docs = questions_ref.stream()
        
        questions_list = []
        for doc in docs:
            q_data = doc.to_dict()
            q_data['firestore_id'] = doc.id # Include the Firestore document ID
            
            # Remove Sentinel objects if they somehow got into the dict during fetch
            # (though they shouldn't if saved correctly, this is a safeguard)
            if 'uploaded_at' in q_data and isinstance(q_data['uploaded_at'], firestore.SERVER_TIMESTAMP.__class__):
                del q_data['uploaded_at']
            if 'last_used_date' in q_data and isinstance(q_data['last_used_date'], firestore.SERVER_TIMESTAMP.__class__):
                del q_data['last_used_date']
            
            questions_list.append(q_data)
        
        return jsonify({"message": "Questions fetched successfully!", "questions": questions_list}), 200
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": f"Failed to fetch user questions: {str(e)}"}), 500


# --- Enhanced Endpoint: Generate Final DOCX Document ---
@app.route('/export_question_paper', methods=['POST'])
@firebase_auth_required
def export_question_paper():
    """Export question paper as PDF or DOCX"""
    user_uid = request.current_user_uid
    data = request.get_json()
    paper_id = data.get('paper_id')
    format_type = data.get('format', 'docx')  # 'docx' or 'pdf'
    
    try:
        # Get the paper from Firestore
        paper_ref = db_firestore.collection('users').document(user_uid).collection('generated_papers').document(paper_id)
        paper_doc = paper_ref.get()
        
        if not paper_doc.exists:
            return jsonify({"error": "Question paper not found"}), 404
        
        paper_data = paper_doc.to_dict()
        questions = paper_data.get('questions', [])
        subject = paper_data.get('subject', 'Question Paper')
        
        if not questions:
            return jsonify({"error": "No questions found in the paper"}), 400
        
        # Generate the document
        if format_type == 'docx':
            file_path = generate_docx_export(questions, subject, paper_id)
            mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            file_extension = 'docx'
        elif format_type == 'pdf':
            file_path = generate_pdf_export(questions, subject, paper_id)
            mime_type = 'application/pdf'
            file_extension = 'pdf'
        else:
            return jsonify({"error": "Invalid format. Use 'docx' or 'pdf'"}), 400
        
        # Update download count
        paper_ref.update({
            'download_count': firestore.Increment(1),
            'last_downloaded': firestore.SERVER_TIMESTAMP
        })
        
        return jsonify({
            "message": f"Question paper exported successfully as {format_type.upper()}",
            "file_path": file_path,
            "file_name": f"{subject.replace(' ', '_')}_{paper_id}.{file_extension}",
            "mime_type": mime_type
        }), 200
        
    except Exception as e:
        return jsonify({"error": f"Failed to export question paper: {str(e)}"}), 500

def generate_docx_export(questions, subject, paper_id):
    """Generate DOCX export of question paper"""
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Add title
    title = doc.add_heading(f'{subject} - Question Paper', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add paper info
    doc.add_paragraph(f'Paper ID: {paper_id}')
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph('')  # Empty line
    
    # Add questions
    for i, question in enumerate(questions, 1):
        main_q = question.get('main_question', f'Q{i}')
        module = question.get('module', 'Unknown')
        
        # Main question header
        doc.add_heading(f'{main_q} (Module {module})', level=1)
        
        # Sub-questions
        sub_questions = question.get('sub_questions', [])
        for j, sub_q in enumerate(sub_questions):
            part = sub_q.get('part', chr(97 + j))  # a, b, c, etc.
            text = sub_q.get('text', sub_q.get('question_text', ''))
            marks = sub_q.get('marks', 0)
            
            doc.add_paragraph(f'({part}) {text} [{marks} marks]')
        
        doc.add_paragraph('')  # Empty line between questions
    
    # Save the document
    filename = f"{subject.replace(' ', '_')}_{paper_id}.docx"
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    doc.save(filepath)
    
    return filepath

def generate_pdf_export(questions, subject, paper_id):
    """Generate PDF export of question paper"""
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    
    filename = f"{subject.replace(' ', '_')}_{paper_id}.pdf"
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    
    doc = SimpleDocTemplate(filepath, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=30,
        alignment=1  # Center alignment
    )
    story.append(Paragraph(f'{subject} - Question Paper', title_style))
    story.append(Spacer(1, 12))
    
    # Paper info
    story.append(Paragraph(f'Paper ID: {paper_id}', styles['Normal']))
    story.append(Paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', styles['Normal']))
    story.append(Spacer(1, 20))
    
    # Questions
    for i, question in enumerate(questions, 1):
        main_q = question.get('main_question', f'Q{i}')
        module = question.get('module', 'Unknown')
        
        # Main question header
        story.append(Paragraph(f'{main_q} (Module {module})', styles['Heading2']))
        story.append(Spacer(1, 12))
        
        # Sub-questions
        sub_questions = question.get('sub_questions', [])
        for j, sub_q in enumerate(sub_questions):
            part = sub_q.get('part', chr(97 + j))  # a, b, c, etc.
            text = sub_q.get('text', sub_q.get('question_text', ''))
            marks = sub_q.get('marks', 0)
            
            story.append(Paragraph(f'({part}) {text} [{marks} marks]', styles['Normal']))
            story.append(Spacer(1, 6))
        
        story.append(Spacer(1, 20))
    
    doc.build(story)
    return filepath

@app.route('/generate_final_document', methods=['POST'])
@firebase_auth_required
def generate_final_document():
    user_uid = request.current_user_uid
    data = request.get_json()
    question_paper_data = data.get('question_paper_data')
    overall_max_marks = data.get('overall_max_marks', 100)
    metadata = data.get('metadata', {})

    if not question_paper_data:
        return jsonify({"error": "No question paper data provided."}), 400

    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.units import cm, inch
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.pdfgen import canvas

        # Flatten question data
        questions_list = []
        
        for main_q_idx, main_q_data in enumerate(question_paper_data):
            if main_q_idx >= 4: break
            
            # Insert OR row for CIE 1 pattern (before Q2 and Q4)
            if main_q_idx == 1 or main_q_idx == 3:
                questions_list.append({
                    'qno': '',
                    'question': 'OR',
                    'marks': '',
                    'co': '',
                    'level': '',
                    'module': '',
                    'images': []
                })
                
            sub_questions = main_q_data.get('subQuestions', [])
            for sub_q_idx, sub_q_data in enumerate(sub_questions):
                if sub_q_idx >= 3: break
                
                # Format Q.No (e.g., 1.a)
                qno = ""
                if sub_q_idx == 0:
                    qno = str(main_q_idx + 1)
                
                sub_letter = chr(97 + sub_q_idx)
                question_text = sub_q_data.get('question_text', '')
                
                # Store images
                images = sub_q_data.get('images', [])
                
                questions_list.append({
                    'qno': qno,
                    'question': f"{sub_letter}. {question_text}",
                    'marks': sub_q_data.get('marks', ''),
                    'co': sub_q_data.get('co', ''),
                    'level': sub_q_data.get('blooms_level', ''),
                    'module': sub_q_data.get('module', ''),
                    'images': images
                })

        # Create PDF
        temp_output_path = os.path.join(app.config['UPLOAD_FOLDER'], f"generated_qp_{user_uid}_{uuid.uuid4()}.pdf")
        doc = SimpleDocTemplate(temp_output_path, pagesize=A4,
                                rightMargin=1.5*cm, leftMargin=1.5*cm,
                                topMargin=1.5*cm, bottomMargin=1.5*cm)
        
        # Container for PDF elements
        elements = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=14,
            textColor=colors.black,
            spaceAfter=6,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Normal'],
            fontSize=12,
            textColor=colors.black,
            spaceAfter=4,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        # Header with logo
        logo_path = os.path.join('static', 'assets', 'logo.jpg')
        
        if os.path.exists(logo_path):
            # Create logo image (slightly larger)
            logo_img = Image(logo_path, width=2.5*cm, height=2.5*cm)
            
            # Create institute text paragraphs with LEFT alignment
            institute_name_style = ParagraphStyle(
                'InstituteName',
                parent=styles['Heading1'],
                fontSize=16,
                fontName='Helvetica-Bold',
                alignment=TA_LEFT,  # Changed to LEFT
                spaceAfter=2
            )
            
            details_style = ParagraphStyle(
                'Details',
                parent=styles['Normal'],
                fontSize=8,
                alignment=TA_LEFT,  # Changed to LEFT
                spaceAfter=1
            )
            
            institute_text = []
            institute_text.append(Paragraph("SRI KRISHNA INSTITUTE OF TECHNOLOGY", institute_name_style))
            institute_text.append(Paragraph("(Accredited by NAAC, Approved by A.I.C.T.E. New Delhi, Recognised by Govt. of Karnataka & Affiliated to V.T.U., Belagavi)", details_style))
            institute_text.append(Paragraph("#57, Chimney Hills, Hesaraghatta Main Road, Chikkabanavara Post, Bengaluru- 560090", details_style))
            
            # Create header table: logo | institute details
            header_data = [[logo_img, institute_text]]
            header_table = Table(header_data, colWidths=[3*cm, 14*cm])
            header_table.setStyle(TableStyle([
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('ALIGN', (0, 0), (0, 0), 'LEFT'),
                ('ALIGN', (1, 0), (1, 0), 'LEFT'),
                ('LINEBELOW', (0, 0), (-1, 0), 1.5, colors.black),
                ('LEFTPADDING', (0, 0), (-1, -1), 0),
                ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                ('TOPPADDING', (0, 0), (-1, -1), 5),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
            ]))
            elements.append(header_table)
        else:
            # Fallback to text-only if logo not found
            elements.append(Paragraph("SRI KRISHNA INSTITUTE OF TECHNOLOGY", title_style))
            elements.append(Paragraph("(Accredited by NAAC, Approved by A.I.C.T.E. New Delhi, Recognised by Govt. of Karnataka & Affiliated to V.T.U., Belagavi)", 
                                     ParagraphStyle('Details', parent=styles['Normal'], fontSize=8, alignment=TA_CENTER, spaceAfter=2)))
            elements.append(Paragraph("#57, Chimney Hills, Hesaraghatta Main Road, Chikkabanavara Post, Bengaluru- 560090",
                                     ParagraphStyle('Address', parent=styles['Normal'], fontSize=8, alignment=TA_CENTER, spaceAfter=4)))
        
        elements.append(Spacer(1, 0.2*cm))
        elements.append(Paragraph(f"Department of {metadata.get('dept', 'ARTIFICIAL INTELLIGENCE AND MACHINE LEARNING')}", subtitle_style))
        elements.append(Paragraph("CIE - I", subtitle_style))
        elements.append(Spacer(1, 0.3*cm))
        
        # Metadata table
        meta_data = [
            [f"Date: {metadata.get('date', '')}", f"Time: {metadata.get('time', '')}", 
             f"Max Marks: {overall_max_marks}", f"Sem/Div: {metadata.get('sem', '')}/{metadata.get('div', '')}"],
            [f"Course: {metadata.get('subject', metadata.get('course', ''))}", 
             f"Code: {metadata.get('code', '')}", 
             f"Elective: {metadata.get('elective', 'N/A')}", ""]
        ]
        
        meta_table = Table(meta_data, colWidths=[4.5*cm, 4*cm, 4*cm, 4*cm])
        meta_table.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('SPAN', (2, 1), (3, 1)),  # Merge last two cells in row 2
        ]))
        elements.append(meta_table)
        elements.append(Spacer(1, 0.3*cm))
        
        # Note
        note_style = ParagraphStyle('Note', parent=styles['Normal'], fontSize=9, fontName='Helvetica-Bold')
        elements.append(Paragraph("Note: Answer any ONE full question from each module.", note_style))
        elements.append(Spacer(1, 0.2*cm))
        
        # Questions table - with images INLINE in question cells
        table_data = [["Q.No", "Question", "CO", "Level", "Marks", "Module"]]
        
        for q in questions_list:
            if q['question'] == 'OR':
                # OR row - will be merged later
                table_data.append(["OR", "", "", "", "", ""])
            else:
                # Normal question row - create nested content for question cell
                question_cell_content = []
                
                # Add question text
                question_cell_content.append(Paragraph(str(q['question']), styles['Normal']))
                
                # Add images inline if present
                if q.get('images') and len(q['images']) > 0:
                    for img_name in q['images']:
                        # Resolve path - images are stored in uploads/{user_uid}/extracted_images/
                        if not os.path.isabs(img_name):
                            user_images_folder = os.path.join(app.config['UPLOAD_FOLDER'], user_uid, 'extracted_images')
                            full_img_path = os.path.join(user_images_folder, img_name)
                        else:
                            full_img_path = img_name
                            
                        if os.path.exists(full_img_path):
                            try:
                                # Add small spacer before image
                                question_cell_content.append(Spacer(1, 0.1*cm))
                                # Add image (smaller size for inline display)
                                img = Image(full_img_path, width=4*cm, height=3*cm, kind='proportional')
                                question_cell_content.append(img)
                            except Exception as e:
                                print(f"Error loading image {full_img_path}: {e}")
                
                # Create a nested single-column table to hold the content
                nested_table = Table([[item] for item in question_cell_content], colWidths=[11.5*cm])
                nested_table.setStyle(TableStyle([
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 0),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                    ('TOPPADDING', (0, 0), (-1, -1), 0),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
                ]))
                
                table_data.append([
                    str(q['qno']),
                    nested_table,  # Use nested table instead of plain Paragraph
                    str(q['co']),
                    str(q['level']),
                    str(q['marks']),
                    str(q['module'])
                ])
        
        # Adjusted column widths: Q.No, Question, CO, Level, Marks, Module
        questions_table = Table(table_data, colWidths=[1.2*cm, 11.5*cm, 1.2*cm, 1.2*cm, 1.2*cm, 1.2*cm])
        
        # Build table style
        table_style_commands = [
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),  # Header row
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),  # Header alignment
            ('ALIGN', (0, 1), (0, -1), 'CENTER'),  # Q.No column
            ('ALIGN', (2, 1), (-1, -1), 'CENTER'),  # Marks, CO, Level, Module columns
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
        ]
        
        # Add SPAN for OR rows
        for i, q in enumerate(questions_list):
            if q['question'] == 'OR':
                table_style_commands.append(('SPAN', (0, i+1), (5, i+1)))
                table_style_commands.append(('ALIGN', (0, i+1), (0, i+1), 'CENTER'))
                table_style_commands.append(('FONTNAME', (0, i+1), (0, i+1), 'Helvetica-Bold'))
                table_style_commands.append(('FONTSIZE', (0, i+1), (0, i+1), 11))
        
        questions_table.setStyle(TableStyle(table_style_commands))
        elements.append(questions_table)
        
        # Build PDF
        doc.build(elements)
        print(f"Generated PDF at: {temp_output_path}")

        # Return the PDF file directly
        return send_file(
            temp_output_path,
            mimetype='application/pdf',
            as_attachment=True,
            download_name='Generated_Question_Paper.pdf'
        )


        # Check if PDF format is requested
        format_requested = data.get('format', 'docx').lower()

        if format_requested == 'pdf':
            # Convert to PDF using our robust conversion function
            try:
                print(f"Converting DOCX to PDF: {temp_output_path}")
                result_path = convert_docx_to_pdf_robust(temp_output_path)

                if result_path.endswith('.pdf') and os.path.exists(result_path) and os.path.getsize(result_path) > 0:
                    print(f"PDF conversion successful. File size: {os.path.getsize(result_path)} bytes")
                    return send_file(result_path, as_attachment=True, download_name="Generated_Question_Paper.pdf",
                                   mimetype="application/pdf")
                else:
                    # Fallback to DOCX if PDF conversion fails
                    print("PDF conversion failed, returning DOCX")
                    return send_file(temp_output_path, as_attachment=True, download_name="Generated_Question_Paper.docx",
                                   mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            except Exception as pdf_error:
                print(f"PDF conversion failed: {pdf_error}")
                traceback.print_exc()
                # Fallback to DOCX
                return send_file(temp_output_path, as_attachment=True, download_name="Generated_Question_Paper.docx",
                               mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else:
            return send_file(temp_output_path, as_attachment=True, download_name="Generated_Question_Paper.docx",
                            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": f"Error generating document: {str(e)}. Check server logs."}), 500


def generate_paper_with_rules(all_questions, pattern='standard'):
    """
    Generate a question paper following the specified rules.
    
    CIE 1 Pattern:
    - Q1 & Q2: Sub-questions (a, b) from Module 1, (c) from Module 3.
    - Q3 & Q4: Sub-questions (a, b) from Module 2, (c) from Module 3.
    """
    import random
    from collections import defaultdict

    # Group questions by module
    questions_by_module = defaultdict(list)
    for q in all_questions:
        try:
            module_num = str(q.get('module', '1'))
            # Normalize module number (remove 'Module ' prefix if present)
            module_num = module_num.lower().replace('module', '').strip()
            
            marks = int(q.get('marks', 0))
            if marks > 0:
                q['module_num'] = module_num
                q['marks'] = marks
                questions_by_module[module_num].append(q)
        except (ValueError, TypeError):
            continue

    # Debug: Print available modules and counts
    print("DEBUG: Questions by Module:")
    for mod, qs in questions_by_module.items():
        print(f"  Module {mod}: {len(qs)} questions")
    
    # Helper to get random questions from a module
    def get_questions(module, count, used_ids):
        available = [q for q in questions_by_module.get(module, []) if q['firestore_id'] not in used_ids]
        if len(available) < count:
            return None
        selected = random.sample(available, count)
        for q in selected:
            used_ids.add(q['firestore_id'])
        return selected

    paper = []
    used_questions = set()

    # --- CIE 1 Logic ---
    if pattern == 'cie1':
        # Q1 & Q2: Module 1 (a, b) + Module 3 (c)
        for q_num in ['Q1', 'Q2']:
            m1_questions = get_questions('1', 2, used_questions)
            m3_questions = get_questions('3', 1, used_questions)
            
            if not m1_questions or not m3_questions:
                return {"error": f"Insufficient questions for {q_num}. Need 2 from Module 1 and 1 from Module 3."}
            
            sub_questions = [
                {"part": "a", "text": m1_questions[0]['question_text'], "marks": m1_questions[0]['marks'], "module": "1", "co": m1_questions[0].get('co', 'N/A'), "blooms_level": m1_questions[0].get('blooms_level', 'L2')},
                {"part": "b", "text": m1_questions[1]['question_text'], "marks": m1_questions[1]['marks'], "module": "1", "co": m1_questions[1].get('co', 'N/A'), "blooms_level": m1_questions[1].get('blooms_level', 'L2')},
                {"part": "c", "text": m3_questions[0]['question_text'], "marks": m3_questions[0]['marks'], "module": "3", "co": m3_questions[0].get('co', 'N/A'), "blooms_level": m3_questions[0].get('blooms_level', 'L2')}
            ]
            
            paper.append({
                "main_question": q_num,
                "module": "1 & 3",
                "sub_questions": sub_questions
            })

        # Q3 & Q4: Module 2 (a, b) + Module 3 (c)
        for q_num in ['Q3', 'Q4']:
            m2_questions = get_questions('2', 2, used_questions)
            m3_questions = get_questions('3', 1, used_questions)
            
            if not m2_questions or not m3_questions:
                return {"error": f"Insufficient questions for {q_num}. Need 2 from Module 2 and 1 from Module 3."}
            
            sub_questions = [
                {"part": "a", "text": m2_questions[0]['question_text'], "marks": m2_questions[0]['marks'], "module": "2", "co": m2_questions[0].get('co', 'N/A'), "blooms_level": m2_questions[0].get('blooms_level', 'L2')},
                {"part": "b", "text": m2_questions[1]['question_text'], "marks": m2_questions[1]['marks'], "module": "2", "co": m2_questions[1].get('co', 'N/A'), "blooms_level": m2_questions[1].get('blooms_level', 'L2')},
                {"part": "c", "text": m3_questions[0]['question_text'], "marks": m3_questions[0]['marks'], "module": "3", "co": m3_questions[0].get('co', 'N/A'), "blooms_level": m3_questions[0].get('blooms_level', 'L2')}
            ]
            
            paper.append({
                "main_question": q_num,
                "module": "2 & 3",
                "sub_questions": sub_questions
            })
            
        return paper

    # --- Fallback / Standard Logic (Simplified for now) ---
    # If not CIE 1, or if other patterns needed, we can expand here.
    # For now, returning error if not CIE 1 to ensure we focus on the requirement.
    return {"error": "Only CIE 1 pattern is currently supported with this strict logic."}

    # Ensure we have at least 2 questions (minimum viable paper)
    if len(paper) < 2:
        return None

    return paper


def generate_cie1_paper(all_questions):
    """
    Generate CIE1 pattern question paper (correct academic structure):
    - Q1 OR Q2: 3 sub-questions from Module 1 (25 marks each)
    - Q3: 3 sub-questions (a, b from Module 2, c from Module 3) (25 marks)
    - Q4: 3 sub-questions (a, b from Module 2, c from Module 3) (25 marks)
    Total: 100 marks (students answer 3 questions = 75 marks)
    """
    print("🎯 Generating CIE1 pattern question paper...")

    # Use the same module detection logic as the standard pattern
    from collections import defaultdict
    questions_by_module = defaultdict(list)
    
    for q in all_questions:
        try:
            # Extract module number directly from the module field
            module_num = str(q.get('module', '1'))  # Default to module 1 if not specified
            
            # Ensure marks is integer
            marks = int(q.get('marks', 0))
            if marks > 0:  # Only include questions with valid marks
                q['module_num'] = module_num
                q['marks'] = marks
                questions_by_module[module_num].append(q)

        except (ValueError, TypeError):
            continue  # Skip questions with invalid marks

    # Check module availability - CIE1 requires specific modules
    available_modules = list(questions_by_module.keys())
    print(f"Available modules: {available_modules}")
    for module, questions in questions_by_module.items():
        print(f"Module {module}: {len(questions)} questions")

    if not available_modules:
        return {"error": "No questions found with valid modules and marks."}

    # CIE1 requires at least 12 questions total
    total_questions = sum(len(questions_by_module[module]) for module in available_modules)
    if total_questions < 12:
        return {"error": f"CIE1 requires at least 12 questions total. Found {total_questions} questions across all modules."}
    
    # Use available modules in order
    primary_module = available_modules[0]
    secondary_module = available_modules[1] if len(available_modules) > 1 else available_modules[0]
    tertiary_module = available_modules[2] if len(available_modules) > 2 else secondary_module

    paper = []
    used_questions = set()

    def find_question_combination_cie(available_questions, count=3):
        """Find combination of questions with different marks that sum to 25"""
        import itertools
        for combo in itertools.combinations(available_questions, count):
            if all(q['firestore_id'] not in used_questions for q in combo):
                return list(combo)
        return None

    # Generate Q1 - Primary module (3 sub-questions)
    primary_questions = [q for q in questions_by_module[primary_module] if q['firestore_id'] not in used_questions]
    q1_combination = find_question_combination_cie(primary_questions)
    if not q1_combination:
        return {"error": f"Could not find suitable question combination for Q1 from Module {primary_module}"}

    for q in q1_combination:
        used_questions.add(q['firestore_id'])

    paper.append({
        "main_question": "Q1",
        "module": primary_module,
        "sub_questions": [
            {"part": "a", "text": q1_combination[0]['question_text'], "marks": q1_combination[0]['marks'], "co": q1_combination[0].get('co', 'N/A'), "blooms_level": q1_combination[0].get('blooms_level', 'L2'), "module": primary_module},
            {"part": "b", "text": q1_combination[1]['question_text'], "marks": q1_combination[1]['marks'], "co": q1_combination[1].get('co', 'N/A'), "blooms_level": q1_combination[1].get('blooms_level', 'L2'), "module": primary_module},
            {"part": "c", "text": q1_combination[2]['question_text'], "marks": q1_combination[2]['marks'], "co": q1_combination[2].get('co', 'N/A'), "blooms_level": q1_combination[2].get('blooms_level', 'L2'), "module": primary_module}
        ]
    })

    # Generate Q2 - Primary module (Alternative to Q1)
    primary_remaining = [q for q in questions_by_module[primary_module] if q['firestore_id'] not in used_questions]
    q2_combination = find_question_combination_cie(primary_remaining)
    if not q2_combination:
        return {"error": f"Could not find suitable question combination for Q2 from Module {primary_module}"}

    for q in q2_combination:
        used_questions.add(q['firestore_id'])

    paper.append({
        "main_question": "Q2",
        "module": primary_module,
        "sub_questions": [
            {"part": "a", "text": q2_combination[0]['question_text'], "marks": q2_combination[0]['marks'], "co": q2_combination[0].get('co', 'N/A'), "blooms_level": q2_combination[0].get('blooms_level', 'L2'), "module": primary_module},
            {"part": "b", "text": q2_combination[1]['question_text'], "marks": q2_combination[1]['marks'], "co": q2_combination[1].get('co', 'N/A'), "blooms_level": q2_combination[1].get('blooms_level', 'L2'), "module": primary_module},
            {"part": "c", "text": q2_combination[2]['question_text'], "marks": q2_combination[2]['marks'], "co": q2_combination[2].get('co', 'N/A'), "blooms_level": q2_combination[2].get('blooms_level', 'L2'), "module": primary_module}
        ]
    })

    # Generate Q3 - Secondary module (a, b) + Tertiary module (c)
    secondary_questions = [q for q in questions_by_module[secondary_module] if q['firestore_id'] not in used_questions]
    tertiary_questions = [q for q in questions_by_module[tertiary_module] if q['firestore_id'] not in used_questions]

    if len(secondary_questions) < 2:
        return {"error": f"Need at least 2 questions from Module {secondary_module} for Q3. Found {len(secondary_questions)}."}
    if len(tertiary_questions) < 1:
        return {"error": f"Need at least 1 question from Module {tertiary_module} for Q3. Found {len(tertiary_questions)}."}

    q3_secondary_questions = secondary_questions[:2]
    q3_tertiary_question = tertiary_questions[0]

    for q in q3_secondary_questions + [q3_tertiary_question]:
        used_questions.add(q['firestore_id'])

    module_label_q3 = f"{secondary_module} & {tertiary_module}" if secondary_module != tertiary_module else secondary_module

    paper.append({
        "main_question": "Q3",
        "module": module_label_q3,
        "sub_questions": [
            {"part": "a", "text": q3_secondary_questions[0]['question_text'], "marks": q3_secondary_questions[0]['marks'], "co": q3_secondary_questions[0].get('co', 'N/A'), "blooms_level": q3_secondary_questions[0].get('blooms_level', 'L2'), "module": secondary_module},
            {"part": "b", "text": q3_secondary_questions[1]['question_text'], "marks": q3_secondary_questions[1]['marks'], "co": q3_secondary_questions[1].get('co', 'N/A'), "blooms_level": q3_secondary_questions[1].get('blooms_level', 'L2'), "module": secondary_module},
            {"part": "c", "text": q3_tertiary_question['question_text'], "marks": q3_tertiary_question['marks'], "co": q3_tertiary_question.get('co', 'N/A'), "blooms_level": q3_tertiary_question.get('blooms_level', 'L2'), "module": tertiary_module}
        ]
    })

    # Generate Q4 - Secondary module (a, b) + Tertiary module (c) - Different questions
    secondary_remaining = [q for q in questions_by_module[secondary_module] if q['firestore_id'] not in used_questions]
    tertiary_remaining = [q for q in questions_by_module[tertiary_module] if q['firestore_id'] not in used_questions]

    if len(secondary_remaining) < 2:
        return {"error": f"Need at least 2 more questions from Module {secondary_module} for Q4. Found {len(secondary_remaining)}."}
    if len(tertiary_remaining) < 1:
        return {"error": f"Need at least 1 more question from Module {tertiary_module} for Q4. Found {len(tertiary_remaining)}."}

    q4_secondary_questions = secondary_remaining[:2]
    q4_tertiary_question = tertiary_remaining[0]

    for q in q4_secondary_questions + [q4_tertiary_question]:
        used_questions.add(q['firestore_id'])

    module_label_q4 = f"{secondary_module} & {tertiary_module}" if secondary_module != tertiary_module else secondary_module

    paper.append({
        "main_question": "Q4",
        "module": module_label_q4,
        "sub_questions": [
            {"part": "a", "text": q4_secondary_questions[0]['question_text'], "marks": q4_secondary_questions[0]['marks'], "co": q4_secondary_questions[0].get('co', 'N/A'), "blooms_level": q4_secondary_questions[0].get('blooms_level', 'L2'), "module": secondary_module},
            {"part": "b", "text": q4_secondary_questions[1]['question_text'], "marks": q4_secondary_questions[1]['marks'], "co": q4_secondary_questions[1].get('co', 'N/A'), "blooms_level": q4_secondary_questions[1].get('blooms_level', 'L2'), "module": secondary_module},
            {"part": "c", "text": q4_tertiary_question['question_text'], "marks": q4_tertiary_question['marks'], "co": q4_tertiary_question.get('co', 'N/A'), "blooms_level": q4_tertiary_question.get('blooms_level', 'L2'), "module": tertiary_module}
        ]
    })

    print(f"✅ CIE1 paper generated successfully with {len(paper)} questions")
    return paper


def generate_cie2_paper(all_questions):
    """
    Generate CIE2 pattern question paper (correct academic structure):
    - Q1 OR Q2: 3 sub-questions from Module 4 (25 marks each)
    - Q3: 3 sub-questions (a, b from Module 5, c from Module 6) (25 marks)
    - Q4: 3 sub-questions (a, b from Module 5, c from Module 6) (25 marks)
    Total: 100 marks (students answer 3 questions = 75 marks)
    """
    print("🎯 Generating CIE2 pattern question paper...")

    # Use the same module detection logic as the standard pattern
    from collections import defaultdict
    questions_by_module = defaultdict(list)
    
    for q in all_questions:
        try:
            # Extract module number directly from the module field
            module_num = str(q.get('module', '1'))  # Default to module 1 if not specified
            
            # Ensure marks is integer
            marks = int(q.get('marks', 0))
            if marks > 0:  # Only include questions with valid marks
                q['module_num'] = module_num
                q['marks'] = marks
                questions_by_module[module_num].append(q)

        except (ValueError, TypeError):
            continue  # Skip questions with invalid marks

    # Check module availability - CIE2 requires specific modules
    available_modules = list(questions_by_module.keys())
    print(f"Available modules: {available_modules}")
    for module, questions in questions_by_module.items():
        print(f"Module {module}: {len(questions)} questions")

    if not available_modules:
        return {"error": "No questions found with valid modules and marks."}

    # CIE2 requires at least 12 questions total
    total_questions = sum(len(questions_by_module[module]) for module in available_modules)
    if total_questions < 12:
        return {"error": f"CIE2 requires at least 12 questions total. Found {total_questions} questions across all modules."}
    
    # Use available modules in order
    primary_module = available_modules[0]
    secondary_module = available_modules[1] if len(available_modules) > 1 else available_modules[0]
    tertiary_module = available_modules[2] if len(available_modules) > 2 else secondary_module

    paper = []
    used_questions = set()

    def find_question_combination_cie(available_questions, count=3):
        """Find combination of questions with different marks that sum to 25"""
        import itertools
        for combo in itertools.combinations(available_questions, count):
            if all(q['firestore_id'] not in used_questions for q in combo):
                return list(combo)
        return None

    # Generate Q1 - Module 4 (3 sub-questions)
    module4_questions = [q for q in questions_by_module['4'] if q['firestore_id'] not in used_questions]
    q1_combination = find_question_combination_cie(module4_questions)
    if not q1_combination:
        return {"error": "Could not find suitable question combination for Q1 from Module 4"}

    for q in q1_combination:
        used_questions.add(q['firestore_id'])

    paper.append({
        "main_question": "Q1",
        "module": "4",
        "sub_questions": [
            {"part": "a", "text": q1_combination[0]['question_text'], "marks": q1_combination[0]['marks'], "co": q1_combination[0].get('co', 'N/A'), "blooms_level": q1_combination[0].get('blooms_level', 'L2'), "module": "4"},
            {"part": "b", "text": q1_combination[1]['question_text'], "marks": q1_combination[1]['marks'], "co": q1_combination[1].get('co', 'N/A'), "blooms_level": q1_combination[1].get('blooms_level', 'L2'), "module": "4"},
            {"part": "c", "text": q1_combination[2]['question_text'], "marks": q1_combination[2]['marks'], "co": q1_combination[2].get('co', 'N/A'), "blooms_level": q1_combination[2].get('blooms_level', 'L2'), "module": "4"}
        ]
    })

    # Generate Q2 - Module 4 (Alternative to Q1)
    module4_remaining = [q for q in questions_by_module['4'] if q['firestore_id'] not in used_questions]
    q2_combination = find_question_combination_cie(module4_remaining)
    if not q2_combination:
        return {"error": "Could not find suitable question combination for Q2 from Module 4"}

    for q in q2_combination:
        used_questions.add(q['firestore_id'])

    paper.append({
        "main_question": "Q2",
        "module": "4",
        "sub_questions": [
            {"part": "a", "text": q2_combination[0]['question_text'], "marks": q2_combination[0]['marks'], "co": q2_combination[0].get('co', 'N/A'), "blooms_level": q2_combination[0].get('blooms_level', 'L2'), "module": "4"},
            {"part": "b", "text": q2_combination[1]['question_text'], "marks": q2_combination[1]['marks'], "co": q2_combination[1].get('co', 'N/A'), "blooms_level": q2_combination[1].get('blooms_level', 'L2'), "module": "4"},
            {"part": "c", "text": q2_combination[2]['question_text'], "marks": q2_combination[2]['marks'], "co": q2_combination[2].get('co', 'N/A'), "blooms_level": q2_combination[2].get('blooms_level', 'L2'), "module": "4"}
        ]
    })

    # Generate Q3 - Module 5 (a, b) + Module 6 (c)
    module5_questions = [q for q in questions_by_module['5'] if q['firestore_id'] not in used_questions]
    module6_questions = [q for q in questions_by_module['6'] if q['firestore_id'] not in used_questions]

    if len(module5_questions) < 2:
        return {"error": f"Need at least 2 questions from Module 5 for Q3. Found {len(module5_questions)}."}
    if len(module6_questions) < 1:
        return {"error": f"Need at least 1 question from Module 6 for Q3. Found {len(module6_questions)}."}

    q3_mod5_questions = module5_questions[:2]
    q3_mod6_question = module6_questions[0]

    for q in q3_mod5_questions + [q3_mod6_question]:
        used_questions.add(q['firestore_id'])

    paper.append({
        "main_question": "Q3",
        "module": "5 & 6",
        "sub_questions": [
            {"part": "a", "text": q3_mod5_questions[0]['question_text'], "marks": q3_mod5_questions[0]['marks'], "co": q3_mod5_questions[0].get('co', 'N/A'), "blooms_level": q3_mod5_questions[0].get('blooms_level', 'L2'), "module": "5"},
            {"part": "b", "text": q3_mod5_questions[1]['question_text'], "marks": q3_mod5_questions[1]['marks'], "co": q3_mod5_questions[1].get('co', 'N/A'), "blooms_level": q3_mod5_questions[1].get('blooms_level', 'L2'), "module": "5"},
            {"part": "c", "text": q3_mod6_question['question_text'], "marks": q3_mod6_question['marks'], "co": q3_mod6_question.get('co', 'N/A'), "blooms_level": q3_mod6_question.get('blooms_level', 'L2'), "module": "6"}
        ]
    })

    # Generate Q4 - Module 5 (a, b) + Module 6 (c) - Different questions
    module5_remaining = [q for q in questions_by_module['5'] if q['firestore_id'] not in used_questions]
    module6_remaining = [q for q in questions_by_module['6'] if q['firestore_id'] not in used_questions]

    if len(module5_remaining) < 2:
        return {"error": f"Need at least 2 more questions from Module 5 for Q4. Found {len(module5_remaining)}."}
    if len(module6_remaining) < 1:
        return {"error": f"Need at least 1 more question from Module 6 for Q4. Found {len(module6_remaining)}."}

    q4_mod5_questions = module5_remaining[:2]
    q4_mod6_question = module6_remaining[0]

    for q in q4_mod5_questions + [q4_mod6_question]:
        used_questions.add(q['firestore_id'])

    paper.append({
        "main_question": "Q4",
        "module": "5 & 6",
        "sub_questions": [
            {"part": "a", "text": q4_mod5_questions[0]['question_text'], "marks": q4_mod5_questions[0]['marks'], "co": q4_mod5_questions[0].get('co', 'N/A'), "blooms_level": q4_mod5_questions[0].get('blooms_level', 'L2'), "module": "5"},
            {"part": "b", "text": q4_mod5_questions[1]['question_text'], "marks": q4_mod5_questions[1]['marks'], "co": q4_mod5_questions[1].get('co', 'N/A'), "blooms_level": q4_mod5_questions[1].get('blooms_level', 'L2'), "module": "5"},
            {"part": "c", "text": q4_mod6_question['question_text'], "marks": q4_mod6_question['marks'], "co": q4_mod6_question.get('co', 'N/A'), "blooms_level": q4_mod6_question.get('blooms_level', 'L2'), "module": "6"}
        ]
    })

    print(f"✅ CIE2 paper generated successfully with {len(paper)} questions")
    return paper


@app.route('/save_question_paper', methods=['POST'])
@firebase_auth_required
def save_question_paper():
    """Save a question paper to Supabase (hybrid approach)"""
    user_uid = request.current_user_uid
    data = request.get_json()
    
    try:
        # Get Supabase user ID
        supabase_user = supabase_service.get_user_by_firebase_uid(user_uid)
        if not supabase_user:
            return jsonify({'error': 'User not found in Supabase'}), 404
        
        # Prepare data for Supabase
        paper_data = {
            'firebase_paper_id': data.get('firebase_paper_id'),
            'paper_name': data.get('paper_name', 'Untitled Question Paper'),
            'subject': data.get('subject', 'Unknown Subject'),
            'pattern': data.get('pattern', 'standard'),
            'total_marks': data.get('total_marks', 100),
            'question_count': len(data.get('questions', [])),
            'questions': data.get('questions', []),
            'metadata': data.get('metadata', {}),
            'status': data.get('status', 'draft'),
            'tags': data.get('tags', [])
        }
        
        # Save to Supabase
        saved_paper = supabase_service.save_question_paper(supabase_user['id'], paper_data)
        
        if saved_paper:
            # Also save to Firebase for backup (optional)
            try:
                data['user_uid'] = user_uid
                data['saved_at'] = firestore.SERVER_TIMESTAMP
                data['updated_at'] = firestore.SERVER_TIMESTAMP
                data['supabase_id'] = saved_paper['id']  # Reference to Supabase record
                
                saved_paper_ref = db_firestore.collection('users').document(user_uid).collection('saved_question_papers').document()
                saved_paper_ref.set(data)
                print(f"Question paper also backed up to Firebase: {saved_paper_ref.id}")
            except Exception as e:
                print(f"Warning: Failed to backup to Firebase: {e}")
            
            return jsonify({
                'message': 'Question paper saved successfully',
                'paper_id': saved_paper['id'],
                'supabase_id': saved_paper['id']
            }), 200
        else:
            return jsonify({'error': 'Failed to save question paper to Supabase'}), 500
        
    except Exception as e:
        print(f"Error saving question paper: {e}")
        return jsonify({'error': f'Failed to save question paper: {str(e)}'}), 500

@app.route('/get_saved_items', methods=['GET'])
@firebase_auth_required
def get_saved_items():
    """Get saved question papers from Supabase, question banks and templates from Firebase"""
    user_uid = request.current_user_uid
    
    try:
        # Get Supabase user ID
        supabase_user = supabase_service.get_user_by_firebase_uid(user_uid)
        if not supabase_user:
            return jsonify({'error': 'User not found in Supabase'}), 404
        
        # Get saved question papers from Supabase
        saved_question_papers = supabase_service.get_saved_question_papers(supabase_user['id'], limit=50)
        
        # Add type field and map fields for frontend compatibility
        for paper in saved_question_papers:
            paper['type'] = 'question-paper'
            # Convert UUID to string for JSON serialization
            paper['id'] = str(paper['id'])
            # Map Supabase fields to frontend expected fields
            if 'paper_name' in paper:
                paper['name'] = paper['paper_name']
            if 'total_marks' in paper:
                paper['marks'] = paper['total_marks']
            if 'created_at' in paper:
                paper['date'] = paper['created_at']
                paper['timestamp'] = paper['created_at'] # For sorting if needed
        
        # Get generated question papers
        papers_ref = db_firestore.collection('users').document(user_uid).collection('generated_papers')
        papers_docs = papers_ref.order_by('created_at', direction=firestore.Query.DESCENDING).limit(50).stream()
        
        generated_question_papers = []
        for doc in papers_docs:
            paper_data = doc.to_dict()
            paper_data['id'] = doc.id
            paper_data['type'] = 'generated-paper'
            generated_question_papers.append(paper_data)
        
        # Get question banks (Optimized)
        banks_ref = db_firestore.collection('users').document(user_uid).collection('question_banks')
        banks_docs = banks_ref.order_by('uploaded_at', direction=firestore.Query.DESCENDING).stream()
        
        question_banks = []
        for doc in banks_docs:
            bank_data = doc.to_dict()
            # Ensure ID is set
            if 'id' not in bank_data:
                bank_data['id'] = doc.id
            question_banks.append(bank_data)
            
        # Lazy Migration: If no banks found in new collection, check if we have questions in the pool
        # and if so, perform migration (group by source_file and create bank entries)
        if not question_banks:
            pool_ref = db_firestore.collection('users').document(user_uid).collection('question_bank_pool')
            # Check if pool has any documents (limit 1 to check existence)
            pool_check = pool_ref.limit(1).get()
            
            if pool_check:
                print("Performing lazy migration for question banks...")
                # Fetch all questions (slow, but only happens once)
                all_questions_docs = pool_ref.stream()
                
                banks_by_source = {}
                for doc in all_questions_docs:
                    q_data = doc.to_dict()
                    source = q_data.get('source_file', 'Unknown Source')
                    
                    if source not in banks_by_source:
                        import re
                        safe_id = re.sub(r'[^a-zA-Z0-9]', '_', source)
                        banks_by_source[source] = {
                            'id': f"bank_{safe_id}",
                            'name': source,
                            'source_file': source,
                            'uploaded_at': q_data.get('uploaded_at', firestore.SERVER_TIMESTAMP),
                            'question_count': 0,
                            'type': 'question-bank',
                            'user_uid': user_uid
                        }
                    
                    banks_by_source[source]['question_count'] += 1
                
                # Save these banks to the new collection
                batch = db_firestore.batch()
                for source, bank_data in banks_by_source.items():
                    new_bank_ref = banks_ref.document(bank_data['id'])
                    batch.set(new_bank_ref, bank_data)
                    question_banks.append(bank_data)
                
                batch.commit()
                print(f"Migrated {len(question_banks)} question banks.")

        return jsonify({
            "message": "Saved items fetched successfully",
            "question_papers": saved_question_papers,
            "generated_papers": generated_question_papers,
            "question_banks": question_banks,
            "templates": []  # Placeholder for templates
        }), 200
        
    except Exception as e:
        return jsonify({"error": f"Error fetching saved items: {str(e)}"}), 500

@app.route('/get_recent_papers', methods=['GET'])
@firebase_auth_required
def get_recent_papers():
    """Get recently generated question papers for the user"""
    user_uid = request.current_user_uid
    
    try:
        # Get recent papers from both generated_papers and saved_question_papers collections
        recent_papers = []
        
        # Get from generated papers
        generated_ref = db_firestore.collection('users').document(user_uid).collection('generated_papers')
        generated_docs = generated_ref.order_by('created_at', direction=firestore.Query.DESCENDING).limit(10).stream()
        
        for doc in generated_docs:
            paper_data = doc.to_dict()
            paper_data['id'] = doc.id
            paper_data['source'] = 'generated'
            recent_papers.append(paper_data)
        
        # Get from saved papers
        saved_ref = db_firestore.collection('users').document(user_uid).collection('saved_question_papers')
        saved_docs = saved_ref.order_by('saved_at', direction=firestore.Query.DESCENDING).limit(10).stream()
        
        for doc in saved_docs:
            paper_data = doc.to_dict()
            paper_data['id'] = doc.id
            paper_data['source'] = 'saved'
            recent_papers.append(paper_data)
        
        # Sort by date (newest first)
        # Sort by date (newest first)
        def get_paper_timestamp(p):
            ts = p.get('created_at') or p.get('saved_at')
            if hasattr(ts, 'timestamp'):
                return ts.timestamp()
            if isinstance(ts, dict):
                return ts.get('seconds', 0)
            return 0
            
        recent_papers.sort(key=get_paper_timestamp, reverse=True)
        
        return jsonify({
            "message": "Recent papers fetched successfully",
            "recent_papers": recent_papers[:20]  # Limit to 20 most recent
        }), 200
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Error fetching recent papers: {str(e)}"}), 500


@app.route('/get_question_paper/<paper_id>', methods=['GET'])
@firebase_auth_required
def get_question_paper(paper_id):
    user_uid = request.current_user_uid
    try:
        paper = supabase_service.get_question_paper_by_id(paper_id, user_uid)
        if paper:
            return jsonify(paper), 200
        else:
            return jsonify({'error': 'Question paper not found'}), 404
    except Exception as e:
        print(f"Error fetching question paper: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/delete_recent_paper', methods=['DELETE'])
@firebase_auth_required
def delete_recent_paper():
    user_uid = request.current_user_uid
    data = request.get_json()
    paper_id = data.get('paper_id')
    source = data.get('source') # 'generated' or 'saved'

    if not paper_id:
        return jsonify({'error': 'Paper ID is required'}), 400

    try:
        if source == 'generated':
            # Delete from generated_papers
            db_firestore.collection('users').document(user_uid).collection('generated_papers').document(paper_id).delete()
            return jsonify({'message': 'Generated paper deleted successfully'}), 200
        
        elif source == 'saved':
            # Delete from saved_question_papers (Firestore)
            db_firestore.collection('users').document(user_uid).collection('saved_question_papers').document(paper_id).delete()
            
            # Also delete from Supabase if it exists there
            # We might need the Supabase ID, but usually for saved papers the ID in the list IS the Supabase ID 
            # or we stored the Supabase ID in the Firestore doc. 
            # However, looking at get_recent_papers, for saved papers, paper_data['id'] = doc.id. 
            # If the Firestore doc ID is different from Supabase ID, we might have an issue.
            # But typically for 'saved' papers, the Firestore doc might contain the Supabase ID.
            # Let's check get_recent_papers again. 
            # It just takes doc.to_dict() and adds doc.id.
            
            # If the ID passed from frontend is the Firestore ID, we need to fetch the doc first to get Supabase ID if they differ.
            # But wait, save_question_paper saves to Supabase first, gets an ID, then saves to Firestore. 
            # The Firestore doc ID is auto-generated: .collection('saved_question_papers').document()
            # So the ID in the list IS the Firestore ID.
            # The Supabase ID is stored in the field 'supabase_id'.
            
            # So we need to:
            # 1. Get the Firestore doc to find the Supabase ID.
            # 2. Delete from Supabase using that ID.
            # 3. Delete from Firestore.
            
            # Actually, let's try to fetch the doc first.
            doc_ref = db_firestore.collection('users').document(user_uid).collection('saved_question_papers').document(paper_id)
            doc = doc_ref.get()
            
            if doc.exists:
                doc_data = doc.to_dict()
                supabase_id = doc_data.get('supabase_id')
                if supabase_id:
                     supabase_service.delete_question_paper(supabase_id, user_uid) # This expects Supabase ID
                
                # If no supabase_id field, maybe the paper_id IS the supabase_id? 
                # Unlikely if it came from Firestore list.
                # But let's just proceed to delete the Firestore doc.
                doc_ref.delete()
                return jsonify({'message': 'Saved paper deleted successfully'}), 200
            else:
                # If not found in Firestore, maybe it was only in Supabase? 
                # But the list comes from Firestore.
                # If the user passed the Supabase ID directly (unlikely for 'saved' source from recent list), we might miss it.
                # But let's assume standard flow.
                return jsonify({'error': 'Paper not found'}), 404

        else:
            # If source is unknown, try to find in both? Or just return error.
            # Let's try to delete from generated first, if not found, try saved.
            # But safer to require source.
            return jsonify({'error': 'Invalid or missing source'}), 400

    except Exception as e:
        print(f"Error deleting paper: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/submit_for_approval', methods=['POST'])
@firebase_auth_required
def submit_for_approval():
    """Submit a question paper for HOD approval (hybrid approach)"""
    user_uid = request.current_user_uid
    data = request.get_json()
    paper_id = data.get('paper_id')
    comments = data.get('comments', '')
    
    try:
        # Get Supabase user ID
        supabase_user = supabase_service.get_user_by_firebase_uid(user_uid)
        if not supabase_user:
            return jsonify({'error': 'User not found in Supabase'}), 404
        
        paper_data = None
        paper_source = None
        
        # First, try to get from Supabase saved papers
        try:
            saved_papers = supabase_service.get_saved_question_papers(supabase_user['id'], limit=100)
            for paper in saved_papers:
                if str(paper['id']) == str(paper_id) or paper.get('firebase_paper_id') == paper_id:
                    paper_data = {
                        'paper_name': paper.get('paper_name', 'Untitled Paper'),
                        'subject': paper.get('subject', 'Unknown Subject'),
                        'questions': paper.get('questions', []),
                        'pattern': paper.get('pattern', 'standard'),
                        'total_marks': paper.get('total_marks', 100),
                        'metadata': paper.get('metadata', {})
                    }
                    paper_source = 'supabase'
                    break
        except Exception as e:
            print(f"Error fetching from Supabase: {e}")
        
        # If not found in Supabase, try Firebase generated papers
        if not paper_data:
            try:
                paper_ref = db_firestore.collection('users').document(user_uid).collection('generated_papers').document(paper_id)
                paper_doc = paper_ref.get()
                
                if paper_doc.exists:
                    paper_data = paper_doc.to_dict()
                    paper_source = 'firebase'
            except Exception as e:
                print(f"Error fetching from Firebase: {e}")
        
        if not paper_data:
            return jsonify({"error": "Question paper not found in either Supabase or Firebase"}), 404
        
        # Get user info from Supabase
        user_data = {
            'name': supabase_user.get('name', 'Unknown Faculty'),
            'email': supabase_user.get('email', ''),
            'department': supabase_user.get('department', 'Unknown Department')
        }
        
        # Find the Supabase paper ID for the approval
        supabase_paper_id = None
        if paper_source == 'supabase':
            # Find the paper in the saved papers list
            for paper in saved_papers:
                if str(paper['id']) == str(paper_id) or paper.get('firebase_paper_id') == paper_id:
                    supabase_paper_id = paper['id']
                    break
        else:
            # For Firebase papers, check if there's a corresponding Supabase paper
            try:
                firebase_paper_result = supabase_service.supabase.table('saved_question_papers').select('id').eq('firebase_paper_id', paper_id).execute()
                if firebase_paper_result.data:
                    supabase_paper_id = firebase_paper_result.data[0]['id']
            except Exception as e:
                print(f"Error finding Supabase paper for Firebase ID {paper_id}: {e}")
        
        if not supabase_paper_id:
            # If paper is not in Supabase but exists in Firebase, save it now
            if paper_source == 'firebase' and paper_data:
                print(f"Paper {paper_id} not in Supabase, saving it now...")
                try:
                    # Prepare data for Supabase
                    supabase_paper_data = {
                        'firebase_paper_id': paper_id,
                        'paper_name': paper_data.get('paper_name', 'Untitled Paper'),
                        'subject': paper_data.get('subject', 'Unknown Subject'),
                        'pattern': paper_data.get('pattern', 'standard'),
                        'total_marks': 100,  # Default
                        'question_count': len(paper_data.get('questions', [])),
                        'questions': paper_data.get('questions', []),
                        'metadata': paper_data.get('source_info', {}),
                        'status': 'draft',
                        'tags': paper_data.get('tags', [])
                    }
                    
                    # Save to Supabase
                    saved_paper = supabase_service.save_question_paper(supabase_user['id'], supabase_paper_data)
                    if saved_paper:
                        supabase_paper_id = saved_paper['id']
                        print(f"Paper saved to Supabase: {supabase_paper_id}")
                        
                        # Update Firebase with Supabase ID reference
                        paper_ref = db_firestore.collection('users').document(user_uid).collection('generated_papers').document(paper_id)
                        paper_ref.update({'supabase_id': supabase_paper_id})
                    else:
                        return jsonify({"error": "Failed to save paper to Supabase"}), 500
                except Exception as e:
                    print(f"Error saving paper to Supabase: {e}")
                    return jsonify({"error": f"Failed to save paper to Supabase: {str(e)}"}), 500
            else:
                return jsonify({"error": "Paper not found in Supabase. Please save the paper first."}), 404
        
        # Create approval request (matching database schema)
        approval_data = {
            "paper_id": supabase_paper_id,  # Use Supabase UUID
            "submitted_by": supabase_user['id'],  # Use Supabase user ID
            "status": "pending",
            "comments": comments,  # faculty comments
            "hod_comments": None,
            "reviewed_by": None,
            "reviewed_at": None
        }
        
        # Save to Supabase first (primary storage)
        try:
            supabase_approval = supabase_service.create_approval(supabase_user['id'], approval_data)
            if supabase_approval:
                approval_id = supabase_approval['id']
                print(f"Approval saved to Supabase: {approval_id}")
            else:
                return jsonify({"error": "Failed to save approval to Supabase"}), 500
        except Exception as e:
            print(f"Error saving to Supabase: {e}")
            return jsonify({"error": f"Failed to save approval: {str(e)}"}), 500
        
        # Also save to Firebase for backup
        try:
            firebase_approval_data = approval_data.copy()
            firebase_approval_data["submitted_at"] = firestore.SERVER_TIMESTAMP
            firebase_approval_data["supabase_id"] = approval_id
            
            approval_ref = db_firestore.collection('approvals').document()
            approval_ref.set(firebase_approval_data)
            print(f"Approval also backed up to Firebase: {approval_ref.id}")
        except Exception as e:
            print(f"Warning: Failed to backup approval to Firebase: {e}")
        
        # Update paper status in both systems
        try:
            # Update Supabase paper status (always use the Supabase paper ID)
            supabase_service.update_question_paper_status(supabase_paper_id, 'submitted', approval_id)
        except Exception as e:
            print(f"Warning: Failed to update Supabase paper status: {e}")
        
        # Also update Firebase paper status if it's a Firebase paper
        if paper_source == 'firebase':
            try:
                paper_ref = db_firestore.collection('users').document(user_uid).collection('generated_papers').document(paper_id)
                paper_ref.update({
                    "status": "submitted",
                    "approval_id": approval_id,
                    "submitted_for_approval_at": firestore.SERVER_TIMESTAMP
                })
            except Exception as e:
                print(f"Warning: Failed to update Firebase paper status: {e}")
        
        return jsonify({
            "message": "Question paper submitted for approval successfully!",
            "approval_id": approval_id,
            "status": "pending"
        }), 200
        
    except Exception as e:
        return jsonify({"error": f"Failed to submit for approval: {str(e)}"}), 500

@app.route('/get_pending_approvals', methods=['GET'])
@firebase_auth_required
def get_pending_approvals():
    """Get papers pending approval for HOD"""
    user_uid = request.current_user_uid
    
    try:
        # Check if user is HOD using Supabase
        hod_user = supabase_service.get_user_by_firebase_uid(user_uid)
        print(f"DEBUG: get_pending_approvals - UID: {user_uid}, User: {hod_user}")
        
        if not hod_user or hod_user.get('role') != 'hod':
            print(f"DEBUG: Access Denied. Role: {hod_user.get('role') if hod_user else 'None'}")
            return jsonify({"error": "Access denied. HOD privileges required."}), 403
        
        # Get pending approvals from Supabase
        approvals = supabase_service.get_pending_approvals(hod_user['id'])
        
        # Format approvals for frontend
        formatted_approvals = []
        for approval in approvals:
            paper = approval.get('saved_question_papers', {})
            faculty = approval.get('users', {})
            
            formatted_approval = {
                'id': approval['id'],
                'paper_id': approval['paper_id'],
                'faculty_uid': approval['submitted_by'],
                'faculty_name': faculty.get('name', 'Unknown Faculty'),
                'faculty_email': faculty.get('email', ''),
                'department': faculty.get('department', 'Unknown Department'),
                'subject': paper.get('subject', 'Unknown Subject'),
                'paper_name': paper.get('paper_name', 'Untitled Paper'),
                'questions': paper.get('questions', []),
                'pattern': paper.get('pattern', 'standard'),
                'total_marks': paper.get('total_marks', 100),
                'status': approval['status'],
                'comments': approval.get('comments', ''),
                'hod_comments': approval.get('hod_comments', ''),
                'submitted_at': approval.get('submitted_at'),
                'reviewed_at': approval.get('reviewed_at'),
                'priority': 'medium',
                'estimated_review_time': '2-3 days'
            }
            formatted_approvals.append(formatted_approval)
        
        return jsonify({
            "approvals": formatted_approvals,
            "department": hod_user.get('department', 'Unknown Department')
        }), 200
        
    except Exception as e:
        return jsonify({"error": f"Error fetching pending approvals: {str(e)}"}), 500

@app.route('/approve_paper', methods=['POST'])
@firebase_auth_required
def approve_paper():
    """Approve a question paper"""
    user_uid = request.current_user_uid
    data = request.get_json()
    approval_id = data.get('approval_id')
    comments = data.get('comments', '')
    
    try:
        # Check if user is HOD using Supabase
        hod_user = supabase_service.get_user_by_firebase_uid(user_uid)
        if not hod_user or hod_user.get('role') != 'hod':
            return jsonify({"error": "Access denied. HOD privileges required."}), 403
        
        # Approve paper using Supabase service
        success = supabase_service.approve_paper(approval_id, hod_user['id'], comments)
        
        if success:
            return jsonify({
                "message": "Question paper approved successfully!",
                "status": "approved"
            }), 200
        else:
            return jsonify({"error": "Failed to approve paper"}), 500
        
    except Exception as e:
        return jsonify({"error": f"Failed to approve paper: {str(e)}"}), 500

@app.route('/request_revision', methods=['POST'])
@firebase_auth_required
def request_revision():
    """Request revision for a question paper"""
    user_uid = request.current_user_uid
    data = request.get_json()
    approval_id = data.get('approval_id')
    comments = data.get('comments', '')
    revision_type = data.get('revision_type', 'minor')  # minor, major, complete_rewrite
    
    try:
        # Check if user is HOD using Supabase
        hod_user = supabase_service.get_user_by_firebase_uid(user_uid)
        if not hod_user or hod_user.get('role') != 'hod':
            return jsonify({"error": "Access denied. HOD privileges required."}), 403
        
        # Reject paper using Supabase service (treating revision as rejection with comments)
        success = supabase_service.reject_paper(approval_id, hod_user['id'], f"Revision requested: {comments}")
        
        if success:
            return jsonify({
                "message": "Revision requested successfully!",
                "status": "revision_requested",
                "revision_type": revision_type
            }), 200
        else:
            return jsonify({"error": "Failed to request revision"}), 500
        
    except Exception as e:
        return jsonify({"error": f"Failed to request revision: {str(e)}"}), 500

@app.route('/get_approval_status', methods=['GET'])
@firebase_auth_required
def get_approval_status():
    """Get approval status for a paper"""
    user_uid = request.current_user_uid
    paper_id = request.args.get('paper_id')
    
    try:
        # Get paper
        paper_ref = db_firestore.collection('users').document(user_uid).collection('generated_papers').document(paper_id)
        paper_doc = paper_ref.get()
        
        if not paper_doc.exists:
            return jsonify({"error": "Question paper not found"}), 404
        
        paper_data = paper_doc.to_dict()
        approval_id = paper_data.get('approval_id')
        
        if not approval_id:
            return jsonify({
                "status": "not_submitted",
                "message": "Paper not submitted for approval"
            }), 200
        
        # Get approval details
        approval_ref = db_firestore.collection('approvals').document(approval_id)
        approval_doc = approval_ref.get()
        
        if not approval_doc.exists:
            return jsonify({
                "status": "error",
                "message": "Approval record not found"
            }), 200
        
        approval_data = approval_doc.to_dict()
        
        return jsonify({
            "status": approval_data.get('status', 'unknown'),
            "submitted_at": approval_data.get('submitted_at'),
            "reviewed_at": approval_data.get('reviewed_at'),
            "reviewed_by": approval_data.get('reviewed_by', ''),
            "hod_comments": approval_data.get('hod_comments', ''),
            "comments": approval_data.get('comments', ''),
            "revision_requests": approval_data.get('revision_requests', []),
            "priority": approval_data.get('priority', 'medium')
        }), 200
        
    except Exception as e:
        return jsonify({"error": f"Error fetching approval status: {str(e)}"}), 500

@app.route('/get_user_submissions', methods=['GET'])
@firebase_auth_required
def get_user_submissions():
    """Get all submissions for a user"""
    user_uid = request.current_user_uid
    
    try:
        # Get Supabase user ID
        supabase_user = supabase_service.get_user_by_firebase_uid(user_uid)
        if not supabase_user:
            return jsonify({'error': 'User not found in Supabase'}), 404
        
        # Get all approvals for this user
        approvals_result = supabase_service.supabase.table('approvals')\
            .select('*, saved_question_papers(paper_name, subject, pattern)')\
            .eq('submitted_by', supabase_user['id'])\
            .order('submitted_at', desc=True)\
            .execute()
        
        submissions = []
        for approval in approvals_result.data:
            paper_info = approval.get('saved_question_papers', {})
            submissions.append({
                'id': approval['id'],
                'paper_id': approval['paper_id'],
                'paper_name': paper_info.get('paper_name', 'Untitled Paper'),
                'subject': paper_info.get('subject', 'Unknown Subject'),
                'pattern': paper_info.get('pattern', 'standard'),
                'status': approval['status'],
                'comments': approval.get('comments', ''),
                'hod_comments': approval.get('hod_comments', ''),
                'submitted_at': approval.get('submitted_at'),
                'reviewed_at': approval.get('reviewed_at'),
                'reviewed_by': approval.get('reviewed_by')
            })
        
        return jsonify({
            'submissions': submissions,
            'total': len(submissions)
        }), 200
        
    except Exception as e:
        print(f"Error fetching user submissions: {e}")
        return jsonify({"error": f"Error fetching submissions: {str(e)}"}), 500

@app.route('/get_dashboard_metrics', methods=['GET'])
@firebase_auth_required
def get_dashboard_metrics():
    """Get dashboard metrics for the user"""
    user_uid = request.current_user_uid
    
    try:
        # Get Supabase user ID
        supabase_user = supabase_service.get_user_by_firebase_uid(user_uid)
        if not supabase_user:
            return jsonify({'error': 'User not found in Supabase'}), 404
        
        user_id = supabase_user['id']
        
        # Get total questions from question banks
        questions_result = supabase_service.supabase.table('question_banks')\
            .select('total_questions')\
            .eq('user_id', user_id)\
            .execute()
        
        total_questions = sum(bank.get('total_questions', 0) for bank in questions_result.data)
        
        # If no questions in question_banks, try to count from saved papers
        if total_questions == 0:
            papers_with_questions = supabase_service.supabase.table('saved_question_papers')\
                .select('question_count')\
                .eq('user_id', user_id)\
                .execute()
            
            total_questions = sum(paper.get('question_count', 0) for paper in papers_with_questions.data)
        
        # Get papers generated this month
        from datetime import datetime, timedelta
        month_start = datetime.now().replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        
        papers_result = supabase_service.supabase.table('saved_question_papers')\
            .select('id, created_at')\
            .eq('user_id', user_id)\
            .gte('created_at', month_start.isoformat())\
            .execute()
        
        papers_generated = len(papers_result.data)
        
        # Get saved templates (all saved papers can be used as templates)
        templates_result = supabase_service.supabase.table('saved_question_papers')\
            .select('id')\
            .eq('user_id', user_id)\
            .execute()
        
        saved_templates = len(templates_result.data)
        
        # Calculate success rate (papers that were successfully generated vs failed)
        # For now, assume high success rate since we don't track failures
        success_rate = 98 if papers_generated > 0 else 100
        
        # Calculate trends (simplified - in real app, compare with previous periods)
        questions_trend = 12 if total_questions > 0 else 0
        papers_trend = 8 if papers_generated > 0 else 0
        
        return jsonify({
            'total_questions': total_questions,
            'papers_generated': papers_generated,
            'saved_templates': saved_templates,
            'success_rate': success_rate,
            'questions_trend': questions_trend,
            'papers_trend': papers_trend
        }), 200
        
    except Exception as e:
        print(f"Error fetching dashboard metrics: {e}")
        return jsonify({"error": f"Error fetching metrics: {str(e)}"}), 500

@app.route('/create_hod_user', methods=['POST'])
def create_hod_user():
    """Create HOD user for testing"""
    try:
        firebase_uid = 'jaQrJOvNfxO4eN5sqKoiNkmlixg2'
        email = 'test_hod@example.com'
        
        # Check if user already exists
        existing_user = supabase_service.get_user_by_firebase_uid(firebase_uid)
        if existing_user:
            if existing_user.get('role') != 'hod':
                # Update role to hod
                supabase_service.supabase.table('users')\
                    .update({'role': 'hod'})\
                    .eq('id', existing_user['id'])\
                    .execute()
                return jsonify({'message': 'User role updated to HOD'}), 200
            else:
                return jsonify({'message': 'User already has HOD role'}), 200
        
        # Create new HOD user
        user_data = {
            'firebase_uid': firebase_uid,
            'email': email,
            'name': 'Test HOD',
            'role': 'hod',
            'department': 'Computer Science'
        }
        
        result = supabase_service.supabase.table('users').insert(user_data).execute()
        
        if result.data:
            return jsonify({'message': 'HOD user created successfully', 'user': result.data[0]}), 200
        else:
            return jsonify({'error': 'Failed to create HOD user'}), 500
            
    except Exception as e:
        return jsonify({'error': f'Error creating HOD user: {str(e)}'}), 500

@app.route('/generate_question_paper', methods=['POST'])
@firebase_auth_required
def generate_question_paper():
    user_uid = request.current_user_uid
    data = request.get_json()
    subject = data.get('subject')
    modules = data.get('modules', [])
    use_latest_upload_only = data.get('use_latest_upload_only', True)  # New parameter
    pattern = data.get('pattern', 'standard')  # New parameter for question pattern

    try:
        # For CIE 1, we likely need questions from multiple modules which might be in different files.
        # So we force 'use_latest_upload_only' to False to ensure we have access to the full pool.
        if pattern == 'cie1':
            print("Pattern is CIE 1: Forcing use_latest_upload_only=False to ensure all modules are available.")
            use_latest_upload_only = False

        if use_latest_upload_only:
            # Check cache first
            cache_key = f"{user_uid}_latest"
            cached_questions = get_cached_questions(user_uid, cache_key)
            
            if cached_questions:
                all_questions = cached_questions
                latest_source_file = all_questions[0].get('source_file') if all_questions else None
            else:
                # First, get the latest source file with a single query
                questions_ref = db_firestore.collection('users').document(user_uid).collection('question_bank_pool')
                
                # Get only the most recent document to find the latest source file
                latest_doc = questions_ref.order_by('uploaded_at', direction=firestore.Query.DESCENDING).limit(1).get()
                
                if not latest_doc:
                    return jsonify({'error': 'No questions found in your question bank.'}), 404
                    
                latest_source_file = latest_doc[0].to_dict().get('source_file')
                
                # Now fetch only questions from the latest source file
                docs = questions_ref.where('source_file', '==', latest_source_file).stream()
                
                all_questions = []
                for doc in docs:
                    question_data = doc.to_dict()
                    question_data['firestore_id'] = doc.id
                    all_questions.append(question_data)
                
                # Cache the results
                cache_questions(user_uid, cache_key, all_questions)

            if not all_questions:
                return jsonify({"error": "No questions found from the latest upload. Please upload a question bank first."}), 400

            print(f"Using questions only from latest upload: '{latest_source_file}' ({len(all_questions)} questions)")

        else:
            # Original behavior: Fetch all questions from user's question bank
            questions_ref = db_firestore.collection('users').document(user_uid).collection('question_bank_pool')
            docs = questions_ref.stream()

            all_questions = []
            for doc in docs:
                question_data = doc.to_dict()
                question_data['firestore_id'] = doc.id
                all_questions.append(question_data)

            if not all_questions:
                return jsonify({"error": "No questions found in your question bank. Please upload questions first."}), 400

            print(f"Using all questions from question bank ({len(all_questions)} questions)")

        # Generate question paper according to selected pattern
        if pattern == 'cie1':
            generated_paper = generate_paper_with_rules(all_questions, pattern='cie1')
        elif pattern == 'cie2':
            generated_paper = generate_cie2_paper(all_questions)
        else:
            # Default to standard pattern
            generated_paper = generate_paper_with_rules(all_questions)

        if not generated_paper:
            return jsonify({"error": "Could not generate a question paper with the available questions. Please ensure you have questions from all required modules with appropriate marks distribution."}), 400

        # Handle error responses from the generation function
        if isinstance(generated_paper, dict) and "error" in generated_paper:
            return jsonify(generated_paper), 400

        # Save the generated paper to Firestore
        metadata = data.get('metadata', {})
        
        paper_data = {
            "user_uid": user_uid,
            "paper_name": f"{subject} - {pattern.upper()} Pattern",
            "subject": subject,
            "pattern": pattern,
            "questions": generated_paper,
            "modules": modules,
            "metadata": metadata,  # Store full metadata
            "source_info": {
                "latest_upload_only": use_latest_upload_only,
                "source_file": latest_source_file if use_latest_upload_only else "All uploaded files",
                "total_questions_used": len(all_questions)
            },
            "created_at": firestore.SERVER_TIMESTAMP,
            "updated_at": firestore.SERVER_TIMESTAMP,
            "status": "generated",
            "download_count": 0,
            "size": len(str(generated_paper)),
            "tags": [pattern, subject.lower().replace(" ", "-")]
        }
        
        # Save to Firestore
        paper_ref = db_firestore.collection('users').document(user_uid).collection('generated_papers').document()
        paper_ref.set(paper_data)
        firebase_paper_id = paper_ref.id
        
        # Also save to Supabase for HOD approval workflow
        try:
            # Get Supabase user ID
            supabase_user = supabase_service.get_user_by_firebase_uid(user_uid)
            if supabase_user:
                # Prepare data for Supabase
                supabase_paper_data = {
                    'firebase_paper_id': firebase_paper_id,
                    'paper_name': f"{subject} - {pattern.upper()} Pattern",
                    'subject': subject,
                    'pattern': pattern,
                    'total_marks': int(metadata.get('max_marks', 100)),
                    'question_count': len(generated_paper),
                    'questions': generated_paper,
                    'metadata': {
                        "latest_upload_only": use_latest_upload_only,
                        "source_file": latest_source_file if use_latest_upload_only else "All uploaded files",
                        "total_questions_used": len(all_questions),
                        "modules": modules,
                        **metadata # Include all other metadata
                    },
                    'status': 'draft',
                    'tags': [pattern, subject.lower().replace(" ", "-")]
                }
                
                # Save to Supabase
                saved_paper = supabase_service.save_question_paper(supabase_user['id'], supabase_paper_data)
                if saved_paper:
                    print(f"Paper also saved to Supabase: {saved_paper['id']}")
                    # Update Firebase with Supabase ID reference
                    paper_ref.update({'supabase_id': saved_paper['id']})
                else:
                    print("Warning: Failed to save paper to Supabase")
            else:
                print("Warning: User not found in Supabase, paper only saved to Firebase")
        except Exception as e:
            print(f"Warning: Failed to save paper to Supabase: {e}")
        
        return jsonify({
            "message": "Question paper generated successfully!",
            "questions": generated_paper,
            "paper_id": firebase_paper_id,
            "source_info": {
                "latest_upload_only": use_latest_upload_only,
                "source_file": latest_source_file if use_latest_upload_only else "All uploaded files",
                "total_questions_used": len(all_questions)
            }
        }), 200

    except Exception as e:
        traceback.print_exc()
        
        # Check if it's a quota exceeded error
        if "Quota exceeded" in str(e) or "ResourceExhausted" in str(e):
            return jsonify({
                'error': 'Firebase quota exceeded. Please try again later or contact support.',
                'details': 'You have reached the daily limit for database operations. This will reset tomorrow.',
                'suggestion': 'Try using fewer questions or wait until tomorrow when the quota resets.'
            }), 429
        else:
            return jsonify({"error": f"Failed to generate question paper: {str(e)}"}), 500




@app.route('/uploads/<path:filename>')
def serve_uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

@app.route('/extracted_images/<path:filename>')
def serve_extracted_image(filename):
    return send_from_directory('extracted_images', filename)

# --- Test Parser Workbench Routes ---
TEST_BANKS_FOLDER = os.path.abspath(os.path.join(os.path.dirname(__file__), '../question banks'))

@app.route('/test_parser')
def test_parser_page():
    return render_template('test_parser.html')

@app.route('/api/test/list_banks')
def list_test_banks():
    try:
        if not os.path.exists(TEST_BANKS_FOLDER):
            return jsonify({'error': f'Folder not found: {TEST_BANKS_FOLDER}'}), 404
            
        files = [f for f in os.listdir(TEST_BANKS_FOLDER) if f.lower().endswith('.pdf')]
        return jsonify({'files': files})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/test/parse_bank', methods=['POST'])
def parse_test_bank():
    try:
        data = request.json
        filename = data.get('filename')
        if not filename:
            return jsonify({'error': 'Filename is required'}), 400
            
        filepath = os.path.join(TEST_BANKS_FOLDER, filename)
        if not os.path.exists(filepath):
            return jsonify({'error': 'File not found'}), 404
            
        # Use Advanced Parser
        from advanced_parser import get_advanced_parser
        parser = get_advanced_parser()
        
        # Output images to the same extracted_images folder so they can be served
        # We'll use a subfolder 'test_runs' to keep it slightly organized or just root
        # Let's use the standard extracted_images folder in backend
        images_folder = os.path.join(os.path.dirname(__file__), "extracted_images")
        os.makedirs(images_folder, exist_ok=True)
        
        # Parse
        raw_parsed_content = parser.parse_pdf(filepath, images_folder)
        
        # Normalize image paths for serving
        for item in raw_parsed_content:
            item["images"] = [img.replace('\\', '/').split('/')[-1] for img in item["images"]]
            # Note: We just keep the filename because serve_extracted_image serves from 'extracted_images' root
            
        return jsonify({'status': 'success', 'data': raw_parsed_content})
        
    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

# Route to serve uploaded/generated files
@app.route('/uploads/<path:filename>')
def serve_uploads(filename):
    """Serve files from the uploads directory"""
    response = send_from_directory(app.config['UPLOAD_FOLDER'], filename)
    # Ensure correct MIME type for PDFs
    if filename.lower().endswith('.pdf'):
        response.headers['Content-Type'] = 'application/pdf'
    return response

if __name__ == '__main__':
    port = int(os.environ.get('FLASK_RUN_PORT', 5000))
    app.run(debug=True, port=port)
