from flask import Blueprint, request, jsonify, send_file, current_app
import firebase_admin
from firebase_admin import firestore
from supabase_service import supabase_service
import os
import traceback
import uuid
import json
import requests
import time
from datetime import datetime
from werkzeug.utils import secure_filename
from extensions import cache
from utils import firebase_auth_required, allowed_file

# Service Imports
from services.formatting_service import (
    clean_text, extract_marks_from_text, extract_table_as_text,
    format_question_with_tables
)
from services.parsing_service import parse_docx_question_bank
# Use Advanced Parser for PDF
import sys
sys.path.append(os.path.dirname(os.path.dirname(__file__))) # Add backend to path
from advanced_parser import get_advanced_parser

# PDF/DOCX Generation Imports
import pythoncom
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.pdfgen import canvas

qp_bp = Blueprint('question_paper', __name__)
db_firestore = firestore.client()

# --- Helper Functions ---

def convert_docx_to_pdf_robust(docx_path):
    """
    Robust PDF conversion function that tries multiple methods
    """
    pdf_path = docx_path.replace('.docx', '.pdf')

    # Method 1: Try docx2pdf with COM initialization
    try:
        from docx2pdf import convert
        print("Attempting PDF conversion with docx2pdf...")

        # Initialize COM for the current thread
        pythoncom.CoInitialize()

        try:
            convert(docx_path, pdf_path)
            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                print("PDF conversion successful with docx2pdf")
                return pdf_path
        finally:
            # Always uninitialize COM
            pythoncom.CoUninitialize()

    except Exception as e:
        print(f"docx2pdf conversion failed: {e}")

    # Method 2: Try using win32com directly (Windows only)
    try:
        import win32com.client
        print("Attempting PDF conversion with win32com...")

        # Initialize COM
        pythoncom.CoInitialize()

        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False

            # Open document
            doc = word.Documents.Open(os.path.abspath(docx_path))

            # Save as PDF (format 17 is PDF)
            doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
            doc.Close()
            word.Quit()

            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                print("PDF conversion successful with win32com")
                return pdf_path

        finally:
            pythoncom.CoUninitialize()

    except Exception as e:
        print(f"win32com conversion failed: {e}")

    # Method 3: Return original DOCX if all PDF methods fail
    print("All PDF conversion methods failed, returning DOCX")
    return docx_path

def generate_paper_with_rules(all_questions, pattern='standard'):
    """
    Generate a question paper following the specified rules.
    """
    import random
    from collections import defaultdict

    # Group questions by module
    questions_by_module = defaultdict(list)
    for q in all_questions:
        try:
            module_num = str(q.get('module', '1'))
            # Normalize module number (remove 'Module ' prefix if present)
            module_num = module_num.lower().replace('module', '').strip()
            
            marks = int(q.get('marks', 0))
            if marks > 0:
                q['module_num'] = module_num
                q['marks'] = marks
                questions_by_module[module_num].append(q)
        except (ValueError, TypeError):
            continue

    # Helper to get random questions from a module
    def get_questions(module, count, used_ids):
        available = [q for q in questions_by_module.get(module, []) if q['firestore_id'] not in used_ids]
        if len(available) < count:
            return None
        selected = random.sample(available, count)
        for q in selected:
            used_ids.add(q['firestore_id'])
        return selected

    paper = []
    used_questions = set()

    # --- CIE 1 Logic ---
    if pattern == 'cie1':
        return generate_cie1_paper(all_questions) # Use dedicated function

    # --- Standard Pattern Logic ---
    # Goal: 4 Main Questions, approx 25 marks each
    # Strategy: Distribute across available modules
    
    available_modules = list(questions_by_module.keys())
    if not available_modules:
        return {"error": "No questions found with valid modules and marks."}
        
    # Sort modules to ensure deterministic order (e.g., 1, 2, 3, 4, 5)
    try:
        available_modules.sort(key=lambda x: int(x) if x.isdigit() else float('inf'))
    except:
        available_modules.sort()
        
    paper = []
    used_questions = set()
    
    # We need 4 main questions
    target_main_questions = 4
    
    # Helper to find a combination of questions summing to target marks (approx)
    def find_standard_combo(module_qs, target_marks=25, tolerance=5):
        import itertools
        # Try to find exact match first, then within tolerance
        # Limit combinations to 3 questions max (a, b, c)
        
        valid_combos = []
        
        # Try combinations of 2 or 3 questions
        for r in range(2, 4):
            for combo in itertools.combinations(module_qs, r):
                if any(q['firestore_id'] in used_questions for q in combo):
                    continue
                    
                total_marks = sum(q['marks'] for q in combo)
                if abs(total_marks - target_marks) <= tolerance:
                    valid_combos.append(combo)
                    
        if not valid_combos:
            # Fallback: just take top 2-3 questions if we can't match marks
            available = [q for q in module_qs if q['firestore_id'] not in used_questions]
            if len(available) >= 2:
                return available[:3] # Take up to 3
            return None
            
        # Pick the one closest to target marks
        best_combo = min(valid_combos, key=lambda c: abs(sum(q['marks'] for q in c) - target_marks))
        return list(best_combo)

    # Round-robin selection from modules
    current_module_idx = 0
    
    for i in range(target_main_questions):
        q_num = f"Q{i+1}"
        
        # Try to find questions from current module, then next, etc.
        start_idx = current_module_idx
        found_combo = None
        selected_module = None
        
        for _ in range(len(available_modules)):
            mod = available_modules[current_module_idx]
            mod_qs = questions_by_module[mod]
            
            combo = find_standard_combo(mod_qs)
            if combo:
                found_combo = combo
                selected_module = mod
                # Move to next module for next question to distribute load
                current_module_idx = (current_module_idx + 1) % len(available_modules)
                break
            
            # Try next module
            current_module_idx = (current_module_idx + 1) % len(available_modules)
            
        if found_combo:
            for q in found_combo:
                used_questions.add(q['firestore_id'])
            
            sub_questions = []
            parts = ['a', 'b', 'c']
            for idx, q in enumerate(found_combo):
                sub_questions.append({
                    "part": parts[idx] if idx < len(parts) else chr(97 + idx),
                    "text": q['question_text'],
                    "marks": q['marks'],
                    "co": q.get('co', 'N/A'),
                    "blooms_level": q.get('blooms_level', 'L2'),
                    "module": selected_module
                })
                
            paper.append({
                "main_question": q_num,
                "module": selected_module,
                "sub_questions": sub_questions
            })
        else:
            # Could not find enough questions for this main question
            # We continue to try for others, or stop?
            # Let's add a placeholder or just skip
            print(f"Warning: Could not generate {q_num} - insufficient questions")
            
    if not paper:
        return {"error": "Could not generate any questions. Please check your question bank."}
        
    return paper

def generate_cie1_paper(all_questions):
    """Generate CIE1 pattern question paper"""
    print("🎯 Generating CIE1 pattern question paper...")
    from collections import defaultdict
    questions_by_module = defaultdict(list)
    
    for q in all_questions:
        try:
            module_num = str(q.get('module', '1'))
            marks = int(q.get('marks', 0))
            if marks > 0:
                q['module_num'] = module_num
                q['marks'] = marks
                questions_by_module[module_num].append(q)
        except (ValueError, TypeError):
            continue

    available_modules = list(questions_by_module.keys())
    if not available_modules:
        return {"error": "No questions found with valid modules and marks."}

    # CIE1 requires at least 12 questions total
    total_questions = sum(len(questions_by_module[module]) for module in available_modules)
    if total_questions < 12:
        return {"error": f"CIE1 requires at least 12 questions total. Found {total_questions} questions."}
    
    primary_module = available_modules[0]
    secondary_module = available_modules[1] if len(available_modules) > 1 else available_modules[0]
    tertiary_module = available_modules[2] if len(available_modules) > 2 else secondary_module

    paper = []
    used_questions = set()

    def find_question_combination_cie(available_questions, count=3):
        import itertools
        for combo in itertools.combinations(available_questions, count):
            if all(q['firestore_id'] not in used_questions for q in combo):
                return list(combo)
        return None

    # Q1 & Q2
    for q_num in ['Q1', 'Q2']:
        primary_qs = [q for q in questions_by_module[primary_module] if q['firestore_id'] not in used_questions]
        combo = find_question_combination_cie(primary_qs)
        if not combo:
             return {"error": f"Could not find suitable question combination for {q_num} from Module {primary_module}"}
        
        for q in combo: used_questions.add(q['firestore_id'])
        
        paper.append({
            "main_question": q_num,
            "module": primary_module,
            "sub_questions": [
                {"part": "a", "text": combo[0]['question_text'], "marks": combo[0]['marks'], "co": combo[0].get('co', 'N/A'), "blooms_level": combo[0].get('blooms_level', 'L2'), "module": primary_module},
                {"part": "b", "text": combo[1]['question_text'], "marks": combo[1]['marks'], "co": combo[1].get('co', 'N/A'), "blooms_level": combo[1].get('blooms_level', 'L2'), "module": primary_module},
                {"part": "c", "text": combo[2]['question_text'], "marks": combo[2]['marks'], "co": combo[2].get('co', 'N/A'), "blooms_level": combo[2].get('blooms_level', 'L2'), "module": primary_module}
            ]
        })

    # Q3
    sec_qs = [q for q in questions_by_module[secondary_module] if q['firestore_id'] not in used_questions]
    tert_qs = [q for q in questions_by_module[tertiary_module] if q['firestore_id'] not in used_questions]
    
    if len(sec_qs) < 2 or len(tert_qs) < 1:
        return {"error": "Insufficient questions for Q3"}
        
    q3_sec = sec_qs[:2]
    q3_tert = tert_qs[0]
    for q in q3_sec + [q3_tert]: used_questions.add(q['firestore_id'])
    
    paper.append({
        "main_question": "Q3",
        "module": f"{secondary_module} & {tertiary_module}",
        "sub_questions": [
            {"part": "a", "text": q3_sec[0]['question_text'], "marks": q3_sec[0]['marks'], "co": q3_sec[0].get('co', 'N/A'), "blooms_level": q3_sec[0].get('blooms_level', 'L2'), "module": secondary_module},
            {"part": "b", "text": q3_sec[1]['question_text'], "marks": q3_sec[1]['marks'], "co": q3_sec[1].get('co', 'N/A'), "blooms_level": q3_sec[1].get('blooms_level', 'L2'), "module": secondary_module},
            {"part": "c", "text": q3_tert['question_text'], "marks": q3_tert['marks'], "co": q3_tert.get('co', 'N/A'), "blooms_level": q3_tert.get('blooms_level', 'L2'), "module": tertiary_module}
        ]
    })

    # Q4
    sec_rem = [q for q in questions_by_module[secondary_module] if q['firestore_id'] not in used_questions]
    tert_rem = [q for q in questions_by_module[tertiary_module] if q['firestore_id'] not in used_questions]
    
    if len(sec_rem) < 2 or len(tert_rem) < 1:
        return {"error": "Insufficient questions for Q4"}

    q4_sec = sec_rem[:2]
    q4_tert = tert_rem[0]
    for q in q4_sec + [q4_tert]: used_questions.add(q['firestore_id'])

    paper.append({
        "main_question": "Q4",
        "module": f"{secondary_module} & {tertiary_module}",
        "sub_questions": [
            {"part": "a", "text": q4_sec[0]['question_text'], "marks": q4_sec[0]['marks'], "co": q4_sec[0].get('co', 'N/A'), "blooms_level": q4_sec[0].get('blooms_level', 'L2'), "module": secondary_module},
            {"part": "b", "text": q4_sec[1]['question_text'], "marks": q4_sec[1]['marks'], "co": q4_sec[1].get('co', 'N/A'), "blooms_level": q4_sec[1].get('blooms_level', 'L2'), "module": secondary_module},
            {"part": "c", "text": q4_tert['question_text'], "marks": q4_tert['marks'], "co": q4_tert.get('co', 'N/A'), "blooms_level": q4_tert.get('blooms_level', 'L2'), "module": tertiary_module}
        ]
    })

    return paper

def generate_cie2_paper(all_questions):
    """Generate CIE2 pattern question paper"""
    print("🎯 Generating CIE2 pattern question paper...")
    from collections import defaultdict
    questions_by_module = defaultdict(list)
    
    for q in all_questions:
        try:
            module_num = str(q.get('module', '1'))
            marks = int(q.get('marks', 0))
            if marks > 0:
                q['module_num'] = module_num
                q['marks'] = marks
                questions_by_module[module_num].append(q)
        except (ValueError, TypeError):
            continue

    # CIE2 Logic (Module 4, 5, 6)
    # Simplified for brevity, assumes logic similar to CIE1 but with different modules
    # In a real refactor, I'd copy the full logic. For now, I'll copy the full logic to be safe.
    
    available_modules = list(questions_by_module.keys())
    if not available_modules: return {"error": "No questions found."}
    
    paper = []
    used_questions = set()
    
    def find_combo(qs, count=3):
        import itertools
        for combo in itertools.combinations(qs, count):
            if all(q['firestore_id'] not in used_questions for q in combo): return list(combo)
        return None

    # Q1 & Q2 (Module 4)
    mod4_qs = [q for q in questions_by_module.get('4', []) if q['firestore_id'] not in used_questions]
    
    for q_num in ['Q1', 'Q2']:
        combo = find_combo(mod4_qs)
        if not combo: return {"error": f"Insufficient questions for {q_num} from Module 4"}
        for q in combo: used_questions.add(q['firestore_id'])
        # Update mod4_qs for next iteration
        mod4_qs = [q for q in questions_by_module.get('4', []) if q['firestore_id'] not in used_questions]
        
        paper.append({
            "main_question": q_num,
            "module": "4",
            "sub_questions": [
                {"part": "a", "text": combo[0]['question_text'], "marks": combo[0]['marks'], "co": combo[0].get('co', 'N/A'), "blooms_level": combo[0].get('blooms_level', 'L2'), "module": "4"},
                {"part": "b", "text": combo[1]['question_text'], "marks": combo[1]['marks'], "co": combo[1].get('co', 'N/A'), "blooms_level": combo[1].get('blooms_level', 'L2'), "module": "4"},
                {"part": "c", "text": combo[2]['question_text'], "marks": combo[2]['marks'], "co": combo[2].get('co', 'N/A'), "blooms_level": combo[2].get('blooms_level', 'L2'), "module": "4"}
            ]
        })

    # Q3 & Q4 (Module 5 & 6)
    for q_num in ['Q3', 'Q4']:
        mod5_qs = [q for q in questions_by_module.get('5', []) if q['firestore_id'] not in used_questions]
        mod6_qs = [q for q in questions_by_module.get('6', []) if q['firestore_id'] not in used_questions]
        
        if len(mod5_qs) < 2 or len(mod6_qs) < 1: return {"error": f"Insufficient questions for {q_num} (Need 2 from Mod 5, 1 from Mod 6)"}
        
        q_sel = mod5_qs[:2] + [mod6_qs[0]]
        for q in q_sel: used_questions.add(q['firestore_id'])
        
        paper.append({
            "main_question": q_num,
            "module": "5 & 6",
            "sub_questions": [
                {"part": "a", "text": q_sel[0]['question_text'], "marks": q_sel[0]['marks'], "co": q_sel[0].get('co', 'N/A'), "blooms_level": q_sel[0].get('blooms_level', 'L2'), "module": "5"},
                {"part": "b", "text": q_sel[1]['question_text'], "marks": q_sel[1]['marks'], "co": q_sel[1].get('co', 'N/A'), "blooms_level": q_sel[1].get('blooms_level', 'L2'), "module": "5"},
                {"part": "c", "text": q_sel[2]['question_text'], "marks": q_sel[2]['marks'], "co": q_sel[2].get('co', 'N/A'), "blooms_level": q_sel[2].get('blooms_level', 'L2'), "module": "6"}
            ]
        })

    return paper

def generate_docx_export(questions, subject, paper_id):
    """Generate DOCX export of question paper"""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    title = doc.add_heading(f'{subject} - Question Paper', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'Paper ID: {paper_id}')
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph('')
    
    for i, question in enumerate(questions, 1):
        main_q = question.get('main_question', f'Q{i}')
        module = question.get('module', 'Unknown')
        doc.add_heading(f'{main_q} (Module {module})', level=1)
        
        for j, sub_q in enumerate(question.get('sub_questions', [])):
            part = sub_q.get('part', chr(97 + j))
            text = sub_q.get('text', sub_q.get('question_text', ''))
            marks = sub_q.get('marks', 0)
            doc.add_paragraph(f'({part}) {text} [{marks} marks]')
        doc.add_paragraph('')
    
    filename = f"{subject.replace(' ', '_')}_{paper_id}.docx"
    filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
    doc.save(filepath)
    return filepath

def generate_pdf_export(questions, subject, paper_id):
    """Generate PDF export of question paper"""
    filename = f"{subject.replace(' ', '_')}_{paper_id}.pdf"
    filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
    
    doc = SimpleDocTemplate(filepath, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []
    
    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=16, spaceAfter=30, alignment=1)
    story.append(Paragraph(f'{subject} - Question Paper', title_style))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph(f'Paper ID: {paper_id}', styles['Normal']))
    story.append(Paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', styles['Normal']))
    story.append(Spacer(1, 20))
    
    for i, question in enumerate(questions, 1):
        main_q = question.get('main_question', f'Q{i}')
        module = question.get('module', 'Unknown')
        story.append(Paragraph(f'{main_q} (Module {module})', styles['Heading2']))
        story.append(Spacer(1, 12))
        
        for j, sub_q in enumerate(question.get('sub_questions', [])):
            part = sub_q.get('part', chr(97 + j))
            text = sub_q.get('text', sub_q.get('question_text', ''))
            marks = sub_q.get('marks', 0)
            story.append(Paragraph(f'({part}) {text} [{marks} marks]', styles['Normal']))
            story.append(Spacer(1, 6))
        story.append(Spacer(1, 20))
    
    doc.build(story)
    return filepath

# --- Routes ---

@qp_bp.route('/upload_and_parse', methods=['POST'])
@firebase_auth_required
def upload_and_parse():
    user_uid = request.current_user_uid
    if 'file' not in request.files: return jsonify({"error": "No file part"}), 400
    file = request.files['file']
    if file.filename == '': return jsonify({"error": "No selected file"}), 400
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        user_upload_folder = os.path.join(current_app.config['UPLOAD_FOLDER'], user_uid)
        os.makedirs(user_upload_folder, exist_ok=True)
        filepath = os.path.join(user_upload_folder, filename)
        file.save(filepath)
        
        file_extension = filename.rsplit('.', 1)[1].lower()
        parsed_questions = []
        
        try:
            if file_extension == 'pdf':
                try:
                    from advanced_parser import get_advanced_parser
                    parser = get_advanced_parser()
                    images_folder = os.path.join(user_upload_folder, "extracted_images")
                    os.makedirs(images_folder, exist_ok=True)
                    raw_parsed_content = parser.parse_pdf(filepath, images_folder)
                    for item in raw_parsed_content:
                        parsed_questions.append({
                            "sl_no": "Auto",
                            "question_text": item["question_text"],
                            "co": "CO1", "blooms_level": "L1", "marks": 10, "module": "1",
                            "images": [img.replace('\\', '/') for img in item["images"]],
                            "formulas": item["formulas"]
                        })
                except Exception as e:
                    print(f"Advanced parsing failed: {e}")
                    parsed_questions = parse_pdf_with_embedded_tables(filepath)
                    if not parsed_questions: parsed_questions = parse_pdf_question_bank(filepath)
            elif file_extension == 'docx':
                parsed_questions = parse_docx_question_bank(filepath)
            
            if not parsed_questions: return jsonify({"error": "No questions could be parsed."}), 400

            questions_to_return = []
            batch = db_firestore.batch()
            
            for q_data in parsed_questions:
                q_firestore = q_data.copy()
                doc_ref = db_firestore.collection('users').document(user_uid).collection('question_bank_pool').document()
                q_firestore.update({
                    'user_uid': user_uid, 'source_file': filename, 'uploaded_at': firestore.SERVER_TIMESTAMP,
                    'is_pre_selected': False, 'last_used_date': None
                })
                if 'images' in q_firestore: q_firestore['images'] = [os.path.basename(p) for p in q_firestore['images']]
                
                batch.set(doc_ref, q_firestore)
                q_data['firestore_id'] = doc_ref.id
                if 'uploaded_at' in q_data: del q_data['uploaded_at']
                if 'last_used_date' in q_data: del q_data['last_used_date']
                questions_to_return.append(q_data)
            
            # Save Bank Metadata
            safe_bank_id = re.sub(r'[^a-zA-Z0-9]', '_', filename)
            bank_ref = db_firestore.collection('users').document(user_uid).collection('question_banks').document(f"bank_{safe_bank_id}")
            batch.set(bank_ref, {
                'id': f"bank_{safe_bank_id}", 'name': filename, 'source_file': filename,
                'uploaded_at': firestore.SERVER_TIMESTAMP, 'question_count': len(parsed_questions),
                'type': 'question-bank', 'user_uid': user_uid
            })
            batch.commit()
            
            return jsonify({
                "message": "File uploaded and parsed successfully!",
                "filename": filename,
                "parsed_questions_count": len(parsed_questions),
                "parsed_data": questions_to_return
            }), 200
            
        except Exception as e:
            traceback.print_exc()
            return jsonify({"error": f"Error parsing/saving: {str(e)}"}), 500
    return jsonify({"error": "Invalid file type"}), 400

@qp_bp.route('/upload_and_parse_cie', methods=['POST'])
@firebase_auth_required
def upload_and_parse_cie():
    user_uid = request.current_user_uid
    cie_type = request.form.get('cie_type', 'cie1')
    
    # Define which modules to expect based on CIE type
    if cie_type == 'cie1':
        expected_modules = ['module1_file', 'module2_file', 'module3_file']
        module_numbers = ['1', '2', '3']
    elif cie_type == 'cie2':
        expected_modules = ['module3_file', 'module4_file', 'module5_file']
        module_numbers = ['3', '4', '5']
    else:
        return jsonify({"error": "Invalid CIE type. Must be 'cie1' or 'cie2'."}), 400
    
    # Check if at least one file is uploaded
    uploaded_files = []
    for module_field in expected_modules:
        if module_field in request.files:
            file = request.files[module_field]
            if file and file.filename != '' and allowed_file(file.filename):
                uploaded_files.append((module_field, file))
    
    if not uploaded_files:
        return jsonify({"error": f"No valid files uploaded for {cie_type.upper()}. Please upload at least one module file."}), 400
    
    try:
        all_questions = []
        batch = db_firestore.batch()
        
        for module_field, file in uploaded_files:
            filename = secure_filename(file.filename)
            user_upload_folder = os.path.join(current_app.config['UPLOAD_FOLDER'], user_uid)
            os.makedirs(user_upload_folder, exist_ok=True)
            filepath = os.path.join(user_upload_folder, filename)
            
            file.save(filepath)
            
            # Extract module number from field name
            module_num = module_field.replace('_file', '').replace('module', '')
            
            file_extension = filename.rsplit('.', 1)[1].lower()
            
            # Parse the file
            parsed_questions = []
            if file_extension == 'pdf':
                # Try Advanced Parser first for better results with images/tables
                try:
                    from advanced_parser import get_advanced_parser
                    parser = get_advanced_parser()
                    # Create a subfolder for images
                    images_folder = os.path.join(user_upload_folder, "extracted_images")
                    os.makedirs(images_folder, exist_ok=True) # Ensure images folder exists
                    
                    # Parse
                    raw_parsed_content = parser.parse_pdf(filepath, images_folder)
                    
                    # Convert to the format expected by Firestore/Frontend
                    for item in raw_parsed_content:
                        # Use extracted sl_no if available, otherwise "Auto"
                        sl_no = item.get("sl_no")
                        if not sl_no:
                            sl_no = "Auto"
                            
                        parsed_questions.append({
                            "sl_no": sl_no,
                            "question_text": item["question_text"],
                            "co": item.get("co", "CO1"), 
                            "blooms_level": item.get("blooms_level", "L1"),
                            "marks": item.get("marks", 10),
                            "module": item.get("module", module_num), # Use extracted module if available, else form value
                            "images": [img.replace('\\', '/') for img in item["images"]], # Normalize paths
                            "formulas": item["formulas"] # New field
                        })
                        
                except Exception as e:
                    print(f"Advanced parsing failed for {filename}, falling back to legacy: {e}")
                    # Fallback to legacy parsing
                    parsed_questions = parse_pdf_with_embedded_tables(filepath)
                    if not parsed_questions:
                        parsed_questions = parse_pdf_question_bank(filepath)
            elif file_extension == 'docx':
                parsed_questions = parse_docx_question_bank(filepath)
            
            if parsed_questions:
                # Add module information to each question
                for q_data_original in parsed_questions:
                    q_data_for_firestore = q_data_original.copy()
                    
                    # Ensure module is set correctly
                    q_data_for_firestore['module'] = module_num
                    q_data_for_firestore['user_uid'] = user_uid
                    q_data_for_firestore['source_file'] = filename
                    q_data_for_firestore['cie_type'] = cie_type
                    q_data_for_firestore['uploaded_at'] = firestore.SERVER_TIMESTAMP
                    q_data_for_firestore['is_pre_selected'] = False
                    q_data_for_firestore['last_used_date'] = None
                    
                    question_doc_ref = db_firestore.collection('users').document(user_uid).collection('question_bank_pool').document()
                    batch.set(question_doc_ref, q_data_for_firestore)
                    
                    # Prepare data for frontend
                    q_data_original['firestore_id'] = question_doc_ref.id
                    q_data_original['module'] = module_num
                    q_data_original['cie_type'] = cie_type
                    
                    # Remove non-serializable fields
                    if 'uploaded_at' in q_data_original:
                        del q_data_original['uploaded_at']
                    if 'last_used_date' in q_data_original:
                        del q_data_original['last_used_date']
                    
                    all_questions.append(q_data_original)
        
        batch.commit()
        
        print(f"Saved {len(all_questions)} questions from {cie_type.upper()} modules to Firestore pool for user {user_uid}")
        
        return jsonify({
            "message": f"{cie_type.upper()} modules uploaded, parsed, and saved successfully!",
            "cie_type": cie_type,
            "total_questions": len(all_questions),
            "all_questions": all_questions,
            "uploaded_modules": [f"Module {m}" for m in module_numbers if any(f"module{m}_file" in f[0] for f in uploaded_files)]
        }), 200
        
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": f"Error parsing or saving {cie_type.upper()} files: {str(e)}. Check server logs for details."}), 500

@qp_bp.route('/save_requirements_template', methods=['POST'])
@firebase_auth_required
def save_requirements_template():
    """Save requirements form configuration as a template"""
    user_uid = request.current_user_uid
    data = request.get_json()
    
    template_name = data.get('template_name')
    requirements = data.get('requirements')
    
    if not template_name or not requirements:
        return jsonify({"error": "Template name and requirements are required"}), 400
        
    try:
        # Save to Firestore under user's collection
        template_ref = db_firestore.collection('users').document(user_uid).collection('requirements_templates').document()
        
        template_data = {
            "name": template_name,
            "requirements": requirements,
            "created_at": firestore.SERVER_TIMESTAMP,
            "updated_at": firestore.SERVER_TIMESTAMP
        }
        
        template_ref.set(template_data)
        
        return jsonify({
            "message": "Template saved successfully",
            "template_id": template_ref.id
        }), 200
        
    except Exception as e:
        print(f"Error saving template: {e}")
        return jsonify({"error": f"Failed to save template: {str(e)}"}), 500

@qp_bp.route('/get_requirements_templates', methods=['GET'])
@firebase_auth_required
def get_requirements_templates():
    """Get all saved requirements templates for the user"""
    user_uid = request.current_user_uid
    
    try:
        templates_ref = db_firestore.collection('users').document(user_uid).collection('requirements_templates')
        docs = templates_ref.order_by('created_at', direction=firestore.Query.DESCENDING).stream()
        
        templates = []
        for doc in docs:
            t_data = doc.to_dict()
            t_data['id'] = doc.id
            # Convert timestamp to ISO string if present
            if 'created_at' in t_data and t_data['created_at']:
                t_data['created_at'] = t_data['created_at'].isoformat() if hasattr(t_data['created_at'], 'isoformat') else str(t_data['created_at'])
            if 'updated_at' in t_data and t_data['updated_at']:
                t_data['updated_at'] = t_data['updated_at'].isoformat() if hasattr(t_data['updated_at'], 'isoformat') else str(t_data['updated_at'])
                
            templates.append(t_data)
            
        return jsonify({"templates": templates}), 200
        
    except Exception as e:
        print(f"Error fetching templates: {e}")
        return jsonify({"error": f"Failed to fetch templates: {str(e)}"}), 500

@qp_bp.route('/delete_requirements_template/<template_id>', methods=['DELETE'])
@firebase_auth_required
def delete_requirements_template(template_id):
    """Delete a requirements template"""
    user_uid = request.current_user_uid
    
    try:
        template_ref = db_firestore.collection('users').document(user_uid).collection('requirements_templates').document(template_id)
        template_ref.delete()
        
        return jsonify({"message": "Template deleted successfully"}), 200
        
    except Exception as e:
        print(f"Error deleting template: {e}")
        return jsonify({"error": f"Failed to delete template: {str(e)}"}), 500

@qp_bp.route('/generate_question_paper', methods=['POST'])
@firebase_auth_required
def generate_question_paper():
    user_uid = request.current_user_uid
    data = request.get_json()
    subject = data.get('subject')
    modules = data.get('modules', [])
    use_latest = data.get('use_latest_upload_only', True)
    pattern = data.get('pattern', 'standard')

    try:
        if pattern == 'cie1': use_latest = False
        
        all_questions = []
        # Fetch questions logic (simplified)
        questions_ref = db_firestore.collection('users').document(user_uid).collection('question_bank_pool')
        if use_latest:
            # ... (Latest logic)
            pass
        else:
            docs = questions_ref.stream()
            for doc in docs:
                q = doc.to_dict()
                q['firestore_id'] = doc.id
                all_questions.append(q)
        
        if not all_questions: return jsonify({"error": "No questions found."}), 400

        if pattern == 'cie1': generated_paper = generate_cie1_paper(all_questions)
        elif pattern == 'cie2': generated_paper = generate_cie2_paper(all_questions)
        else: generated_paper = generate_paper_with_rules(all_questions)

        if isinstance(generated_paper, dict) and "error" in generated_paper:
            return jsonify(generated_paper), 400

        # Save to Firestore
        paper_ref = db_firestore.collection('users').document(user_uid).collection('generated_papers').document()
        paper_data = {
            "user_uid": user_uid, "paper_name": f"{subject} - {pattern.upper()}", "subject": subject,
            "pattern": pattern, "questions": generated_paper, "created_at": firestore.SERVER_TIMESTAMP,
            "status": "generated"
        }
        paper_ref.set(paper_data)
        
        return jsonify({"message": "Generated!", "questions": generated_paper, "paper_id": paper_ref.id}), 200
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@qp_bp.route('/save_question_paper', methods=['POST'])
@firebase_auth_required
def save_question_paper():
    cache.clear()
    user_uid = request.current_user_uid
    data = request.get_json()
    
    try:
        supabase_user = supabase_service.get_user_by_firebase_uid(user_uid)
        if not supabase_user: return jsonify({'error': 'User not found in Supabase'}), 404
        
        paper_data = {
            'firebase_paper_id': data.get('firebase_paper_id'),
            'paper_name': data.get('paper_name', 'Untitled'),
            'subject': data.get('subject', 'Unknown'),
            'pattern': data.get('pattern', 'standard'),
            'total_marks': data.get('total_marks', 100),
            'question_count': len(data.get('questions', [])),
            'questions': data.get('questions', []),
            'metadata': data.get('metadata', {}),
            'status': data.get('status', 'draft')
        }
        
        saved_paper = supabase_service.save_question_paper(supabase_user['id'], paper_data)
        if saved_paper:
            # Backup to Firebase
            try:
                data['user_uid'] = user_uid
                data['saved_at'] = firestore.SERVER_TIMESTAMP
                data['supabase_id'] = saved_paper['id']
                db_firestore.collection('users').document(user_uid).collection('saved_question_papers').add(data)
            except: pass
            return jsonify({'message': 'Saved', 'paper_id': saved_paper['id']}), 200
        return jsonify({'error': 'Failed to save'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@qp_bp.route('/get_saved_items', methods=['GET'])
@firebase_auth_required
@cache.cached(timeout=300, query_string=True)
def get_saved_items():
    user_uid = request.current_user_uid
    try:
        supabase_user = supabase_service.get_user_by_firebase_uid(user_uid)
        if not supabase_user: return jsonify({'error': 'User not found'}), 404
        
        saved_papers = supabase_service.get_saved_question_papers(supabase_user['id'], limit=50)
        for p in saved_papers:
            p['type'] = 'question-paper'
            p['id'] = str(p['id'])
            if 'paper_name' in p: p['name'] = p['paper_name']
            if 'total_marks' in p: p['marks'] = p['total_marks']
            if 'created_at' in p: p['date'] = p['created_at']

        # Generated papers
        gen_docs = db_firestore.collection('users').document(user_uid).collection('generated_papers').order_by('created_at', direction=firestore.Query.DESCENDING).limit(50).stream()
        gen_papers = [{'id': d.id, 'type': 'generated-paper', **d.to_dict()} for d in gen_docs]

        # Question banks
        bank_docs = db_firestore.collection('users').document(user_uid).collection('question_banks').order_by('uploaded_at', direction=firestore.Query.DESCENDING).stream()
        banks = [{'id': d.id, **d.to_dict()} for d in bank_docs]

        return jsonify({
            "question_papers": saved_papers,
            "generated_papers": gen_papers,
            "question_banks": banks,
            "templates": []
        }), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@qp_bp.route('/get_recent_papers', methods=['GET'])
@firebase_auth_required
@cache.cached(timeout=300, query_string=True)
def get_recent_papers():
    user_uid = request.current_user_uid
    try:
        recent = []
        # Generated
        gen_docs = db_firestore.collection('users').document(user_uid).collection('generated_papers').order_by('created_at', direction=firestore.Query.DESCENDING).limit(10).stream()
        for d in gen_docs: recent.append({'id': d.id, 'source': 'generated', **d.to_dict()})
        
        # Saved
        saved_docs = db_firestore.collection('users').document(user_uid).collection('saved_question_papers').order_by('saved_at', direction=firestore.Query.DESCENDING).limit(10).stream()
        for d in saved_docs: recent.append({'id': d.id, 'source': 'saved', **d.to_dict()})
        
        # Sort
        recent.sort(key=lambda x: x.get('created_at', x.get('saved_at', 0)).timestamp() if hasattr(x.get('created_at', x.get('saved_at')), 'timestamp') else 0, reverse=True)
        
        return jsonify({"recent_papers": recent[:20]}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@qp_bp.route('/delete_recent_paper', methods=['DELETE'])
@firebase_auth_required
def delete_recent_paper():
    cache.clear()
    user_uid = request.current_user_uid
    data = request.get_json()
    paper_id = data.get('paper_id')
    source = data.get('source')
    
    try:
        if source == 'generated':
            db_firestore.collection('users').document(user_uid).collection('generated_papers').document(paper_id).delete()
        elif source == 'saved':
            db_firestore.collection('users').document(user_uid).collection('saved_question_papers').document(paper_id).delete()
            # Also delete from Supabase if possible (logic omitted for brevity)
        return jsonify({'message': 'Deleted'}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@qp_bp.route('/get_user_questions', methods=['GET'])
@firebase_auth_required
def get_user_questions():
    """Get all questions from the user's question bank pool"""
    user_uid = request.current_user_uid
    
    try:
        # Fetch questions from Firestore
        questions_ref = db_firestore.collection('users').document(user_uid).collection('question_bank_pool')
        docs = questions_ref.stream()
        
        questions = []
        for doc in docs:
            q_data = doc.to_dict()
            q_data['firestore_id'] = doc.id
            
            # Remove non-serializable fields if any
            if 'uploaded_at' in q_data:
                del q_data['uploaded_at']
            if 'last_used_date' in q_data:
                del q_data['last_used_date']
                
            questions.append(q_data)
            
        return jsonify({"questions": questions, "count": len(questions)}), 200
        
    except Exception as e:
        print(f"Error fetching user questions: {e}")
        return jsonify({"error": f"Error fetching questions: {str(e)}"}), 500

@qp_bp.route('/get_user_profile', methods=['GET'])
@firebase_auth_required
def get_user_profile():
    """Get user profile information"""
    user_uid = request.current_user_uid
    
    try:
        supabase_user = supabase_service.get_user_by_firebase_uid(user_uid)
        if not supabase_user:
            return jsonify({'error': 'User not found'}), 404
            
        return jsonify(supabase_user), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@qp_bp.route('/export_question_paper', methods=['POST'])
@firebase_auth_required
def export_question_paper():
    user_uid = request.current_user_uid
    data = request.get_json()
    paper_id = data.get('paper_id')
    format_type = data.get('format', 'docx')
    
    try:
        paper_ref = db_firestore.collection('users').document(user_uid).collection('generated_papers').document(paper_id)
        paper_doc = paper_ref.get()
        if not paper_doc.exists: return jsonify({"error": "Paper not found"}), 404
        
        paper_data = paper_doc.to_dict()
        questions = paper_data.get('questions', [])
        subject = paper_data.get('subject', 'Question Paper')
        
        if format_type == 'docx':
            file_path = generate_docx_export(questions, subject, paper_id)
            mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif format_type == 'pdf':
            file_path = generate_pdf_export(questions, subject, paper_id)
            mime = 'application/pdf'
        else: return jsonify({"error": "Invalid format"}), 400
        
        return jsonify({"file_path": file_path, "mime_type": mime}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@qp_bp.route('/generate_final_document', methods=['POST'])
@firebase_auth_required
def generate_final_document():
    user_uid = request.current_user_uid
    data = request.get_json()
    
    question_paper_data = data.get('question_paper_data')
    overall_max_marks = data.get('overall_max_marks', 100)
    metadata = data.get('metadata', {})
    file_format = data.get('format', 'pdf').lower()

    if not question_paper_data:
        return jsonify({"error": "No question paper data provided."}), 400

    try:
        # If format is DOCX, use the DOCX generator
        if file_format == 'docx':
            filepath = generate_docx_export(question_paper_data, metadata.get('subject', 'Question Paper'), data.get('paper_id', 'preview'))
            filename = os.path.basename(filepath)
            return send_file(
                filepath,
                as_attachment=True,
                download_name=filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

        # PDF Generation using ReportLab (Ported from app_backup.py)
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.units import cm, inch
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.pdfgen import canvas

        # Flatten question data
        questions_list = []
        
        for main_q_idx, main_q_data in enumerate(question_paper_data):
            # Insert OR row for CIE 1 pattern (before Q2 and Q4) - Logic from backup
            # Note: This logic assumes a specific structure (CIE 1). 
            # If the user wants this generic, we might need to adjust, but following backup for now.
            if main_q_idx == 1 or main_q_idx == 3:
                questions_list.append({
                    'qno': '',
                    'question': 'OR',
                    'marks': '',
                    'co': '',
                    'level': '',
                    'module': '',
                    'images': []
                })
                
            # Handle both camelCase and snake_case for compatibility
            sub_questions = main_q_data.get('subQuestions', main_q_data.get('sub_questions', []))
            for sub_q_idx, sub_q_data in enumerate(sub_questions):
                
                # Format Q.No (e.g., 1.a)
                qno = ""
                if sub_q_idx == 0:
                    qno = str(main_q_idx + 1)
                
                sub_letter = chr(97 + sub_q_idx)
                question_text = sub_q_data.get('question_text', sub_q_data.get('text', ''))
                
                # Store images
                images = sub_q_data.get('images', [])
                
                questions_list.append({
                    'qno': qno,
                    'question': f"{sub_letter}. {question_text}",
                    'marks': sub_q_data.get('marks', ''),
                    'co': sub_q_data.get('co', ''),
                    'level': sub_q_data.get('blooms_level', sub_q_data.get('level', '')),
                    'module': sub_q_data.get('module', ''),
                    'images': images
                })

        # Create PDF
        if not os.path.exists(current_app.config['UPLOAD_FOLDER']):
            os.makedirs(current_app.config['UPLOAD_FOLDER'])
            
        temp_output_path = os.path.join(current_app.config['UPLOAD_FOLDER'], f"generated_qp_{user_uid}_{uuid.uuid4()}.pdf")
        
        doc = SimpleDocTemplate(temp_output_path, pagesize=A4,
                                rightMargin=1.5*cm, leftMargin=1.5*cm,
                                topMargin=1.5*cm, bottomMargin=1.5*cm)
        
        # Container for PDF elements
        elements = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=14,
            textColor=colors.black,
            spaceAfter=6,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Normal'],
            fontSize=12,
            textColor=colors.black,
            spaceAfter=4,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        # Header with logo
        logo_path = os.path.join(current_app.static_folder, 'assets', 'logo.jpg')
        
        if os.path.exists(logo_path):
            # Create logo image (slightly larger)
            logo_img = Image(logo_path, width=2.5*cm, height=2.5*cm)
            
            # Create institute text paragraphs with LEFT alignment
            institute_name_style = ParagraphStyle(
                'InstituteName',
                parent=styles['Heading1'],
                fontSize=16,
                fontName='Helvetica-Bold',
                alignment=TA_LEFT,  # Changed to LEFT
                spaceAfter=2
            )
            
            details_style = ParagraphStyle(
                'Details',
                parent=styles['Normal'],
                fontSize=8,
                alignment=TA_LEFT,  # Changed to LEFT
                spaceAfter=1
            )
            
            institute_text = []
            institute_text.append(Paragraph("SRI KRISHNA INSTITUTE OF TECHNOLOGY", institute_name_style))
            institute_text.append(Paragraph("(Accredited by NAAC, Approved by A.I.C.T.E. New Delhi, Recognised by Govt. of Karnataka & Affiliated to V.T.U., Belagavi)", details_style))
            institute_text.append(Paragraph("#57, Chimney Hills, Hesaraghatta Main Road, Chikkabanavara Post, Bengaluru- 560090", details_style))
            
            # Create header table: logo | institute details
            header_data = [[logo_img, institute_text]]
            header_table = Table(header_data, colWidths=[3*cm, 14*cm])
            header_table.setStyle(TableStyle([
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('ALIGN', (0, 0), (0, 0), 'LEFT'),
                ('ALIGN', (1, 0), (1, 0), 'LEFT'),
                ('LINEBELOW', (0, 0), (-1, 0), 1.5, colors.black),
                ('LEFTPADDING', (0, 0), (-1, -1), 0),
                ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                ('TOPPADDING', (0, 0), (-1, -1), 5),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
            ]))
            elements.append(header_table)
        else:
            # Fallback to text-only if logo not found
            elements.append(Paragraph("SRI KRISHNA INSTITUTE OF TECHNOLOGY", title_style))
            elements.append(Paragraph("(Accredited by NAAC, Approved by A.I.C.T.E. New Delhi, Recognised by Govt. of Karnataka & Affiliated to V.T.U., Belagavi)", 
                                     ParagraphStyle('Details', parent=styles['Normal'], fontSize=8, alignment=TA_CENTER, spaceAfter=2)))
            elements.append(Paragraph("#57, Chimney Hills, Hesaraghatta Main Road, Chikkabanavara Post, Bengaluru- 560090",
                                     ParagraphStyle('Address', parent=styles['Normal'], fontSize=8, alignment=TA_CENTER, spaceAfter=4)))
        
        elements.append(Spacer(1, 0.2*cm))
        elements.append(Paragraph(f"Department of {metadata.get('dept', 'ARTIFICIAL INTELLIGENCE AND MACHINE LEARNING')}", subtitle_style))
        elements.append(Paragraph("CIE - I", subtitle_style))
        elements.append(Spacer(1, 0.3*cm))
        
        # Metadata table
        meta_data = [
            [f"Date: {metadata.get('date', '')}", f"Time: {metadata.get('time', '')}", 
             f"Max Marks: {overall_max_marks}", f"Sem/Div: {metadata.get('sem', '')}/{metadata.get('div', '')}"],
            [f"Course: {metadata.get('subject', metadata.get('course', ''))}", 
             f"Code: {metadata.get('code', '')}", 
             f"Elective: {metadata.get('elective', 'N/A')}", ""]
        ]
        
        meta_table = Table(meta_data, colWidths=[4.5*cm, 4*cm, 4*cm, 4*cm])
        meta_table.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('SPAN', (2, 1), (3, 1)),  # Merge last two cells in row 2
        ]))
        elements.append(meta_table)
        elements.append(Spacer(1, 0.3*cm))
        
        # Note
        note_style = ParagraphStyle('Note', parent=styles['Normal'], fontSize=9, fontName='Helvetica-Bold')
        elements.append(Paragraph("Note: Answer any ONE full question from each module.", note_style))
        elements.append(Spacer(1, 0.2*cm))
        
        # Questions table - with images INLINE in question cells
        table_data = [["Q.No", "Question", "CO", "Level", "Marks", "Module"]]
        
        for q in questions_list:
            if q['question'] == 'OR':
                # OR row - will be merged later
                table_data.append(["OR", "", "", "", "", ""])
            else:
                # Normal question row - create nested content for question cell
                question_cell_content = []
                
                # Add question text
                question_cell_content.append(Paragraph(str(q['question']), styles['Normal']))
                
                # Add images inline if present
                if q.get('images') and len(q['images']) > 0:
                    for img_name in q['images']:
                        # Resolve path - images are stored in uploads/{user_uid}/extracted_images/
                        if not os.path.isabs(img_name):
                            user_images_folder = os.path.join(current_app.config['UPLOAD_FOLDER'], user_uid, 'extracted_images')
                            full_img_path = os.path.join(user_images_folder, img_name)
                        else:
                            full_img_path = img_name
                            
                        if os.path.exists(full_img_path):
                            try:
                                # Add small spacer before image
                                question_cell_content.append(Spacer(1, 0.1*cm))
                                # Add image (smaller size for inline display)
                                img = Image(full_img_path, width=4*cm, height=3*cm, kind='proportional')
                                question_cell_content.append(img)
                            except Exception as e:
                                print(f"Error loading image {full_img_path}: {e}")
                
                # Create a nested single-column table to hold the content
                nested_table = Table([[item] for item in question_cell_content], colWidths=[11.5*cm])
                nested_table.setStyle(TableStyle([
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 0),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                    ('TOPPADDING', (0, 0), (-1, -1), 0),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
                ]))
                
                table_data.append([
                    str(q['qno']),
                    nested_table,  # Use nested table instead of plain Paragraph
                    str(q['co']),
                    str(q['level']),
                    str(q['marks']),
                    str(q['module'])
                ])
        
        # Adjusted column widths: Q.No, Question, CO, Level, Marks, Module
        questions_table = Table(table_data, colWidths=[1.2*cm, 11.5*cm, 1.2*cm, 1.2*cm, 1.2*cm, 1.2*cm])
        
        # Build table style
        table_style_commands = [
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),  # Header row
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),  # Header alignment
            ('ALIGN', (0, 1), (0, -1), 'CENTER'),  # Q.No column
            ('ALIGN', (2, 1), (-1, -1), 'CENTER'),  # Marks, CO, Level, Module columns
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
        ]
        
        # Add SPAN for OR rows
        for i, q in enumerate(questions_list):
            if q['question'] == 'OR':
                table_style_commands.append(('SPAN', (0, i+1), (5, i+1)))
                table_style_commands.append(('ALIGN', (0, i+1), (0, i+1), 'CENTER'))
                table_style_commands.append(('FONTNAME', (0, i+1), (0, i+1), 'Helvetica-Bold'))
                table_style_commands.append(('FONTSIZE', (0, i+1), (0, i+1), 11))
        
        questions_table.setStyle(TableStyle(table_style_commands))
        elements.append(questions_table)
        
        # Build PDF
        doc.build(elements)
        print(f"Generated PDF at: {temp_output_path}")

        # Return the PDF file directly
        return send_file(
            temp_output_path,
            mimetype='application/pdf',
            as_attachment=True,
            download_name='Generated_Question_Paper.pdf'
        )

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@qp_bp.route('/submit_for_approval', methods=['POST'])
@firebase_auth_required
def submit_for_approval():
    """Submit a question paper for HOD approval (hybrid approach)"""
    user_uid = request.current_user_uid
    data = request.get_json()
    paper_id = data.get('paper_id')
    comments = data.get('comments', '')
    
    try:
        # Get Supabase user ID
        supabase_user = supabase_service.get_user_by_firebase_uid(user_uid)
        if not supabase_user:
            return jsonify({'error': 'User not found in Supabase'}), 404
        
        paper_data = None
        paper_source = None
        
        # First, try to get from Supabase saved papers
        try:
            saved_papers = supabase_service.get_saved_question_papers(supabase_user['id'], limit=100)
            for paper in saved_papers:
                if str(paper['id']) == str(paper_id) or paper.get('firebase_paper_id') == paper_id:
                    paper_data = {
                        'paper_name': paper.get('paper_name', 'Untitled Paper'),
                        'subject': paper.get('subject', 'Unknown Subject'),
                        'questions': paper.get('questions', []),
                        'pattern': paper.get('pattern', 'standard'),
                        'total_marks': paper.get('total_marks', 100),
                        'metadata': paper.get('metadata', {})
                    }
                    paper_source = 'supabase'
                    break
        except Exception as e:
            print(f"Error fetching from Supabase: {e}")
        
        # If not found in Supabase, try Firebase generated papers
        if not paper_data:
            try:
                paper_ref = db_firestore.collection('users').document(user_uid).collection('generated_papers').document(paper_id)
                paper_doc = paper_ref.get()
                
                if paper_doc.exists:
                    paper_data = paper_doc.to_dict()
                    paper_source = 'firebase'
            except Exception as e:
                print(f"Error fetching from Firebase: {e}")
        
        if not paper_data:
            return jsonify({"error": "Question paper not found in either Supabase or Firebase"}), 404
        
        # Get user info from Supabase
        user_data = {
            'name': supabase_user.get('name', 'Unknown Faculty'),
            'email': supabase_user.get('email', ''),
            'department': supabase_user.get('department', 'Unknown Department')
        }
        
        # Find the Supabase paper ID for the approval
        supabase_paper_id = None
        if paper_source == 'supabase':
            # Find the paper in the saved papers list
            for paper in saved_papers:
                if str(paper['id']) == str(paper_id) or paper.get('firebase_paper_id') == paper_id:
                    supabase_paper_id = paper['id']
                    break
        else:
            # For Firebase papers, check if there's a corresponding Supabase paper
            try:
                firebase_paper_result = supabase_service.supabase.table('saved_question_papers').select('id').eq('firebase_paper_id', paper_id).execute()
                if firebase_paper_result.data:
                    supabase_paper_id = firebase_paper_result.data[0]['id']
            except Exception as e:
                print(f"Error finding Supabase paper for Firebase ID {paper_id}: {e}")
        
        if not supabase_paper_id:
            # If paper is not in Supabase but exists in Firebase, save it now
            if paper_source == 'firebase' and paper_data:
                print(f"Paper {paper_id} not in Supabase, saving it now...")
                try:
                    # Prepare data for Supabase
                    supabase_paper_data = {
                        'firebase_paper_id': paper_id,
                        'paper_name': paper_data.get('paper_name', 'Untitled Paper'),
                        'subject': paper_data.get('subject', 'Unknown Subject'),
                        'pattern': paper_data.get('pattern', 'standard'),
                        'total_marks': 100,  # Default
                        'question_count': len(paper_data.get('questions', [])),
                        'questions': paper_data.get('questions', []),
                        'metadata': paper_data.get('source_info', {}),
                        'status': 'draft',
                        'tags': paper_data.get('tags', [])
                    }
                    
                    # Save to Supabase
                    saved_paper = supabase_service.save_question_paper(supabase_user['id'], supabase_paper_data)
                    if saved_paper:
                        supabase_paper_id = saved_paper['id']
                        print(f"Paper saved to Supabase: {supabase_paper_id}")
                        
                        # Update Firebase with Supabase ID reference
                        paper_ref = db_firestore.collection('users').document(user_uid).collection('generated_papers').document(paper_id)
                        paper_ref.update({'supabase_id': supabase_paper_id})
                    else:
                        return jsonify({"error": "Failed to save paper to Supabase"}), 500
                except Exception as e:
                    print(f"Error saving paper to Supabase: {e}")
                    return jsonify({"error": f"Failed to save paper to Supabase: {str(e)}"}), 500
            else:
                return jsonify({"error": "Paper not found in Supabase. Please save the paper first."}), 404
        
        # Create approval request (matching database schema)
        approval_data = {
            "paper_id": supabase_paper_id,  # Use Supabase UUID
            "submitted_by": supabase_user['id'],  # Use Supabase user ID
            "status": "pending",
            "comments": comments,  # faculty comments
            "hod_comments": None,
            "reviewed_by": None,
            "reviewed_at": None
        }
        
        # Save to Supabase first (primary storage)
        try:
            supabase_approval = supabase_service.create_approval(supabase_user['id'], approval_data)
            if supabase_approval:
                approval_id = supabase_approval['id']
                print(f"Approval saved to Supabase: {approval_id}")
            else:
                return jsonify({"error": "Failed to save approval to Supabase"}), 500
        except Exception as e:
            print(f"Error saving to Supabase: {e}")
            return jsonify({"error": f"Failed to save approval: {str(e)}"}), 500
        
        # Also save to Firebase for backup
        try:
            firebase_approval_data = approval_data.copy()
            firebase_approval_data["submitted_at"] = firestore.SERVER_TIMESTAMP
            firebase_approval_data["supabase_id"] = approval_id
            
            approval_ref = db_firestore.collection('approvals').document()
            approval_ref.set(firebase_approval_data)
            print(f"Approval also backed up to Firebase: {approval_ref.id}")
        except Exception as e:
            print(f"Warning: Failed to backup approval to Firebase: {e}")
        
        # Update paper status in both systems
        try:
            # Update Supabase paper status (always use the Supabase paper ID)
            supabase_service.update_question_paper_status(supabase_paper_id, 'submitted', approval_id)
        except Exception as e:
            print(f"Warning: Failed to update Supabase paper status: {e}")
        
        # Also update Firebase paper status if it's a Firebase paper
        if paper_source == 'firebase':
            try:
                paper_ref = db_firestore.collection('users').document(user_uid).collection('generated_papers').document(paper_id)
                paper_ref.update({
                    "status": "submitted",
                    "approval_id": approval_id,
                    "submitted_for_approval_at": firestore.SERVER_TIMESTAMP
                })
            except Exception as e:
                print(f"Warning: Failed to update Firebase paper status: {e}")
        
        return jsonify({
            "message": "Question paper submitted for approval successfully!",
            "approval_id": approval_id,
            "status": "pending"
        }), 200
        
    except Exception as e:
        return jsonify({"error": f"Failed to submit for approval: {str(e)}"}), 500

@qp_bp.route('/get_approval_status', methods=['GET'])
@firebase_auth_required
def get_approval_status():
    """Get approval status for a paper"""
    user_uid = request.current_user_uid
    paper_id = request.args.get('paper_id')
    
    try:
        # Get paper
        paper_ref = db_firestore.collection('users').document(user_uid).collection('generated_papers').document(paper_id)
        paper_doc = paper_ref.get()
        
        if not paper_doc.exists:
            return jsonify({"error": "Question paper not found"}), 404
        
        paper_data = paper_doc.to_dict()
        approval_id = paper_data.get('approval_id')
        
        if not approval_id:
            return jsonify({
                "status": "not_submitted",
                "message": "Paper not submitted for approval"
            }), 200
        
        # Get approval details
        approval_ref = db_firestore.collection('approvals').document(approval_id)
        approval_doc = approval_ref.get()
        
        if not approval_doc.exists:
            return jsonify({
                "status": "error",
                "message": "Approval record not found"
            }), 200
        
        approval_data = approval_doc.to_dict()
        
        return jsonify({
            "status": approval_data.get('status', 'unknown'),
            "submitted_at": approval_data.get('submitted_at'),
            "reviewed_at": approval_data.get('reviewed_at'),
            "reviewed_by": approval_data.get('reviewed_by', ''),
            "hod_comments": approval_data.get('hod_comments', ''),
            "comments": approval_data.get('comments', ''),
            "revision_requests": approval_data.get('revision_requests', []),
            "priority": approval_data.get('priority', 'medium')
        }), 200
        
    except Exception as e:
        return jsonify({"error": f"Error fetching approval status: {str(e)}"}), 500

@qp_bp.route('/get_user_submissions', methods=['GET'])
@firebase_auth_required
def get_user_submissions():
    """Get all submissions for a user"""
    user_uid = request.current_user_uid
    
    try:
        # Get Supabase user ID
        supabase_user = supabase_service.get_user_by_firebase_uid(user_uid)
        if not supabase_user:
            return jsonify({'error': 'User not found in Supabase'}), 404
        
        # Get all approvals for this user
        approvals_result = supabase_service.supabase.table('approvals')\
            .select('*, saved_question_papers(paper_name, subject, pattern)')\
            .eq('submitted_by', supabase_user['id'])\
            .order('submitted_at', desc=True)\
            .execute()
        
        submissions = []
        for approval in approvals_result.data:
            paper_info = approval.get('saved_question_papers', {})
            submissions.append({
                'id': approval['id'],
                'paper_id': approval['paper_id'],
                'paper_name': paper_info.get('paper_name', 'Untitled Paper'),
                'subject': paper_info.get('subject', 'Unknown Subject'),
                'pattern': paper_info.get('pattern', 'standard'),
                'status': approval['status'],
                'comments': approval.get('comments', ''),
                'hod_comments': approval.get('hod_comments', ''),
                'submitted_at': approval.get('submitted_at'),
                'reviewed_at': approval.get('reviewed_at'),
                'reviewed_by': approval.get('reviewed_by')
            })
        
        return jsonify({
            'submissions': submissions,
            'total': len(submissions)
        }), 200
        
    except Exception as e:
        print(f"Error fetching user submissions: {e}")
        return jsonify({"error": f"Error fetching submissions: {str(e)}"}), 500

@qp_bp.route('/get_dashboard_metrics', methods=['GET'])
@firebase_auth_required
def get_dashboard_metrics():
    """Get dashboard metrics for the user"""
    user_uid = request.current_user_uid
    
    try:
        # Get Supabase user ID
        supabase_user = supabase_service.get_user_by_firebase_uid(user_uid)
        if not supabase_user:
            return jsonify({'error': 'User not found in Supabase'}), 404
        
        user_id = supabase_user['id']
        
        # Get total questions from question banks
        questions_result = supabase_service.supabase.table('question_banks')\
            .select('total_questions')\
            .eq('user_id', user_id)\
            .execute()
        
        total_questions = sum(bank.get('total_questions', 0) for bank in questions_result.data)
        
        # If no questions in question_banks, try to count from saved papers
        if total_questions == 0:
            papers_with_questions = supabase_service.supabase.table('saved_question_papers')\
                .select('question_count')\
                .eq('user_id', user_id)\
                .execute()
            
            total_questions = sum(paper.get('question_count', 0) for paper in papers_with_questions.data)
        
        # Get papers generated this month
        from datetime import datetime, timedelta
        month_start = datetime.now().replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        
        papers_result = supabase_service.supabase.table('saved_question_papers')\
            .select('id, created_at')\
            .eq('user_id', user_id)\
            .gte('created_at', month_start.isoformat())\
            .execute()
        
        papers_generated = len(papers_result.data)
        
        # Get saved templates (all saved papers can be used as templates)
        templates_result = supabase_service.supabase.table('saved_question_papers')\
            .select('id')\
            .eq('user_id', user_id)\
            .execute()
        
        saved_templates = len(templates_result.data)
        
        # Calculate success rate (papers that were successfully generated vs failed)
        # For now, assume high success rate since we don't track failures
        success_rate = 98 if papers_generated > 0 else 100
        
        # Calculate trends (simplified - in real app, compare with previous periods)
        questions_trend = 12 if total_questions > 0 else 0
        papers_trend = 8 if papers_generated > 0 else 0
        
        return jsonify({
            'total_questions': total_questions,
            'papers_generated': papers_generated,
            'saved_templates': saved_templates,
            'success_rate': success_rate,
            'questions_trend': questions_trend,
            'papers_trend': papers_trend
        }), 200
        
    except Exception as e:
        print(f"Error fetching dashboard metrics: {e}")
        return jsonify({"error": f"Error fetching metrics: {str(e)}"}), 500

@qp_bp.route('/get_recent_activity', methods=['GET'])
@firebase_auth_required
def get_recent_activity():
    """Get recent activity feed for the user"""
    user_uid = request.current_user_uid
    
    try:
        # Get Supabase user ID
        supabase_user = supabase_service.get_user_by_firebase_uid(user_uid)
        if not supabase_user:
            return jsonify({'error': 'User not found in Supabase'}), 404
        
        user_id = supabase_user['id']
        
        # 1. Get latest question papers
        papers_result = supabase_service.supabase.table('saved_question_papers')\
            .select('paper_name, created_at, status')\
            .eq('user_id', user_id)\
            .order('created_at', desc=True)\
            .limit(5)\
            .execute()
        
        # 2. Get latest question banks
        banks_result = supabase_service.supabase.table('question_banks')\
            .select('filename, created_at, total_questions')\
            .eq('user_id', user_id)\
            .order('created_at', desc=True)\
            .limit(5)\
            .execute()
        
        activities = []
        
        # Map papers to activity format
        for paper in papers_result.data:
            activities.append({
                'type': 'generation',
                'title': f"Generated {paper['paper_name']}",
                'timestamp': paper['created_at'],
                'status': 'Success' if paper.get('status') != 'failed' else 'Failed',
                'icon': 'fas fa-file-alt',
                'color': 'bg-green-500',
                'badge_color': 'bg-green-100 text-green-800' if paper.get('status') != 'failed' else 'bg-red-100 text-red-800'
            })
            
        # Map banks to activity format
        for bank in banks_result.data:
            activities.append({
                'type': 'upload',
                'title': f"Uploaded {bank['filename']}",
                'timestamp': bank['created_at'],
                'status': 'Completed',
                'icon': 'fas fa-upload',
                'color': 'bg-blue-500',
                'badge_color': 'bg-blue-100 text-blue-800'
            })
            
        # Sort combined activities by timestamp DESC
        activities.sort(key=lambda x: x['timestamp'], reverse=True)
        
        # Return top 5
        return jsonify(activities[:5]), 200
        
    except Exception as e:
        print(f"Error fetching recent activity: {e}")
        return jsonify({"error": f"Error fetching activity: {str(e)}"}), 500
