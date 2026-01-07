import os
import uuid
import random
from docx import Document
from docxtpl import DocxTemplate
from werkzeug.utils import secure_filename
import pdfplumber
from flask import session
from collections import Counter

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
GENERATED_FOLDER = os.path.join(BASE_DIR, 'generated')
TEMPLATE_PATH = os.path.join(BASE_DIR, 'qp_template.docx')


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ['docx', 'pdf']


def parse_docx(path):
    document = Document(path)
    questions = []
    current_module = ""

    for para in document.paragraphs:
        if 'Module' in para.text:
            current_module = para.text.strip()

    for table in document.tables:
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        for row in table.rows[1:]:
            q = {}
            for i, cell in enumerate(row.cells):
                header = headers[i]
                q[header] = cell.text.strip()
            q['module'] = current_module
            questions.append(q)

    print(f"✅ Parsed {len(questions)} questions from DOCX")
    return questions


def parse_pdf(path):
    questions = []
    current_module = "Module-1"
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in tables:
                headers = [h.strip().lower() for h in table[0]]
                for row in table[1:]:
                    if len(row) < len(headers):
                        continue
                    q = {}
                    for i in range(len(headers)):
                        q[headers[i]] = row[i].strip() if row[i] else ""
                    q['module'] = current_module
                    questions.append(q)
    print(f"✅ Parsed {len(questions)} questions from PDF")
    return questions


def filter_questions(questions, levels, modules):
    levels = [l.lower() for l in levels]
    modules = [m.lower() for m in modules]
    filtered = [
        q for q in questions
        if str(q.get('level', '')).lower() in levels and str(q.get('module', '')).lower() in modules
    ]
    print(f"✅ Filtered {len(filtered)} questions after level+module filter")
    return filtered


def group_questions_by_marks(questions, target=25):
    # Clean and validate marks
    clean_questions = []
    for q in questions:
        try:
            q['marks'] = int(str(q.get('marks', '')).strip())
            clean_questions.append(q)
        except:
            print(f"⚠️ Skipping question with invalid marks: {q}")
    questions = clean_questions

    print(f"✅ Cleaned questions: {len(questions)}")
    print("📊 Mark distribution:", Counter([q['marks'] for q in questions]))

    used = set()
    valid_groups = []
    fallback_groups = []

    def group_score(i, j, k):
        return abs((questions[i]['marks'] + questions[j]['marks'] + questions[k]['marks']) - target)

    n = len(questions)
    for i in range(n):
        for j in range(i+1, n):
            for k in range(j+1, n):
                if i in used or j in used or k in used:
                    continue
                total = questions[i]['marks'] + questions[j]['marks'] + questions[k]['marks']
                group = [questions[i], questions[j], questions[k]]
                if total == target:
                    valid_groups.append(group)
                    used.update([i, j, k])
                else:
                    fallback_groups.append((group_score(i, j, k), group))

                if len(valid_groups) == 4:
                    break
            if len(valid_groups) == 4:
                break
        if len(valid_groups) == 4:
            break

    while len(valid_groups) < 4 and fallback_groups:
        fallback_groups.sort(key=lambda x: x[0])  # sort by closest to 25
        _, best_group = fallback_groups.pop(0)
        valid_groups.append(best_group)

    if len(valid_groups) == 0:
        raise ValueError("Could not form any valid question groups that total 25 marks.")

    final = []
    for qn, group in enumerate(valid_groups):
        for idx, q in enumerate(group):
            q_copy = q.copy()
            q_copy['qno'] = f"Q{qn+1} ({chr(97+idx)})"
            final.append(q_copy)

    return final


def generate_docx(questions, metadata, output_name):
    tpl = DocxTemplate(TEMPLATE_PATH)
    context = {
        'metadata': metadata,
        'questions': questions
    }
    output_path = os.path.join(GENERATED_FOLDER, output_name + '.docx')
    tpl.render(context)
    tpl.save(output_path)
    return output_path


def convert_to_pdf(docx_path):
    """
    Robust PDF conversion function that tries multiple methods
    """
    pdf_path = docx_path.replace('.docx', '.pdf')

    # Method 1: Try docx2pdf with COM initialization
    try:
        import pythoncom
        from docx2pdf import convert
        print("Attempting PDF conversion with docx2pdf...")

        # Initialize COM for the current thread
        pythoncom.CoInitialize()

        try:
            convert(docx_path, pdf_path)
            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                print("PDF conversion successful with docx2pdf")
                return pdf_path
        finally:
            # Always uninitialize COM
            pythoncom.CoUninitialize()

    except Exception as e:
        print(f"docx2pdf conversion failed: {e}")

    # Method 2: Try using win32com directly (Windows only)
    try:
        import pythoncom
        import win32com.client
        print("Attempting PDF conversion with win32com...")

        # Initialize COM
        pythoncom.CoInitialize()

        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False

            # Open document
            doc = word.Documents.Open(os.path.abspath(docx_path))

            # Save as PDF (format 17 is PDF)
            doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
            doc.Close()
            word.Quit()

            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                print("PDF conversion successful with win32com")
                return pdf_path

        finally:
            pythoncom.CoUninitialize()

    except Exception as e:
        print(f"win32com conversion failed: {e}")

    # Method 3: Return original DOCX if all PDF methods fail
    print("All PDF conversion methods failed, returning DOCX")
    return docx_path


def handle_upload_and_generate(file, form_data):
    if not allowed_file(file.filename):
        return None, "File type not supported. Please upload DOCX or PDF."

    filename = secure_filename(file.filename)
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    file.save(filepath)

    # Try parsing
    if filename.endswith('.docx'):
        questions = parse_docx(filepath)
    else:
        questions = parse_pdf(filepath)

    if not questions:
        return None, "No questions found in the file. Check formatting."

    selected_levels = form_data.getlist('levels')
    selected_modules = form_data.getlist('modules')
    filtered_questions = filter_questions(questions, selected_levels, selected_modules)

    if not filtered_questions:
        # Save state so user doesn’t need to refill
        session['metadata'] = dict(form_data)
        session['raw_questions'] = questions
        return None, "No questions matched your selected filters. Try different level/module."

    try:
        grouped_questions = group_questions_by_marks(filtered_questions)
    except ValueError as e:
        session['metadata'] = dict(form_data)
        session['raw_questions'] = filtered_questions
        return None, f"Error during grouping: {str(e)}"

    metadata = {
        'dept': form_data.get('dept'),
        'sem': form_data.get('sem'),
        'course': form_data.get('course'),
        'date': form_data.get('date'),
        'time': form_data.get('time'),
        'code': form_data.get('code'),
        'elective': form_data.get('elective'),
        'max_marks': form_data.get('max_marks')
    }

    output_id = str(uuid.uuid4())
    docx_path = generate_docx(grouped_questions, metadata, output_id)

    # Store state for preview/finalize
    session['metadata'] = metadata
    session['questions'] = grouped_questions
    session['raw_questions'] = filtered_questions
    session['output_id'] = output_id

    if form_data.get('format') == 'pdf':
        convert_to_pdf(docx_path)
        return output_id + '.pdf', None
    else:
        return output_id + '.docx', None
