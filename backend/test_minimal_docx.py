import docx
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Test minimal document generation
doc = docx.Document()

# Add simple text
p = doc.add_paragraph("Test Document")
p.runs[0].bold = True

# Add a table
table = doc.add_table(rows=2, cols=3)
table.style = 'Table Grid'

# Header
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Q.No"
hdr_cells[1].text = "Question"
hdr_cells[2].text = "Marks"

# Data row
row_cells = table.rows[1].cells
row_cells[0].text = "1"
row_cells[1].text = "What is AI?"
row_cells[2].text = "10"

# Save
doc.save("test_minimal.docx")
print("Minimal test document created successfully")
