import os
import sys
from docx import Document
from docx.shared import Pt

base_dir = os.path.dirname(os.path.abspath(__file__))
template_path = os.path.join(base_dir, 'qp_template.docx')
output_path = os.path.join(base_dir, 'qp_template_fixed.docx')

if not os.path.exists(template_path):
    print("Template not found!")
    sys.exit(1)

doc = Document(template_path)

# Find the questions table
target_table = None
for table in doc.tables:
    if len(table.rows) > 0:
        row_text = [cell.text.lower() for cell in table.rows[0].cells]
        if any('question' in t for t in row_text):
            target_table = table
            break

if target_table:
    print("Found questions table.")
    
    # We expect headers: Q.No, Questions, CO, Level, Marks, Module
    # We will keep the header row (0) and remove all other rows
    # Then add a single row with Jinja2 tags
    
    # Remove existing data rows
    for i in range(len(target_table.rows) - 1, 0, -1):
        tbl = target_table._tbl
        tr = target_table.rows[i]._tr
        tbl.remove(tr)
        
    # Add the template row
    row = target_table.add_row()
    
    # Define tags for each column
    # Assuming 6 columns based on previous code
    tags = [
        "{{ q.qno }}",
        "{{ q.question }}",
        "{{ q.co }}",
        "{{ q.level }}",
        "{{ q.marks }}",
        "{{ q.module }}"
    ]
    
    # Add the loop start tag to the first cell's paragraph
    # But docxtpl is tricky with table loops. 
    # Standard way: {% tr for q in questions %} in the first cell, 
    # and {% tr endfor %} in the last cell (or same cell if row-based).
    
    # Let's try the row-based approach which is safer for tables
    # We need to insert the loop logic into the XML or use specific docxtpl syntax.
    # Simplest docxtpl table row:
    # Cell 0: {% tr for q in questions %}{{ q.qno }}
    # ...
    # Cell N: {{ q.module }}{% tr endfor %}
    
    for i, tag in enumerate(tags):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        p.text = "" # Clear existing
        run = p.add_run()
        
        if i == 0:
            run.text = "{% tr for q in questions %}" + tag
        elif i == len(tags) - 1:
            run.text = tag + "{% tr endfor %}"
        else:
            run.text = tag
            
        # Set font
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)

    doc.save(output_path)
    print(f"Created fixed template at {output_path}")
    
    # Backup original and rename fixed
    backup_path = os.path.join(base_dir, 'qp_template_backup.docx')
    if os.path.exists(backup_path):
        os.remove(backup_path)
    os.rename(template_path, backup_path)
    os.rename(output_path, template_path)
    print("Replaced original template with fixed version.")

else:
    print("Questions table not found in template.")
